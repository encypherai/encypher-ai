"""
Test script to verify CLI docx to PDF conversion with the new parameters.
"""

import os
import subprocess
from datetime import datetime, timezone
from cryptography.hazmat.primitives.asymmetric import ed25519

# Generate a key pair for testing
private_key = ed25519.Ed25519PrivateKey.generate()
public_key = private_key.public_key()

# Save the private key to a file
private_key_path = "test_private_key.pem"
from cryptography.hazmat.primitives.serialization import Encoding, PrivateFormat, NoEncryption

with open(private_key_path, "wb") as f:
    f.write(private_key.private_bytes(
        encoding=Encoding.PEM,
        format=PrivateFormat.PKCS8,
        encryption_algorithm=NoEncryption()
    ))

# Create a simple test docx file if it doesn't exist
test_docx = "test_document.docx"
if not os.path.exists(test_docx):
    from docx import Document
    doc = Document()
    doc.add_heading('Test Document', 0)
    doc.add_paragraph('This is a test document for CLI conversion with metadata embedding.')
    doc.add_paragraph('It should have metadata embedded in the first letter of words.')
    doc.save(test_docx)
    print(f"Created test document: {test_docx}")

# Output PDF path
output_pdf = "test_cli_output.pdf"

# Current timestamp
timestamp = int(datetime.now(timezone.utc).timestamp())

# Run the CLI command with the new parameters
cmd = [
    "python", "-m", "encypher.examples.cli_example", "convert-to-pdf",
    "--docx-file", test_docx,
    "--output-file", output_pdf,
    "--private-key-file", private_key_path,
    "--signer-id", "test-signer",
    "--target", "first_letter",
    "--timestamp", str(timestamp)
]

print("Running command:", " ".join(cmd))
result = subprocess.run(cmd, capture_output=True, text=True)

# Print the result
print("Exit code:", result.returncode)
print("Output:", result.stdout)
if result.stderr:
    print("Error:", result.stderr)

# Check if the PDF was created
if os.path.exists(output_pdf):
    print(f"PDF created successfully: {output_pdf}")
    print(f"File size: {os.path.getsize(output_pdf)} bytes")
else:
    print(f"Failed to create PDF: {output_pdf}")

# Clean up
print("Cleaning up temporary files...")
if os.path.exists(private_key_path):
    os.remove(private_key_path)
    print(f"Removed: {private_key_path}")
