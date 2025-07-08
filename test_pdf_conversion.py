#!/usr/bin/env python
"""
Simple test for Word to PDF conversion with metadata embedding.
"""
import os
import sys
import tempfile
from datetime import datetime, timezone
from pathlib import Path
from cryptography.hazmat.primitives.asymmetric.ed25519 import Ed25519PrivateKey

# Add the parent directory to sys.path to ensure imports work
sys.path.insert(0, str(Path(__file__).parent))

# Use correct import paths based on project structure
from encypher.core.unicode_metadata import UnicodeMetadata
from encypher.pdf_generator import EncypherPDF

def main():
    """Test the Word to PDF conversion with metadata embedding."""
    print("Testing Word to PDF conversion with metadata embedding...")
    
    # Create a simple docx file for testing
    try:
        from docx import Document
    except ImportError:
        print("python-docx not installed. Installing...")
        import subprocess
        subprocess.check_call([sys.executable, "-m", "pip", "install", "python-docx"])
        from docx import Document
    
    # Create a test document
    doc = Document()
    doc.add_heading('Test Document with Unicode Characters', 0)
    doc.add_paragraph('This is a test document with some Unicode characters: ✓ ♥ ☺ ☻ ♦ ♣ ♠')
    doc.add_paragraph('Testing metadata embedding in PDF conversion.')
    doc.add_paragraph('The quick brown fox jumps over the lazy dog.')
    
    # Save the test document
    test_docx = os.path.join(tempfile.gettempdir(), 'test_document.docx')
    doc.save(test_docx)
    print(f"Created test document at: {test_docx}")
    
    # Generate output PDF path
    output_pdf = os.path.join(tempfile.gettempdir(), 'test_document_output.pdf')
    
    # Generate a key pair for signing and current timestamp
    private_key = Ed25519PrivateKey.generate()
    signer_id = "test-signer"
    current_timestamp = datetime.now(timezone.utc)
    
    print(f"Converting {test_docx} to PDF with embedded metadata...")
    
    try:
        # Convert docx to PDF with metadata
        pdf_path = EncypherPDF.from_docx(
            docx_file=test_docx,
            output_file=output_pdf,
            private_key=private_key,
            signer_id=signer_id,
            timestamp=current_timestamp
        )
        
        print(f"Successfully generated PDF at: {pdf_path}")
        
        # Extract and verify the text from the PDF
        print("Extracting text from the generated PDF...")
        extracted_text = EncypherPDF.extract_text(pdf_path)
        
        # Check if metadata is present in the extracted text
        metadata = UnicodeMetadata.extract_metadata(extracted_text)
        if metadata:
            print("✅ Metadata successfully embedded and extracted from PDF!")
            print(f"Metadata: {metadata}")
        else:
            print("❌ No metadata found in the extracted text.")
            print("Checking if metadata was embedded as an attachment...")
            
            # Check if the PDF has an embedded metadata file
            import fitz
            doc = fitz.open(pdf_path)
            names = doc.embfile_names()
            
            if "embedded_metadata.txt" in names:
                buffer = doc.embfile_get("embedded_metadata.txt")
                if buffer:
                    embedded_text = buffer["content"].decode("utf-8")
                    metadata = UnicodeMetadata.extract_metadata(embedded_text)
                    if metadata:
                        print("✅ Metadata found in embedded attachment!")
                        print(f"Metadata: {metadata}")
                    else:
                        print("❌ No valid metadata found in embedded attachment.")
            else:
                print("❌ No embedded metadata attachment found in PDF.")
        
        print("\nTest completed.")
        
    except Exception as e:
        print(f"Error during conversion: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()
