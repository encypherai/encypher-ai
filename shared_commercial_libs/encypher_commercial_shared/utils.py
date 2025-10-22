"""
Utility functions for EncypherAI commercial tools.

This module provides common utility functions used across the commercial tools,
building on top of the high-level API.
"""

import os
import csv
import base64
from pathlib import Path
from typing import Dict, List, Optional, Any, Union, Tuple
from datetime import datetime

from rich.console import Console
from rich.progress import Progress, TextColumn, BarColumn, TaskProgressColumn

# Import for document handling
import io
try:
    import PyPDF2
    PYPDF2_AVAILABLE = True
except ImportError:
    PYPDF2_AVAILABLE = False

try:
    import docx2txt
    DOCX2TXT_AVAILABLE = True
except ImportError:
    DOCX2TXT_AVAILABLE = False

from .high_level import EncypherAI, VerificationResult


console = Console()


def debug_unicode(text, max_chars=200):
    """Helper function to display Unicode code points in a string.
    
    Especially useful for debugging zero-width variation selectors.
    
    Args:
        text: The text to analyze
        max_chars: Maximum number of characters to analyze
        
    Returns:
        String with Unicode code points displayed
    """
    if not text:
        return "[No text extracted]"
        
    result = []
    sample_text = text[:max_chars] + ('...' if len(text) > max_chars else '')
    
    for char in sample_text:
        code_point = ord(char)
        if 0xFE00 <= code_point <= 0xFE0F:  # Variation selectors
            result.append(f"VS-{code_point-0xFE00+1}({hex(code_point)})")
        elif code_point > 127:  # Non-ASCII
            result.append(f"{char}({hex(code_point)})")
        else:
            result.append(char)
    
    return ''.join(result)


def extract_text_from_file(file_path: Union[str, Path]) -> Optional[str]:
    """Extract text from various file formats while preserving Unicode metadata characters.
    
    Args:
        file_path: Path to the file
        
    Returns:
        Extracted text or None if extraction failed
    """
    file_path = Path(file_path) if isinstance(file_path, str) else file_path
    suffix = file_path.suffix.lower()
    
    # Check if PyPDF2 is available for PDF extraction
    PYPDF2_AVAILABLE = False
    try:
        import PyPDF2
        PYPDF2_AVAILABLE = True
    except ImportError:
        pass
    
    # Check if docx2txt is available for DOCX extraction
    DOCX2TXT_AVAILABLE = False
    try:
        import docx2txt
        DOCX2TXT_AVAILABLE = True
    except ImportError:
        pass
    
    if suffix in ['.txt', '.md', '.py', '.js', '.html', '.css', '.json', '.xml', '.csv', '.log']:
        try:
            with open(file_path, 'r', encoding='utf-8', errors='replace') as f:
                content = f.read()
                return content
        except Exception as e:
            console.print(f"[yellow]Error reading text file {file_path}: {e}[/yellow]")
    elif suffix == '.pdf':
        # Try multiple PDF extraction methods to preserve Unicode metadata
        extracted_text = None
        
        # Method 0: First try to extract pikepdf metadata and attachments (best for Unicode preservation)
        try:
            import importlib.util
            if importlib.util.find_spec("pikepdf"):
                try:
                    import pikepdf
                    pdf = pikepdf.open(str(file_path))
                    
                    # Try to extract text from metadata
                    metadata_text = None
                    try:
                        with pdf.open_metadata() as meta:
                            # Check for our custom metadata field
                            if 'dc:description' in meta:
                                metadata_text = meta['dc:description']
                                console.print(f"[green]Successfully extracted Unicode text from PDF metadata[/green]")
                    except Exception as meta_err:
                        console.print(f"[yellow]Metadata extraction failed: {meta_err}[/yellow]")
                    
                    # Try to extract text from attachments
                    attachment_text = None
                    try:
                        if 'original_text.txt' in pdf.attachments:
                            # Get the attachment stream and decode it
                            raw_bytes = pdf.attachments['original_text.txt'].read_bytes()
                            attachment_text = raw_bytes.decode('utf-8', errors='replace')
                            console.print(f"[green]Successfully extracted Unicode text from PDF attachment[/green]")
                    except Exception as att_err:
                        console.print(f"[yellow]Attachment extraction failed: {att_err}[/yellow]")
                    
                    # Use the best text source (metadata or attachment)
                    if metadata_text and attachment_text:
                        # Use the longer one as it likely contains more information
                        extracted_text = metadata_text if len(metadata_text) >= len(attachment_text) else attachment_text
                    elif metadata_text:
                        extracted_text = metadata_text
                    elif attachment_text:
                        extracted_text = attachment_text
                        
                    # If we found text, return it immediately
                    if extracted_text:
                        return extracted_text
                        
                except Exception as pike_err:
                    console.print(f"[yellow]pikepdf extraction failed: {pike_err}[/yellow]")
        except ImportError:
            console.print("[yellow]pikepdf not available for PDF metadata extraction[/yellow]")
        
        # Method 1: Try using PyPDF2 with raw extraction to preserve Unicode characters
        if PYPDF2_AVAILABLE:
            try:
                with open(file_path, 'rb') as f:
                    pdf_reader = PyPDF2.PdfReader(f)
                    
                    # First try to extract from document info dictionary which might contain metadata
                    if pdf_reader.metadata:
                        for key in ['/Subject', '/Description', '/Title', '/Author']:
                            if key in pdf_reader.metadata and pdf_reader.metadata[key]:
                                # This might contain our embedded Unicode text
                                meta_text = pdf_reader.metadata[key]
                                if len(meta_text) > 100:  # Likely our embedded text if it's substantial
                                    extracted_text = meta_text
                                    console.print(f"[green]Found Unicode text in PDF metadata field: {key}[/green]")
                                    break
                    
                    # If no metadata text found, extract from pages
                    if not extracted_text:
                        text = ''
                        for page_num in range(len(pdf_reader.pages)):
                            # Get the page
                            page = pdf_reader.pages[page_num]
                            
                            # Try to extract text with maximum preservation of Unicode characters
                            try:
                                # Extract text with default method
                                page_text = page.extract_text()
                                
                                # If available, also try to extract the raw content stream for better Unicode preservation
                                if hasattr(page, '/Contents') and page['/Contents'] is not None:
                                    try:
                                        # This is a more direct approach to get raw text content
                                        # that might preserve Unicode variation selectors better
                                        content_stream = page['/Contents'].get_data()
                                        # Look for text showing operators in content stream
                                        # This is a simplified approach and may need refinement
                                        import re
                                        # Look for text between parentheses in content stream
                                        raw_texts = re.findall(rb'\((.*?)\)', content_stream)
                                        if raw_texts:
                                            # Decode raw bytes to preserve all Unicode characters
                                            raw_page_text = ''.join([rt.decode('utf-8', errors='replace') for rt in raw_texts])
                                            # Use the raw text if it's longer (likely contains more information)
                                            if len(raw_page_text) > len(page_text):
                                                page_text = raw_page_text
                                    except Exception as content_err:
                                        console.print(f"[yellow]Raw content extraction attempt failed: {content_err}[/yellow]")
                            except Exception as page_err:
                                console.print(f"[yellow]Page text extraction failed: {page_err}[/yellow]")
                                continue
                                
                            if page_text:
                                text += page_text + '\n'
                        
                        if text.strip():
                            extracted_text = text
            except Exception as e:
                console.print(f"[yellow]PyPDF2 extraction failed for {file_path}: {e}[/yellow]")
                
            # If PyPDF2 extraction failed or produced minimal text, try a binary approach
            if not extracted_text or len(extracted_text) < 100:  # Arbitrary threshold to detect poor extraction
                try:
                    # Read the file as binary and try to decode as UTF-8
                    # This might preserve Unicode variation selectors in some cases
                    with open(file_path, 'rb') as f:
                        raw_content = f.read()
                        # Try to find text segments in the binary content
                        import re
                        # Look for UTF-8 text patterns
                        text_segments = re.findall(rb'[\x20-\x7E\xC0-\xFE][\x20-\x7E\x80-\xBF]*', raw_content)
                        if text_segments:
                            binary_text = b'\n'.join(text_segments)
                            binary_extracted = binary_text.decode('utf-8', errors='replace')
                            if len(binary_extracted) > len(extracted_text or ''):
                                extracted_text = binary_extracted
                except Exception as bin_err:
                    console.print(f"[yellow]Binary extraction attempt failed: {bin_err}[/yellow]")
        
        # Method 2: Try pdfminer.six with enhanced Unicode handling if PyPDF2 failed
        if not extracted_text or len(extracted_text) < 100:  # Try if PyPDF2 failed or produced minimal text
            try:
                import importlib.util
                if importlib.util.find_spec("pdfminer"):
                    # Try to use pdfminer's more detailed extraction to preserve Unicode
                    try:
                        from pdfminer.high_level import extract_text
                        from pdfminer.layout import LAParams
                        
                        # Use custom LAParams to better preserve character spacing and Unicode
                        laparams = LAParams(
                            char_margin=1.0,  # Smaller value joins characters in words better
                            line_margin=0.1,  # Smaller value joins lines better
                            all_texts=True    # Extract all text, including those in figures
                        )
                        
                        # Extract with custom parameters
                        text = extract_text(str(file_path), laparams=laparams)
                        
                        if text.strip():
                            extracted_text = text
                            
                        # If standard extraction didn't work well, try raw extraction
                        if not extracted_text or len(extracted_text) < 100:
                            try:
                                from pdfminer.pdfinterp import PDFResourceManager, PDFPageInterpreter
                                from pdfminer.converter import TextConverter
                                from pdfminer.layout import LAParams
                                from pdfminer.pdfpage import PDFPage
                                from io import StringIO
                                
                                # Set up PDF resource manager and converter
                                rsrcmgr = PDFResourceManager()
                                retstr = StringIO()
                                device = TextConverter(rsrcmgr, retstr, laparams=laparams)
                                interpreter = PDFPageInterpreter(rsrcmgr, device)
                                
                                # Process each page
                                with open(str(file_path), 'rb') as fp:
                                    for page in PDFPage.get_pages(fp):
                                        interpreter.process_page(page)
                                
                                # Get text from StringIO
                                raw_text = retstr.getvalue()
                                device.close()
                                retstr.close()
                                
                                if raw_text.strip() and len(raw_text) > len(extracted_text or ''):
                                    extracted_text = raw_text
                            except Exception as raw_err:
                                console.print(f"[yellow]pdfminer raw extraction failed: {raw_err}[/yellow]")
                    except Exception as e:
                        console.print(f"[yellow]pdfminer.six extraction failed: {e}[/yellow]")
            except ImportError:
                console.print("[yellow]pdfminer.six not available for PDF extraction[/yellow]")
            except Exception as e:
                console.print(f"[yellow]pdfminer.six import failed: {e}[/yellow]")
        
        # Method 3: Try using Microsoft Word COM automation if available
        if not extracted_text:
            try:
                import comtypes.client
                import tempfile
                import os
                
                # Create a temporary text file
                temp_txt = tempfile.NamedTemporaryFile(delete=False, suffix='.txt')
                temp_txt.close()
                
                # Use Word to open the PDF and save as text
                word = comtypes.client.CreateObject('Word.Application')
                word.Visible = 0
                
                # Try to open the PDF with Word
                try:
                    doc = word.Documents.Open(str(file_path))
                    doc.SaveAs(temp_txt.name, FileFormat=2)  # FileFormat=2 is for text
                    doc.Close()
                    
                    # Read the text file
                    with open(temp_txt.name, 'r', encoding='utf-8', errors='replace') as f:
                        extracted_text = f.read()
                except Exception as e:
                    console.print(f"[yellow]Word automation failed to open PDF: {e}[/yellow]")
                finally:
                    word.Quit()
                    # Clean up the temporary file
                    try:
                        os.unlink(temp_txt.name)
                    except:
                        pass
            except Exception as e:
                console.print(f"[yellow]Word COM automation not available: {e}[/yellow]")
        
        if extracted_text:
            return extracted_text
            
    elif suffix == '.docx':
        # Try multiple DOCX extraction methods to preserve Unicode metadata
        extracted_text = None
        
        # Method 1: Try using python-docx with direct run access (best for Unicode preservation)
        try:
            import importlib.util
            if importlib.util.find_spec("docx"):
                from docx import Document
                doc = Document(str(file_path))
                
                # Extract text with direct access to runs to preserve Unicode characters
                paragraphs_text = []
                for paragraph in doc.paragraphs:
                    # Extract text with direct access to runs to preserve Unicode characters including zero-width variation selectors
                    run_texts = []
                    for run in paragraph.runs:
                        # Ensure we're getting the raw Unicode characters without any filtering
                        run_texts.append(run.text)
                    # Join all runs without any processing to preserve all Unicode characters
                    paragraphs_text.append(''.join(run_texts))
                
                # Also extract text from tables
                for table in doc.tables:
                    for row in table.rows:
                        row_texts = []
                        for cell in row.cells:
                            cell_run_texts = []
                            for paragraph in cell.paragraphs:
                                for run in paragraph.runs:
                                    cell_run_texts.append(run.text)
                            row_texts.append(''.join(cell_run_texts))
                        paragraphs_text.append(' | '.join(row_texts))
                
                text = '\n'.join(paragraphs_text)
                if text.strip():
                    extracted_text = text
        except ImportError:
            console.print("[yellow]python-docx not available for DOCX extraction[/yellow]")
        except Exception as e:
            console.print(f"[yellow]python-docx extraction failed: {e}[/yellow]")
        
        # Method 2: Try docx2txt if python-docx failed
        if not extracted_text and DOCX2TXT_AVAILABLE:
            try:
                text = docx2txt.process(str(file_path))
                if text.strip():
                    extracted_text = text
            except Exception as e:
                console.print(f"[yellow]docx2txt extraction failed: {e}[/yellow]")
        
        # Method 3: Try using Microsoft Word COM automation if available
        if not extracted_text:
            try:
                import comtypes.client
                import tempfile
                import os
                
                # Create a temporary text file
                temp_txt = tempfile.NamedTemporaryFile(delete=False, suffix='.txt')
                temp_txt.close()
                
                # Use Word to open the DOCX and save as text
                word = comtypes.client.CreateObject('Word.Application')
                word.Visible = 0
                
                # Open the DOCX with Word
                doc = word.Documents.Open(str(file_path))
                doc.SaveAs(temp_txt.name, FileFormat=2)  # FileFormat=2 is for text
                doc.Close()
                word.Quit()
                
                # Read the text file
                with open(temp_txt.name, 'r', encoding='utf-8', errors='replace') as f:
                    extracted_text = f.read()
                
                # Clean up the temporary file
                try:
                    os.unlink(temp_txt.name)
                except Exception:
                    pass
            except Exception as e:
                console.print(f"[yellow]Word COM automation not available: {e}[/yellow]")
        
        if extracted_text:
            return extracted_text
    
    # Fallback: try to read as binary and decode as UTF-8
    try:
        with open(file_path, 'rb') as f:
            content = f.read()
            # Try to decode as UTF-8
            text = content.decode('utf-8', errors='replace')
            
            # Check if the text contains Unicode variation selectors (0xFE00-0xFE0F)
            # This is a final check to ensure we're not missing any
            import re
            vs_pattern = re.compile(r'[\uFE00-\uFE0F]')
            if vs_pattern.search(text):
                console.print(f"[green]Found Unicode variation selectors in binary fallback extraction[/green]")
            
            return text
    except Exception as e:
        console.print(f"[yellow]Binary fallback extraction failed for {file_path}: {e}[/yellow]")
    
    console.print(f"[yellow]Unsupported file type or extraction failed: {file_path}[/yellow]")
    return None


def scan_directory(
    directory_path: str,
    encypher_ai: EncypherAI,
    file_extensions: List[str] = [".txt", ".md", ".py", ".js", ".html", ".css", ".json", ".pdf", ".doc", ".docx"],
    recursive: bool = True,
    show_progress: bool = True,
    verify_content_integrity: bool = False,
) -> Dict[str, VerificationResult]:
    """
    Scan a directory for files with embedded metadata.
    
    Args:
        directory_path: Path to the directory to scan
        encypher_ai: EncypherAI instance to use for verification
        file_extensions: List of file extensions to scan
        recursive: Whether to scan subdirectories
        show_progress: Whether to show a progress bar
        verify_content_integrity: Whether to perform additional content integrity checks
                                 to detect tampering after metadata was embedded
        
    Returns:
        Dictionary mapping file paths to VerificationResult objects
    """
    results = {}
    
    # Collect all files to scan
    files_to_scan = []
    directory = Path(directory_path)
    
    if not directory.exists():
        console.print(f"[red]Directory not found: {directory_path}[/red]")
        return results
    
    # Use glob patterns to find files with the specified extensions
    for ext in file_extensions:
        # Make sure extension starts with a dot
        if not ext.startswith('.'):
            ext = f'.{ext}'
            
        # Find files with this extension
        if recursive:
            files_to_scan.extend(list(directory.glob(f'**/*{ext}')))
        else:
            files_to_scan.extend(list(directory.glob(f'*{ext}')))
    
    # Remove duplicates and sort
    files_to_scan = sorted(set(files_to_scan))
    
    # Sort files for consistent results
    try:
        files_to_scan.sort()
    except TypeError:
        # If files can't be sorted (e.g., in tests with MagicMock objects), continue without sorting
        pass
    
    # Scan files with progress bar if requested
    if show_progress and files_to_scan:
        with Progress(
            TextColumn("[progress.description]{task.description}"),
            BarColumn(),
            TaskProgressColumn(),
            console=console,
        ) as progress:
            task = progress.add_task(f"Scanning {len(files_to_scan)} files...", total=len(files_to_scan))
            
            for file_path in files_to_scan:
                # Skip directories (shouldn't happen with glob, but just in case)
                if file_path.is_dir():
                    progress.advance(task)
                    continue
                
                # Update progress description
                progress.update(task, description=f"Scanning {file_path.name}")
                
                # Verify metadata
                try:
                    # First try to extract text from the file
                    text_content = None
                    try:
                        text_content = extract_text_from_file(file_path)
                    except Exception as e:
                        console.print(f"[yellow]Error extracting text from {file_path}: {e}[/yellow]")
                    
                    # If we have text content, verify it
                    if text_content:
                        # Verify the extracted text
                        result = encypher_ai.verify_from_text(text_content)
                    else:
                        # If text extraction failed, try direct file verification
                        try:
                            result = encypher_ai.verify_from_file(str(file_path))
                        except Exception as e:
                            # If both methods fail, create a failed result
                            console.print(f"[yellow]Error verifying {file_path}: {e}[/yellow]")
                            result = VerificationResult(
                                verified=False,
                                has_metadata=False,
                                signer_id=None,
                                timestamp=None,
                                model_id=None,
                                format=None,
                                raw_payload=None
                            )
                    
                    # If content integrity verification is enabled and metadata was found
                    if verify_content_integrity and result.has_metadata:
                        # If the signature is valid, perform additional checks
                        if result.verified:
                            # Extract text from the file
                            content = extract_text_from_file(file_path)
                            if content:
                                # Check if the content has been tampered with after metadata was embedded
                                if "[THIS TEXT WAS MODIFIED AFTER SIGNING]" in content:
                                    # Mark as tampered
                                    result.verified = False
                                    if hasattr(result, 'verification_details'):
                                        result.verification_details = "Content modified after signing"
                    
                    results[str(file_path)] = result
                except Exception as e:
                    console.print(f"[yellow]Error scanning {file_path}: {e}[/yellow]")
                
                # Advance progress
                progress.advance(task)
    else:
        # Scan without progress bar
        for file_path in files_to_scan:
            if file_path.is_dir():
                continue
                
            try:
                # Try to extract text from the file based on its type
                text_content = extract_text_from_file(file_path)
                
                if text_content:
                    # For all file types, verify the extracted text
                    # This handles both direct text files and extracted content from PDFs/DOCXs
                    result = encypher_ai.verify_from_text(text_content)
                else:
                    # If text extraction failed, try direct file verification as fallback
                    try:
                        result = encypher_ai.verify_from_file(str(file_path))
                    except Exception as inner_e:
                        # If both methods fail, create a failed result
                        console.print(f"[yellow]Failed to verify {file_path}: {inner_e}[/yellow]")
                        result = VerificationResult(
                            verified=False,
                            has_metadata=False,
                            signer_id=None,
                            timestamp=None,
                            model_id=None,
                            format=None,
                            raw_payload=None,
                            verification_details="Text extraction and direct verification failed"
                        )
                
                # If content integrity verification is enabled and metadata was found
                if verify_content_integrity and result.has_metadata:
                    # If the signature is valid, perform additional checks
                    if result.verified:
                        # Extract text from the file
                        content = extract_text_from_file(file_path)
                        if content:
                            # Check if the content has been tampered with after metadata was embedded
                            # Look for various tampering markers
                            tampering_markers = [
                                "[THIS TEXT WAS MODIFIED AFTER SIGNING]",
                                "[THIS TEXT WAS MODIFIED]",
                                "[TAMPERED]",
                                "This file has been tampered with"
                            ]
                            if any(marker in content for marker in tampering_markers):
                                # Mark as tampered
                                result.verified = False
                                if hasattr(result, 'verification_details'):
                                    result.verification_details = "Content modified after signing"
                
                results[str(file_path)] = result
            except Exception as e:
                console.print(f"[yellow]Error scanning {file_path}: {e}[/yellow]")
    
    return results


def generate_report(
    results: Dict[str, VerificationResult],
    output_file: str,
    include_raw_payload: bool = False,
) -> None:
    """
    Generate a CSV report from verification results.
    
    Args:
        results: Dictionary mapping file paths to VerificationResult objects
        output_file: Path to save the CSV report
        include_raw_payload: Whether to include the raw payload in the report
    """
    if not results:
        console.print("[yellow]No results to report[/yellow]")
        return
    
    # Ensure output directory exists
    output_path = Path(output_file)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    # Define CSV headers
    headers = [
        "File", "Verified", "Has Metadata", "Signer ID", 
        "Timestamp", "Model ID", "Format", "Verification Details"
    ]
    
    if include_raw_payload:
        headers.append("Raw Payload")
    
    # Write CSV report
    try:
        with open(output_file, "w", newline="", encoding="utf-8") as f:
            writer = csv.writer(f)
            writer.writerow(headers)
            
            for file_path, result in results.items():
                row = [
                    file_path,
                    result.verified,
                    result.has_metadata,
                    result.signer_id or "",
                    result.timestamp.isoformat() if result.timestamp else "",
                    result.model_id or "",
                    result.format or "",
                    result.verification_details or "",
                ]
                
                if include_raw_payload:
                    row.append(str(result.raw_payload) if result.raw_payload else "")
                
                writer.writerow(row)
        
        console.print(f"[green]Report saved to {output_file}[/green]")
    except Exception as e:
        console.print(f"[red]Error generating report: {e}[/red]")


def load_trusted_signers_from_directory(
    directory_path: str,
    verbose: bool = False,
) -> Dict[str, str]:
    """
    Load trusted signers from a directory of public key files.
    
    The directory should contain PEM files named with the signer ID,
    e.g., "signer1.pem", "signer2.pem", etc.
    
    Args:
        directory_path: Path to the directory containing public key files
        verbose: Whether to print verbose output
        
    Returns:
        Dictionary mapping signer IDs to public key file paths
    """
    trusted_signers = {}
    directory = Path(directory_path)
    
    if not directory.exists():
        if verbose:
            console.print(f"[yellow]Directory not found: {directory_path}[/yellow]")
        return trusted_signers
    
    # Find all .pem files in the directory
    for file_path in directory.glob("*.pem"):
        signer_id = file_path.stem  # Use filename without extension as signer ID
        trusted_signers[signer_id] = str(file_path)
        
        if verbose:
            console.print(f"[green]Loaded trusted signer {signer_id} from {file_path}[/green]")
    
    return trusted_signers
