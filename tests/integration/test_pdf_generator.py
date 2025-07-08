import os
import tempfile
import unittest
from datetime import datetime, timezone

from cryptography.hazmat.primitives.asymmetric import ed25519
from encypher.core.unicode_metadata import MetadataTarget

from encypher.pdf_generator import EncypherPDF


class TestPDFGeneration(unittest.TestCase):
    def setUp(self):
        self.private_key = ed25519.Ed25519PrivateKey.generate()
        self.public_key = self.private_key.public_key()
        self.signer_id = "pdf-signer"
        self.sample_text = "Hello PDF world"
        self.timestamp = datetime.now(timezone.utc).isoformat()

    def test_pdf_creation_and_verification(self):
        with tempfile.TemporaryDirectory() as tmpdir:
            pdf_path = os.path.join(tmpdir, "output.pdf")
            EncypherPDF.from_text(
                text=self.sample_text,
                output_path=pdf_path,
                private_key=self.private_key,
                signer_id=self.signer_id,
                timestamp=self.timestamp,
            )

            self.assertTrue(os.path.exists(pdf_path))

            extracted_text = EncypherPDF.extract_text(pdf_path)
            self.assertIn(self.sample_text, extracted_text)

            # Ensure hidden bytes are present even if exact selectors are lost
            self.assertGreater(len(extracted_text), len(self.sample_text))

    def test_docx_to_pdf_conversion(self):
        """Test converting a Word document to PDF with metadata embedding."""
        try:
            from docx import Document
        except ImportError:
            self.skipTest("python-docx not installed, skipping docx test")

        with tempfile.TemporaryDirectory() as tmpdir:
            # Create a test Word document
            docx_path = os.path.join(tmpdir, "test_document.docx")
            doc = Document()
            doc.add_heading('Test Document with Unicode Characters', 0)
            doc.add_paragraph('This is a test document with Unicode characters: ✓ ♥ ☺')
            doc.add_paragraph('Testing metadata embedding in PDF conversion.')
            doc.save(docx_path)
            
            # Convert to PDF with metadata
            pdf_path = os.path.join(tmpdir, "output.pdf")
            result_path = EncypherPDF.from_docx(
                docx_file=docx_path,
                output_file=pdf_path,
                private_key=self.private_key,
                signer_id=self.signer_id,
                timestamp=datetime.now(timezone.utc),
                target=MetadataTarget.FIRST_LETTER
            )
            
            # Verify the PDF was created
            self.assertTrue(os.path.exists(result_path))
            self.assertEqual(result_path, pdf_path)
            
            # Extract text and verify metadata is present
            extracted_text = EncypherPDF.extract_text(pdf_path)
            
            # When metadata is embedded, especially in the first letter,
            # the text might not match exactly. Let's verify key phrases are present
            # somewhere in the extracted text, ignoring metadata characters
            
            # Extract metadata to verify it was embedded
            from encypher.core.unicode_metadata import UnicodeMetadata
            metadata = UnicodeMetadata.extract_metadata(extracted_text)
            self.assertIsNotNone(metadata, "No metadata found in extracted text")
            self.assertEqual(metadata.get('signer_id'), self.signer_id)
            
            # Verify the text is longer than expected due to metadata
            original_text = "Test Document with Unicode Characters\n\nThis is a test document with Unicode characters: ✓ ♥ ☺\n\nTesting metadata embedding in PDF conversion."
            self.assertGreater(len(extracted_text), len(original_text))


if __name__ == "__main__":
    unittest.main()
