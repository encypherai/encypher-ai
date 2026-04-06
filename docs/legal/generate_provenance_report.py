#!/usr/bin/env python3
"""Generate branded DOCX: Global Text Provenance Legislative Report.

Produces:
  docs/legal/Encypher_Global_Text_Provenance_Legislative_Report.docx

Usage:
  python3 docs/legal/generate_provenance_report.py
"""

import os
import sys

sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))

from docx.shared import Pt
from shared.docx_brand import (
    add_body as _add_body,
)
from shared.docx_brand import (
    add_bullet as _add_bullet,
)
from shared.docx_brand import (
    add_callout as _add_callout,
)
from shared.docx_brand import (
    add_heading as _add_heading,
)
from shared.docx_brand import (
    add_meta_line as _add_meta_line,
)
from shared.docx_brand import (
    add_numbered as _add_numbered,
)
from shared.docx_brand import (
    add_page_break as _add_page_break,
)
from shared.docx_brand import (
    add_subtitle as _add_subtitle,
)
from shared.docx_brand import (
    add_table as _add_table,
)
from shared.docx_brand import (
    add_title as _add_title,
)
from shared.docx_brand import (
    new_doc as _new_doc,
)
from shared.docx_brand import (
    setup_header_footer as _setup_header_footer,
)

OUTPUT_DIR = os.path.dirname(__file__)

# ---------------------------------------------------------------------------
# Report content
# ---------------------------------------------------------------------------

# fmt: off

def build_report():
    doc = _new_doc()
    _setup_header_footer(
        doc,
        author="Erik Svilich, CEO & Co-Chair, C2PA Text Task Force",
        email="erik.svilich@encypher.com",
    )

    # -----------------------------------------------------------------------
    # Title page
    # -----------------------------------------------------------------------

    # spacer
    for _ in range(4):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    _add_title(
        doc,
        "Global Regulatory Frameworks for\nAI-Generated Text Provenance",
        size=24,
    )
    _add_subtitle(
        doc,
        "A Comparative Analysis of Legislative Mandates, Technical Standards,\n"
        "and Structural Gaps Across Jurisdictions",
        size=13,
    )

    _add_meta_line(doc, "Prepared by", "Encypher Corporation")
    _add_meta_line(doc, "Author", "Erik Svilich, Co-Chair, C2PA Text Provenance Task Force")
    _add_meta_line(doc, "Date", "April 2026")
    _add_meta_line(doc, "Classification", "External - Advisory Distribution")
    _add_meta_line(doc, "Version", "1.0")

    # spacer
    for _ in range(3):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)

    _add_callout(
        doc,
        "This report provides an independent analysis of global text provenance "
        "regulation. Encypher Corporation co-chairs the C2PA Text Provenance "
        "Task Force and authored C2PA 2.3 Section A.7 (Embedding Manifests "
        "into Unstructured Text). Where the report references Encypher's "
        "public comment on the EU AI Act Code of Practice, this is identified "
        "as such. Technical assessments of C2PA capabilities reflect the "
        "published specification, not proprietary extensions.",
    )

    _add_page_break(doc)

    # -----------------------------------------------------------------------
    # Executive Summary
    # -----------------------------------------------------------------------
    _add_heading(doc, "Executive Summary", level=1)

    _add_body(
        doc,
        "Generative AI systems now produce human-like text at industrial "
        "scale. Policymakers worldwide are responding with frameworks that "
        "require provenance, the verifiable indication of whether content was "
        "generated or modified by AI, alongside the metadata of its creation "
        "and any subsequent alterations.",
    )
    _add_body(
        doc,
        "The global regulatory landscape is deeply fragmented. The European "
        "Union and the People's Republic of China have enacted binding, "
        "comprehensive transparency mandates. South Korea and Brazil have "
        "passed targeted legislation. The United States relies on voluntary "
        "federal guidance while individual states pass sector-specific laws. "
        "Japan, the United Kingdom, Australia, and Canada operate through "
        "soft-law frameworks and administrative advisories.",
    )
    _add_body(
        doc,
        "A fundamental misapprehension complicates these legal frameworks. "
        "Several major legislative efforts treat text provenance as "
        "technically infeasible and quietly exempt text from stringent "
        "latent disclosure requirements, focusing enforcement entirely on "
        "visual and auditory deepfakes. California's AI Transparency Act "
        "(SB 942) is the most prominent example, deliberately excluding "
        "text from its operational disclosure clauses. This premise is "
        "outdated. C2PA 2.3 Section A.7, published January 8, 2026, "
        "encodes provenance manifests as invisible Unicode characters "
        "embedded directly in the text stream, surviving copy-paste, "
        "syndication, and database storage with no minimum content length.",
    )
    _add_body(
        doc,
        "This report dissects these divergent approaches. It evaluates "
        "the technical viability of proposed standards, the depth and "
        "enforcement mechanisms of regional guidelines, and the structural "
        "gaps that emerge when regulatory ambition meets technical reality. "
        "Where applicable, it draws on Encypher's public comment to the EU "
        "AI Act Code of Practice to illustrate how specific technical "
        "capabilities address identified regulatory gaps.",
    )

    _add_page_break(doc)

    # -----------------------------------------------------------------------
    # 1. The Technical Landscape
    # -----------------------------------------------------------------------
    _add_heading(doc, "1. The Technical Landscape: Why Text Is Different", level=1)

    _add_body(
        doc,
        "Legislative drafting on text provenance has been shaped by the "
        "assumption that text, unlike images or audio, cannot carry "
        "reliable provenance signals. This assumption reflects the "
        "limitations of one specific technique, statistical watermarking, "
        "rather than the full range of available methods. Two distinct "
        "approaches exist, and their capabilities differ sharply.",
    )

    _add_heading(doc, "Statistical Watermarking and Its Limits", level=2)
    _add_body(
        doc,
        "Statistical text watermarking operates by altering the probability "
        "distribution of token generation during inference. The resulting "
        "statistical signature is vulnerable to paraphrasing, translation, "
        "and manual editing. Watermarking tools yield confidence scores "
        "rather than definitive attribution, making them problematic for "
        "legal liability frameworks that require certainty.",
    )

    _add_heading(doc, "Standards-Based Invisible Embedding: C2PA Section A.7", level=2)
    _add_body(
        doc,
        "The Coalition for Content Provenance and Authenticity (C2PA) "
        "provides an open technical standard for establishing the origin "
        "and edit history of digital content. C2PA manifests use "
        "cryptographic hashing to bind a signed provenance record to the "
        "digital asset, relying on X.509 certificates for identity "
        "verification and CBOR for data serialization.",
    )
    _add_body(
        doc,
        "C2PA 2.3 Section A.7 (\"Embedding Manifests into Unstructured "
        "Text,\" published January 8, 2026) solves the text distribution "
        "problem directly. Rather than wrapping text in an external "
        "container, Section A.7 encodes the entire C2PA manifest as a "
        "sequence of Unicode Variation Selectors embedded in the text "
        "stream itself. The specification opens by defining its scope: "
        "\"content intended for copy-paste operations across different "
        "systems.\" The encoding uses 256 variation selectors (U+FE00 "
        "through U+FE0F and U+E0100 through U+E01EF) to represent each "
        "byte of the binary manifest, producing characters that are "
        "invisible to readers but preserved through copy-paste, "
        "syndication, database storage, and cross-platform distribution.",
    )
    _add_body(
        doc,
        "The manifest is cryptographically bound to the text content "
        "through a c2pa.hash.data assertion. The text is NFC-normalized "
        "and hashed; the hash is recorded inside the signed manifest with "
        "byte-offset exclusions that mark where the manifest itself sits. "
        "Any modification to the text, even a single character, causes "
        "hash verification to fail. There is no minimum content length: "
        "a manifest can be attached to a one-character chatbot reply, a "
        "single sentence, or a full article.",
    )
    _add_body(
        doc,
        "Section A.7 defines two text-specific validation status codes:",
    )

    _add_table(
        doc,
        ["C2PA Text Status Code", "Meaning"],
        [
            [
                "manifest.text.corruptedWrapper",
                "A C2PATextManifestWrapper magic number (C2PATXT) was "
                "detected, but the wrapper is malformed or incomplete.",
            ],
            [
                "manifest.text.multipleWrappers",
                "More than one valid C2PATextManifestWrapper was found "
                "in the text. Selection is governed by the exclusions "
                "field of the c2pa.hash.data assertion.",
            ],
        ],
        col_widths=[2.6, 4.3],
        font_size=8,
    )

    _add_body(
        doc,
        "This approach resolves the fragility that characterizes "
        "statistical watermarking. The provenance signal is not a "
        "statistical pattern that degrades under paraphrasing; it is a "
        "cryptographically signed manifest encoded as invisible Unicode "
        "characters that travel with the text wherever it goes. The hard "
        "binding ensures definitive verification, not a confidence score.",
        space_after=2,
    )

    _add_heading(doc, "Fuzzy Provenance and Information Retrieval", level=2)
    _add_body(
        doc,
        "The U.S. National Institute of Standards and Technology (NIST) has "
        "introduced \"fuzzy provenance\" as a complementary detection "
        "mechanism. This methodology relies on search engine indexing of "
        "AI model generation logs, allowing users to find exact or "
        "near-exact matches of suspicious text. The OECD has endorsed "
        "information retrieval methods as superior to statistical "
        "watermarking for text, citing their resistance to adversarial "
        "paraphrasing. Fuzzy provenance is most useful as a backstop for "
        "content that has been aggressively paraphrased or where invisible "
        "embedding has been intentionally stripped.",
    )

    _add_callout(
        doc,
        "The technical premise that text provenance is inherently fragile "
        "applies to statistical watermarking, not to standards-based "
        "invisible embedding. C2PA Section A.7's variation selector "
        "encoding survives copy-paste, syndication, database storage, and "
        "cross-platform distribution. Fuzzy provenance via information "
        "retrieval complements this as a detection backstop. Together, "
        "these methods provide a robust multi-layered framework.",
    )

    _add_page_break(doc)

    # -----------------------------------------------------------------------
    # 2. Multilateral Standards
    # -----------------------------------------------------------------------
    _add_heading(doc, "2. Multilateral Standards and International Harmonization", level=1)

    _add_heading(doc, "The G7 Hiroshima AI Process", level=2)
    _add_body(
        doc,
        "The G7 Hiroshima AI Process represents the most significant "
        "multilateral harmonization effort. Under Principle 7, "
        "organizations developing advanced AI systems are directed to "
        "deploy reliable content authentication and provenance mechanisms, "
        "such as watermarking, wherever technically feasible. Provenance "
        "data must include a secure identifier of the service or model that "
        "created the content, though user information is explicitly exempted "
        "to comply with privacy standards.",
    )
    _add_body(
        doc,
        "The framework encourages organizations to develop public tools "
        "and APIs enabling third-party verification of AI-generated content. "
        "Under Principle 3, organizations are urged to publish detailed "
        "transparency reports on model capabilities, limitations, and "
        "domains of appropriate use. The Hiroshima Process is voluntary but "
        "serves as an architectural blueprint for domestic legislative "
        "drafting across member nations.",
    )

    _add_heading(doc, "The OECD Framework", level=2)
    _add_body(
        doc,
        "The OECD Due Diligence Guidance for Responsible AI advocates "
        "for clear, context-aware disclosure of AI-generated content, "
        "tailored to intended audiences. The OECD approach favors "
        "\"human-centric transparency,\" ensuring ordinary users can "
        "understand content provenance without specialized technical "
        "knowledge. OECD policy analysis explicitly acknowledges the "
        "technical superiority of information retrieval methods over "
        "fragile text watermarking.",
    )

    _add_page_break(doc)

    # -----------------------------------------------------------------------
    # 3. Binding Legislative Regimes
    # -----------------------------------------------------------------------
    _add_heading(doc, "3. Binding Legislative Regimes", level=1)

    # -- EU --
    _add_heading(doc, "European Union: The AI Act (Article 50)", level=2)
    _add_body(
        doc,
        "The EU AI Act treats transparency as a shared operational "
        "obligation distributed across the AI value chain. Under Article "
        "50(2), providers of generative AI systems must ensure outputs are "
        "marked in a machine-readable format and are reliably detectable "
        "as artificially generated. The EU AI Office's Code of Practice "
        "dictates a multi-layered marking strategy: digitally signed "
        "metadata, imperceptible internal watermarking, and fallback "
        "fingerprinting mechanisms, applied simultaneously.",
    )
    _add_body(
        doc,
        "Deployers face user-facing labeling obligations under Articles "
        "50(4) and 50(5). Disclosures must be clear, proportionate, and "
        "presented at the moment of first exposure, using a uniform EU-wide "
        "\"AI\" visual cue with explanatory text.",
    )

    _add_heading(doc, "The Editorial Review Exception", level=3)
    _add_body(
        doc,
        "Article 50 contains a significant carve-out for text: if "
        "AI-generated text has undergone \"genuine human review\" and a "
        "natural or legal person assumes editorial responsibility, the "
        "mandatory AI disclosure requirement is waived. This exception "
        "creates compliance ambiguity. Determining what constitutes "
        "\"genuine human review\" will become a heavily litigated standard.",
    )

    _add_callout(
        doc,
        "Encypher's public comment to the EU AI Act Code of Practice "
        "(second draft) identifies nine structural gaps in the framework. "
        "Among them: editorial review claims should be verifiable through "
        "standardized provenance records, not internal documentation alone. "
        "A C2PA manifest can record that human review occurred, who "
        "performed it, and when, transforming a self-reported assertion "
        "into a cryptographically signed attestation.",
    )

    _add_heading(doc, "Key Gaps Identified in the EU Framework", level=3)
    _add_body(
        doc,
        "Encypher's public comment to the EU drafting committee identified "
        "the following structural gaps in the second draft Code of Practice:",
    )
    _add_numbered(
        doc, 1,
        " The Code lacks a designated reference interoperability "
        "standard. Without one, each provider builds proprietary marking. "
        "Cross-provider verification becomes impossible.",
        "Interoperability. ",
    )
    _add_numbered(
        doc, 2,
        " The draft contemplates document-level marking only. Content "
        "workflows routinely produce hybrid documents containing both AI "
        "and human-authored sections. Document-level marking forces a "
        "binary choice that does not match this reality.",
        "Sub-document granularity. ",
    )
    _add_numbered(
        doc, 3,
        " For providers processing content at scale, optional "
        "fingerprinting leaves a detection gap. Watermarks embedded "
        "in text can be stripped by paraphrasing, summarization, and "
        "translation.",
        "Fingerprinting. ",
    )
    _add_numbered(
        doc, 4,
        " If marking is stripped at the first copy-paste, the "
        "compliance obligation is satisfied at publication and violated "
        "everywhere the content reaches readers.",
        "Distribution survival. ",
    )
    _add_numbered(
        doc, 5,
        " Chat-based and API-based AI interactions produce output as "
        "token-by-token streams. Marking protocols that require a "
        "completed document before signing do not work for this content.",
        "Streaming content. ",
    )
    _add_numbered(
        doc, 6,
        " The draft includes a carve-out for short text segments. "
        "Standards-based metadata (C2PA manifests) can be attached to "
        "text of any length, including a single sentence. The technical "
        "premise for the carve-out is incorrect.",
        "Short text exemption. ",
    )
    _add_numbered(
        doc, 7,
        " If verification requires payment, proprietary software, or "
        "a commercial account, the transparency framework creates a "
        "two-tier system.",
        "Free verification. ",
    )

    # -- China --
    _add_heading(doc, "People's Republic of China: AI Labeling Rules", level=2)
    _add_body(
        doc,
        "China has adopted a rigid, prescriptive approach. The Measures for "
        "Labeling Artificial Intelligence-Generated Content (supported by "
        "mandatory national standard GB 45438-2025, effective September 1, "
        "2025) impose absolute traceability requirements. Providers must "
        "place visible labels at the beginning, middle, or end of generated "
        "text, and embed implicit metadata containing the provider's "
        "registered name and a unique cryptographic content ID.",
    )

    _add_heading(doc, "Three-Tiered Platform Classification", level=3)
    _add_body(
        doc,
        "China's most severe regulatory innovation is the epistemological "
        "burden placed on content distribution platforms. Platforms must "
        "implement automated detection and categorize all user-uploaded "
        "content:",
    )
    _add_bullet(
        doc,
        " Platform detection tools identify an implicit label. The "
        "platform appends an unremovable public AI-generated label.",
        "Confirmed AI-Generated:",
    )
    _add_bullet(
        doc,
        " No implicit label detected, but users report the content "
        "as AI-generated. Platform adds a \"possibly AI-generated\" "
        "warning.",
        "Possible AI-Generated:",
    )
    _add_bullet(
        doc,
        " No implicit label, no user report, but explicit labeling or "
        "behavioral evidence suggests AI involvement. Platform must "
        "proactively label the material as \"suspected AI-generated.\"",
        "Suspected AI-Generated:",
    )
    _add_body(
        doc,
        "This system deputizes private platforms as arbiters of content "
        "authenticity at scale. Non-compliant platforms face content "
        "takedowns and license suspensions.",
    )

    # -- South Korea --
    _add_heading(doc, "South Korea: The AI Basic Act", level=2)
    _add_body(
        doc,
        "South Korea's Framework Act on AI Development and Trust (effective "
        "January 22, 2026) applies extraterritorially to AI activities "
        "impacting South Korea's domestic market. Transparency obligations "
        "apply to \"AI deployers,\" defined as entities providing AI "
        "products and services directly to end users. Internal corporate "
        "use of AI does not trigger labeling obligations.",
    )
    _add_body(
        doc,
        "The Transparency Guidelines (revised September 2025) distinguish "
        "between outputs kept within the service environment and outputs "
        "exported outside it, acknowledging the technical difficulty of "
        "maintaining text provenance once content leaves the application "
        "boundary. Advertisers using AI to generate promotional text must "
        "label their content, with legal prohibitions against tampering "
        "with or removing provenance markers.",
    )

    # -- Brazil --
    _add_heading(doc, "Brazil: AI Bill (PL 2338/2023)", level=2)
    _add_body(
        doc,
        "Brazil's AI Bill, approved by the Federal Senate in December 2024, "
        "intertwines text provenance with copyright protection. The "
        "framework mandates permanent identifiers in synthetic content and "
        "requires AI developers to disclose which copyright-protected "
        "materials were used to train their models. Text provenance in "
        "Brazil is the technical mechanism for enforcing a new economic "
        "redistribution model between AI developers and human creators.",
    )

    _add_page_break(doc)

    # -----------------------------------------------------------------------
    # 4. The United States
    # -----------------------------------------------------------------------
    _add_heading(doc, "4. The United States: Federal Guidance and State Mandates", level=1)

    _add_heading(doc, "Federal Posture: From Executive Orders to Fuzzy Provenance", level=2)
    _add_body(
        doc,
        "Executive Order 14110 (2023) mandated development of provenance "
        "tools and cited cryptographic watermarking for text, audio, and "
        "images. The order was revoked in January 2025. The burden of "
        "establishing best practices now rests with NIST, which has shifted "
        "from demanding text watermarking toward fuzzy provenance: "
        "encouraging AI companies to make generation logs indexable by "
        "search engines, with PII stripped, so users can find exact text "
        "matches within one click.",
    )

    _add_heading(doc, "California: SB 942 and the Text Exemption", level=2)
    _add_body(
        doc,
        "The California AI Transparency Act (SB 942, operative January 1, "
        "2026) targets providers with over one million monthly California "
        "users. The law imposes three core requirements: a free public AI "
        "detection tool, optional manifest disclosures (visible labels), "
        "and mandatory latent disclosures (embedded metadata).",
    )

    _add_body(
        doc,
        "A critical regulatory exemption exists in the operational clauses. "
        "While SB 942 defines GenAI systems as capable of generating "
        "\"text, images, video, and audio,\" the clauses requiring the "
        "detection tool, manifest disclosures, and latent metadata apply "
        "exclusively to \"image, video, or audio content, or content that "
        "is any combination thereof.\" Text is deliberately and entirely "
        "exempt from the law's strictest provenance requirements.",
        bold=True,
    )

    _add_callout(
        doc,
        "California's text exemption reflects a legislative calculation "
        "based on the fragility of statistical text watermarking. "
        "C2PA Section A.7 uses an entirely different mechanism: invisible "
        "Unicode variation selectors that encode a cryptographically "
        "signed provenance manifest directly in the text stream. This "
        "encoding survives copy-paste, syndication, database storage, "
        "and cross-platform distribution, with no minimum content length. "
        "A single-character AI chatbot reply can carry a definitive "
        "provenance manifest. The technical premise for the text "
        "exemption, that text marking is infeasible, does not account "
        "for this standard.",
    )

    _add_heading(doc, "New York: AI Transparency in Advertising", level=2)
    _add_body(
        doc,
        "New York's AI Transparency in Advertising Act (S.8420-A) requires "
        "disclosure of AI-generated synthetic performers in marketing "
        "materials. Companion legislation (S6954A) requires providers of "
        "synthetic content creation systems to include provenance data. "
        "State agencies must ensure audio, images, and video carry "
        "provenance data indicating the system used to generate them. "
        "Text outputs are deprioritized.",
    )

    _add_heading(doc, "Utah: SB 131 and Political Communications", level=2)
    _add_body(
        doc,
        "Utah's SB 131 targets synthetic media in political communications "
        "intended to influence voting. The law requires embedded "
        "tamper-evident provenance and visible disclaimers. Violations "
        "carry civil penalties and can serve as aggravating factors in "
        "criminal sentencing.",
    )

    _add_page_break(doc)

    # -----------------------------------------------------------------------
    # 5. Soft Law and Advisories
    # -----------------------------------------------------------------------
    _add_heading(doc, "5. Soft Law, Advisories, and Administrative Directives", level=1)

    _add_heading(doc, "India: The MeitY Advisories", level=2)
    _add_body(
        doc,
        "India's draft IT Rules amendments propose a visible declaration "
        "covering at least 10% of a content's total display area, applying "
        "to text, video, and audio. Platforms face a three-hour takedown "
        "window for non-compliant synthetic content. Civil society "
        "organizations have criticized the 10% rule as unworkable and "
        "disproportionate, arguing it forces tools to permanently deface "
        "works regardless of context.",
    )

    _add_heading(doc, "Japan: Agile Governance", level=2)
    _add_body(
        doc,
        "Japan's AI Promotion Act (May 2025) and AI Business Operator "
        "Guidelines rely on corporate self-regulation. The government "
        "encourages companies to implement watermarks and provenance tools "
        "for transparency and copyright protection, using existing statutes "
        "to handle legal ramifications. Japan adheres to the G7 Hiroshima "
        "Process principles and provides a safe harbor for AI developers.",
    )

    _add_heading(doc, "United Kingdom: Principles-Based Oversight", level=2)
    _add_body(
        doc,
        "The UK has deliberately avoided statutory AI labeling "
        "requirements, citing technological immaturity. The AI Playbook "
        "for UK Government (February 2025) advises agencies to prioritize "
        "data quality, bias mitigation, and human-in-the-loop verification "
        "over technical provenance markers. Parliamentary briefings "
        "acknowledge that no industry-wide standard for text watermarking "
        "exists and that marks can be easily removed.",
    )

    _add_heading(doc, "Australia and Canada", level=2)
    _add_body(
        doc,
        "Australia's AI Technical Standard requires agencies to apply "
        "visual watermarks and metadata to generated media, while warning "
        "against degrading text readability or WCAG accessibility "
        "compliance. Canada's Voluntary Code recommends watermarks and "
        "plain-language AI disclosures, aligning with the G7 consensus "
        "on digital trust and public auditability.",
    )

    _add_page_break(doc)

    # -----------------------------------------------------------------------
    # 6. Comparative Table
    # -----------------------------------------------------------------------
    _add_heading(doc, "6. Global Comparative Reference Table", level=1)

    _add_body(
        doc,
        "The following table synthesizes the status of text provenance "
        "guidelines across all jurisdictions analyzed in this report.",
        space_after=6,
    )

    _add_table(
        doc,
        ["Jurisdiction", "Mechanism", "Type", "Text Coverage"],
        [
            [
                "G7",
                "Hiroshima AI Process",
                "Voluntary",
                "Broad: authentication and watermarking "
                "where feasible. APIs for third-party verification.",
            ],
            [
                "OECD",
                "Due Diligence Guidance",
                "Voluntary",
                "Advisory: audience-tailored disclosure. "
                "Endorses IR methods over watermarking.",
            ],
            [
                "European Union",
                "AI Act (Art. 50) & Code of Practice",
                "Binding",
                "Deep: machine-readable marking required. "
                "Editorial review exception for text.",
            ],
            [
                "China",
                "AI Labeling Rules (GB 45438-2025)",
                "Binding",
                "Prescriptive: explicit labels at fixed "
                "positions. Platform classification system.",
            ],
            [
                "South Korea",
                "AI Basic Act & Guidelines",
                "Binding",
                "Moderate: labeling of generative outputs. "
                "Export vs. in-service distinction.",
            ],
            [
                "Brazil",
                "AI Bill (PL 2338/2023)",
                "Binding",
                "Deep: permanent identifiers. Integrated "
                "with copyright compensation framework.",
            ],
            [
                "USA (Federal)",
                "NIST Guidelines",
                "Voluntary",
                "Technical: fuzzy provenance via search "
                "engine indexing of AI logs.",
            ],
            [
                "USA (California)",
                "SB 942",
                "Binding",
                "Paradoxical: text is deliberately exempt "
                "from disclosure mandates.",
            ],
            [
                "USA (Utah)",
                "SB 131",
                "Binding",
                "Narrow: political communications only. "
                "Severe penalties.",
            ],
            [
                "India",
                "MeitY Draft IT Rules",
                "Proposed",
                "Severe: 10% display area label "
                "applies to text. 3-hour takedown.",
            ],
            [
                "Japan",
                "AI Promotion Act & Guidelines",
                "Soft law",
                "Moderate: encourages watermarks and "
                "provenance tools. Self-regulation.",
            ],
            [
                "United Kingdom",
                "AI Playbook",
                "Soft law",
                "Principle-based: acknowledges watermark "
                "immaturity. Human-in-the-loop focus.",
            ],
            [
                "Australia",
                "AI Technical Standard",
                "Govt binding",
                "Nuanced: requires watermarks. Warns "
                "against text readability degradation.",
            ],
            [
                "Canada",
                "Voluntary Code & Govt Guide",
                "Voluntary",
                "Broad: recommends watermarks and "
                "plain-language AI disclosures.",
            ],
        ],
        col_widths=[1.1, 1.8, 0.7, 3.3],
        font_size=7.5,
    )

    _add_page_break(doc)

    # -----------------------------------------------------------------------
    # 7. Structural Themes and Gap Analysis
    # -----------------------------------------------------------------------
    _add_heading(doc, "7. Structural Themes and Gap Analysis", level=1)

    _add_body(
        doc,
        "Seven structural themes emerge from the comparative analysis. "
        "Each represents a gap that recurs across multiple jurisdictions. "
        "The themes align closely with the observations submitted in "
        "Encypher's public comment to the EU AI Act Code of Practice, "
        "suggesting these are not EU-specific deficiencies but systemic "
        "patterns in how legislators approach text provenance.",
    )

    _add_heading(doc, "7.1. The Text Exemption Pattern", level=2)
    _add_body(
        doc,
        "California (SB 942) entirely exempts text from latent disclosure "
        "mandates. New York deprioritizes text outputs. The EU's editorial "
        "review exception creates a pathway to avoid AI disclosure for "
        "text that has been reviewed by a human editor. The common premise "
        "is that text marking is technically infeasible. C2PA Section A.7 "
        "refutes this premise. The standard encodes cryptographically "
        "signed provenance manifests as invisible Unicode variation "
        "selectors embedded directly in the text stream. The encoding "
        "survives copy-paste, syndication, database storage, and "
        "cross-platform distribution. A manifest can be attached to any "
        "text, including a single-character chatbot reply. No minimum "
        "content length exists.",
    )

    _add_heading(doc, "7.2. No Interoperability Standard", level=2)
    _add_body(
        doc,
        "No jurisdiction designates a reference technical standard for "
        "content provenance. The EU Code of Practice describes a "
        "two-layered architecture (secured metadata plus watermarking) "
        "without specifying which standard implements it. C2PA 2.3 is the "
        "only published multi-stakeholder standard with production "
        "implementations at scale, established AI content assertion "
        "vocabularies (c2pa.ai_generated, c2pa.ai_training), open "
        "governance under the Linux Foundation, and a text-specific "
        "embedding specification (Section A.7) that survives "
        "distribution. Without a designated baseline, providers build "
        "proprietary marking systems and cross-provider verification "
        "becomes impossible.",
    )

    _add_heading(doc, "7.3. Distribution Survival Is Not Addressed", level=2)
    _add_body(
        doc,
        "Content is published, syndicated, aggregated, quoted, and "
        "reshared. No jurisdiction requires marking to survive common "
        "distribution paths. C2PA Section A.7's variation selector "
        "encoding was designed for exactly this scenario: the spec's "
        "stated scope is \"content intended for copy-paste operations "
        "across different systems.\" The encoding persists through "
        "copy-paste, syndication, databases, and cross-platform sharing. "
        "Legislators have not accounted for this capability, treating "
        "distribution survival as an unsolved problem rather than a "
        "feature of the published standard.",
    )

    _add_heading(doc, "7.4. Streaming and Real-Time Content Is Ignored", level=2)
    _add_body(
        doc,
        "The highest-volume category of AI-generated text, chat-based "
        "and API-based interactions, produces output as token-by-token "
        "streams rather than static documents. No jurisdiction explicitly "
        "addresses streaming content in its marking obligations. Marking "
        "protocols that require a completed document before signing do not "
        "work for this content. Incremental authentication during "
        "streaming is technically feasible and production-deployed.",
    )

    _add_heading(doc, "7.5. The Verification Asymmetry", level=2)
    _add_body(
        doc,
        "Marking AI-generated content has value only if third parties can "
        "check the markings. Several jurisdictions require or encourage "
        "detection tools but do not mandate that verification be free, "
        "open, and unauthenticated. If verification requires payment or "
        "proprietary software, the transparency framework creates a "
        "two-tier system: organizations that pay see provenance data; the "
        "general public does not.",
    )

    _add_heading(doc, "7.6. Platform Epistemological Burden", level=2)
    _add_body(
        doc,
        "China and India deputize private platforms as the primary "
        "enforcers of text provenance. Platforms must scan incoming text, "
        "categorize it as confirmed, possible, or suspected AI-generated, "
        "and forcefully append labels. Current AI text detection tools "
        "carry high false-positive rates. Platforms will inevitably flag "
        "genuine human-authored text as \"suspected AI,\" chilling free "
        "expression and degrading public trust in the provenance framework "
        "itself.",
    )

    _add_heading(doc, "7.7. Compliance Fragmentation", level=2)
    _add_body(
        doc,
        "A global AI provider deploying an LLM-powered chatbot must now "
        "architect a geofenced user interface: a 10% visual overlay for "
        "Indian users, rigid positional labels for Chinese users, "
        "human-review audit trails for EU users, and no text requirements "
        "in California. This fragmentation increases engineering, legal, "
        "and operational costs, pricing out smaller developers and "
        "consolidating market power among hyperscalers capable of "
        "maintaining multi-jurisdictional compliance architectures.",
    )

    _add_page_break(doc)

    # -----------------------------------------------------------------------
    # 8. Recommendations
    # -----------------------------------------------------------------------
    _add_heading(doc, "8. Observations for California Policymakers", level=1)

    _add_body(
        doc,
        "California's SB 942 is among the most structurally detailed "
        "AI transparency laws in the United States. Its text exemption, "
        "while reflecting a defensible technical judgment about statistical "
        "watermarking, creates a gap that other jurisdictions have "
        "addressed or are addressing. The following observations are drawn "
        "from the comparative analysis above.",
    )

    _add_numbered(
        doc, 1,
        " SB 942's text exemption rests on the assumption that text "
        "marking is technically infeasible. C2PA Section A.7, published "
        "January 8, 2026, encodes provenance manifests as invisible "
        "Unicode variation selectors directly in the text stream. The "
        "encoding survives copy-paste, syndication, and database storage. "
        "A single-character output can carry a signed manifest. The EU, "
        "China, South Korea, and Brazil all include text in their "
        "provenance mandates.",
        "The technical premise for the text exemption is outdated. ",
    )
    _add_numbered(
        doc, 2,
        " C2PA 2.3 is the only published multi-stakeholder standard "
        "with production implementations, AI content assertion "
        "vocabularies, and open governance. Referencing C2PA as a "
        "baseline technical standard would align California with "
        "emerging international norms and prevent the proliferation "
        "of incompatible proprietary marking systems.",
        "An interoperability standard exists. ",
    )
    _add_numbered(
        doc, 3,
        " California's latent disclosure requirements apply to "
        "images, video, and audio. The highest-volume category of "
        "AI-generated content reaching consumers is text: chatbot "
        "interactions, generated summaries, automated replies, social "
        "media posts. Exempting text creates a transparency gap that "
        "covers the majority of what consumers actually encounter.",
        "Text is the highest-volume AI output. ",
    )
    _add_numbered(
        doc, 4,
        " NIST's fuzzy provenance approach, encouraging AI providers "
        "to make generation logs searchable, complements rather than "
        "replaces file-level metadata. A state-level requirement for "
        "metadata-based text provenance (distinct from statistical "
        "watermarking) would be compatible with the federal direction.",
        "Federal guidance supports, not precludes, state action. ",
    )
    _add_numbered(
        doc, 5,
        " The EU's editorial review exception demonstrates that text "
        "provenance mandates can coexist with editorial freedom. If a "
        "human editor reviews the content and assumes responsibility, "
        "the AI disclosure is waived. A similar mechanism could address "
        "concerns about over-labeling in California.",
        "Editorial exceptions can balance transparency and expression. ",
    )

    _add_page_break(doc)

    # -----------------------------------------------------------------------
    # 9. Conclusion
    # -----------------------------------------------------------------------
    _add_heading(doc, "9. Conclusion", level=1)

    _add_body(
        doc,
        "The global regulatory pursuit of AI-generated text provenance is "
        "defined by a tension between legal ambition and technical reality. "
        "Binding regimes in the EU, China, South Korea, and Brazil impose "
        "specific labeling, metadata, and platform enforcement requirements "
        "that often push beyond the limits of current cryptographic "
        "capabilities. The United States presents a regulatory paradox: "
        "federal agencies pioneer network-based solutions like fuzzy "
        "provenance, while California's most prominent state-level "
        "legislation deliberately exempts text from its strictest "
        "disclosure mandates.",
    )
    _add_body(
        doc,
        "Robust text provenance cannot rely on statistical watermarks, "
        "which degrade under paraphrasing and yield confidence scores "
        "rather than definitive verification. C2PA Section A.7 provides "
        "a fundamentally different mechanism: cryptographically signed "
        "manifests encoded as invisible Unicode variation selectors, "
        "embedded directly in the text stream, surviving copy-paste, "
        "syndication, database storage, and cross-platform distribution "
        "with no minimum content length. Fuzzy provenance via information "
        "retrieval complements this as a detection backstop for "
        "aggressively paraphrased content.",
    )
    _add_body(
        doc,
        "International harmonization, driven by frameworks like the G7 "
        "Hiroshima Process, is necessary to prevent a fractured digital "
        "ecosystem where the legal definition of content authenticity "
        "changes depending on the jurisdiction in which a sentence is "
        "read. The technical standards to support this harmonization "
        "exist today. The question is whether legislators will reference "
        "them.",
    )

    # -----------------------------------------------------------------------
    # Save
    # -----------------------------------------------------------------------
    path = os.path.join(
        OUTPUT_DIR,
        "Encypher_Global_Text_Provenance_Legislative_Report.docx",
    )
    doc.save(path)
    print(f"  Saved: {path}")
    return path

# fmt: on

# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    print("Generating Global Text Provenance Legislative Report...")
    build_report()
    print("Done.")
