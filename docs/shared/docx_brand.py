"""Encypher branded DOCX utilities.

Shared brand constants, header/footer setup, and content helpers for all
DOCX generators. Import from here instead of duplicating brand definitions.

Usage from any generator script under docs/:

    import sys, os
    sys.path.insert(0, os.path.join(os.path.dirname(__file__), ".."))
    from shared.docx_brand import (
        new_doc, setup_header_footer,
        add_heading, add_body, add_body_mixed, add_bullet, add_numbered,
        add_table, add_callout, add_title, add_subtitle, add_page_break,
        add_hyperlink, set_cell_shading,
        DELFT_BLUE, BLUE_NCS, COLUMBIA_BLUE, WHITE, BODY_TEXT, LIGHT_BG,
        FONT, FONT_MONO,
    )
"""

import os

from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn
from docx.shared import Cm, Inches, Pt, RGBColor

# ---------------------------------------------------------------------------
# Brand constants
# ---------------------------------------------------------------------------
DELFT_BLUE = RGBColor(0x1B, 0x2F, 0x50)
BLUE_NCS = RGBColor(0x2A, 0x87, 0xC4)
COLUMBIA_BLUE = RGBColor(0xB7, 0xD5, 0xED)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BODY_TEXT = RGBColor(0x33, 0x33, 0x33)
LIGHT_BG = "EDF4F9"
ACCENT_RED = RGBColor(0xC0, 0x39, 0x2B)

FONT = "Roboto"
FONT_MONO = "Roboto Mono"

_REPO_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), "../.."))
LOGO_PATH = os.path.join(_REPO_ROOT, "apps/marketing-site/public/encypher_full_nobg.png")

SITE_URL = "https://encypher.com/?utm_source=report&utm_medium=docx"
DOCS_URL = "https://docs.encypher.com"
EMAIL = "info@encypher.com"
MAILTO_URL = f"mailto:{EMAIL}"

# Relationship type for external hyperlinks (Office Open XML standard).
_HYPERLINK_REL = "http://schemas.openxmlformats.org/officeDocument/2006/relationships" "/hyperlink"


# ---------------------------------------------------------------------------
# Helpers
# ---------------------------------------------------------------------------


def set_cell_shading(cell, hex_color):
    """Apply a solid background fill to a table cell."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_page_break(doc):
    """Insert a hard page break."""
    from docx.enum.text import WD_BREAK

    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)


def add_hyperlink(paragraph, text, url, color=BLUE_NCS, size_pt=7, bold=False):
    """Append a hyperlink run to *paragraph*.

    Works in body, header, and footer contexts. The link renders with the
    specified color and no underline (clean brand style).
    """
    part = paragraph.part
    r_id = part.relate_to(url, _HYPERLINK_REL, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    run_el = OxmlElement("w:r")

    # -- Run properties --
    rPr = OxmlElement("w:rPr")

    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), FONT)
    rFonts.set(qn("w:hAnsi"), FONT)
    rPr.append(rFonts)

    half_pts = str(int(size_pt * 2))
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), half_pts)
    rPr.append(sz)
    szCs = OxmlElement("w:szCs")
    szCs.set(qn("w:val"), half_pts)
    rPr.append(szCs)

    c = OxmlElement("w:color")
    c.set(qn("w:val"), str(color))
    rPr.append(c)

    if bold:
        rPr.append(OxmlElement("w:b"))

    # No underline -- keep it clean.
    u = OxmlElement("w:u")
    u.set(qn("w:val"), "none")
    rPr.append(u)

    run_el.append(rPr)

    t = OxmlElement("w:t")
    t.set(qn("xml:space"), "preserve")
    t.text = text
    run_el.append(t)

    hyperlink.append(run_el)
    paragraph._p.append(hyperlink)


def _add_plain_run(paragraph, text, color=BODY_TEXT, size_pt=7, bold=False):
    """Append a plain (non-hyperlink) styled run to *paragraph*."""
    run = paragraph.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size_pt)
    run.font.color.rgb = color
    run.bold = bold
    return run


def _add_page_field(paragraph, color=BODY_TEXT, size_pt=7):
    """Insert an auto-updating PAGE number field into *paragraph*."""
    run = paragraph.add_run()
    run.font.name = FONT
    run.font.size = Pt(size_pt)
    run.font.color.rgb = color
    run._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="begin"/>'))
    run._r.append(parse_xml(f'<w:instrText {nsdecls("w")} ' f'xml:space="preserve"> PAGE </w:instrText>'))
    run._r.append(parse_xml(f'<w:fldChar {nsdecls("w")} w:fldCharType="end"/>'))


# ---------------------------------------------------------------------------
# Header / Footer setup
# ---------------------------------------------------------------------------


def setup_header_footer(doc, logo_width=Inches(1.4), author=None, email=None, site_label=None, site_url=None):
    """Configure branded header (logo) and footer (contact + page number).

    Call once after creating the Document and before adding body content.
    The header and footer repeat on every page automatically.

    Footer layout:
        Default:      Encypher Corporation | info@encypher.com | encypher.com
        With author:  Encypher Corporation | Author Name | author@co | encypher.com

    Optional overrides:
        author     - e.g. "Erik Svilich, CEO"  (omit for company-only footer)
        email      - author's email when credited (ignored without author)
        site_label - e.g. "encypher.com"        (defaults to "encypher.com")
        site_url   - e.g. "https://..."         (defaults to module SITE_URL)
    """
    _site_label = site_label or "encypher.com"
    _site_url = site_url or SITE_URL
    for section in doc.sections:
        # -- Header: logo --
        header = section.header
        header.is_linked_to_previous = False
        hp = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
        hp.paragraph_format.space_after = Pt(0)
        hp.paragraph_format.space_before = Pt(0)
        if os.path.exists(LOGO_PATH):
            run = hp.add_run()
            run.add_picture(LOGO_PATH, width=logo_width)
        else:
            print(f"  WARNING: logo not found at {LOGO_PATH}")

        # -- Footer: single-row table [contact info | page number] --
        footer = section.footer
        footer.is_linked_to_previous = False
        for fp in footer.paragraphs:
            fp.clear()

        ft = footer.add_table(rows=1, cols=2, width=Inches(6.5))
        ft.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Table borders: thin top rule only.
        tbl_el = ft._tbl
        tblPr = tbl_el.tblPr if tbl_el.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
        borders = parse_xml(
            f'<w:tblBorders {nsdecls("w")}>'
            f'  <w:top w:val="single" w:sz="4" w:space="0" '
            f'w:color="B7D5ED"/>'
            f'  <w:bottom w:val="none" w:sz="0" w:space="0" '
            f'w:color="auto"/>'
            f'  <w:left w:val="none" w:sz="0" w:space="0" '
            f'w:color="auto"/>'
            f'  <w:right w:val="none" w:sz="0" w:space="0" '
            f'w:color="auto"/>'
            f'  <w:insideH w:val="none" w:sz="0" w:space="0" '
            f'w:color="auto"/>'
            f'  <w:insideV w:val="none" w:sz="0" w:space="0" '
            f'w:color="auto"/>'
            f"</w:tblBorders>"
        )
        tblPr.append(borders)

        # Left cell: Encypher Corporation | email | site
        left_cell = ft.rows[0].cells[0]
        left_cell.width = Inches(5.2)
        lp = left_cell.paragraphs[0]
        lp.paragraph_format.space_before = Pt(2)
        lp.paragraph_format.space_after = Pt(0)

        _add_plain_run(lp, "Encypher Corporation", DELFT_BLUE, bold=True)
        if author:
            _add_plain_run(lp, f"  |  {author}", BODY_TEXT)
            if email:
                _add_plain_run(lp, "  |  ")
                add_hyperlink(lp, email, f"mailto:{email}", color=BODY_TEXT)
        else:
            _add_plain_run(lp, "  |  ")
            add_hyperlink(lp, EMAIL, MAILTO_URL, color=BODY_TEXT)
        _add_plain_run(lp, "  |  ")
        add_hyperlink(lp, _site_label, _site_url, color=BLUE_NCS)

        # Right cell: page number, right-aligned.
        right_cell = ft.rows[0].cells[1]
        right_cell.width = Inches(1.3)
        rp = right_cell.paragraphs[0]
        rp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rp.paragraph_format.space_before = Pt(2)
        rp.paragraph_format.space_after = Pt(0)
        _add_page_field(rp, color=BODY_TEXT)


# ---------------------------------------------------------------------------
# Document creation
# ---------------------------------------------------------------------------


def new_doc():
    """Create a new Document with standard US Letter margins."""
    from docx import Document

    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        section.page_height = Cm(27.94)
        section.page_width = Cm(21.59)
    return doc


# ---------------------------------------------------------------------------
# Content helpers
# ---------------------------------------------------------------------------


def add_title(doc, text, size=22):
    """Add a large branded title with a blue underline rule."""
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.color.rgb = DELFT_BLUE
    run.bold = True
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(4)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}>' f'  <w:bottom w:val="single" w:sz="8" w:space="3" ' f'w:color="2A87C4"/>' f"</w:pBdr>")
    pPr.append(pBdr)
    return p


def add_subtitle(doc, text, size=13):
    """Add an italic subtitle in brand blue."""
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.color.rgb = BLUE_NCS
    run.italic = True
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    return p


def add_heading(doc, text, level=2):
    """Add a branded heading. Keeps with the next paragraph to avoid orphans."""
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    run = p.add_run(text)
    run.font.name = FONT
    run.bold = True
    if level == 1:
        run.font.size = Pt(16)
        run.font.color.rgb = DELFT_BLUE
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(6)
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}>' f'  <w:bottom w:val="single" w:sz="6" w:space="2" ' f'w:color="2A87C4"/>' f"</w:pBdr>")
        pPr.append(pBdr)
    elif level == 2:
        run.font.size = Pt(12)
        run.font.color.rgb = BLUE_NCS
        p.paragraph_format.space_before = Pt(10)
        p.paragraph_format.space_after = Pt(4)
    elif level == 3:
        run.font.size = Pt(10)
        run.font.color.rgb = DELFT_BLUE
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(3)
    return p


def add_body(doc, text, size=9.5, space_after=4, bold=False, italic=False):
    """Add a body paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.color.rgb = BODY_TEXT
    run.bold = bold
    run.italic = italic
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(14)
    return p


def add_body_mixed(doc, parts, size=9.5, space_after=4):
    """Add a paragraph with mixed bold/italic runs.

    *parts*: list of ``(text, bold)`` or ``(text, bold, italic)`` tuples.
    """
    p = doc.add_paragraph()
    for item in parts:
        text = item[0]
        bold = item[1]
        italic = item[2] if len(item) > 2 else False
        run = p.add_run(text)
        run.font.name = FONT
        run.font.size = Pt(size)
        run.font.color.rgb = BODY_TEXT
        run.bold = bold
        run.italic = italic
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(14)
    return p


def add_bullet(doc, text, bold_prefix=None, size=9):
    """Add a bullet-pointed paragraph with optional bold prefix."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(13)
    bullet_run = p.add_run("\u2022  ")
    bullet_run.font.name = FONT
    bullet_run.font.size = Pt(size)
    bullet_run.font.color.rgb = BLUE_NCS
    if bold_prefix:
        br = p.add_run(bold_prefix)
        br.font.name = FONT
        br.font.size = Pt(size)
        br.font.color.rgb = DELFT_BLUE
        br.bold = True
    body_run = p.add_run(text)
    body_run.font.name = FONT
    body_run.font.size = Pt(size)
    body_run.font.color.rgb = BODY_TEXT
    return p


def add_numbered(doc, number, text, bold_prefix=None, size=9):
    """Add a numbered paragraph with optional bold prefix."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.8)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(13)
    num_run = p.add_run(f"{number}.  ")
    num_run.font.name = FONT
    num_run.font.size = Pt(size)
    num_run.font.color.rgb = BLUE_NCS
    num_run.bold = True
    if bold_prefix:
        br = p.add_run(bold_prefix)
        br.font.name = FONT
        br.font.size = Pt(size)
        br.font.color.rgb = DELFT_BLUE
        br.bold = True
    body_run = p.add_run(text)
    body_run.font.name = FONT
    body_run.font.size = Pt(size)
    body_run.font.color.rgb = BODY_TEXT
    return p


def add_table(doc, headers, rows, col_widths=None, font_size=8.5):
    """Add a branded table with header row and alternating row shading."""

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(header)
        run.font.name = FONT
        run.font.size = Pt(font_size)
        run.font.color.rgb = WHITE
        run.bold = True
        set_cell_shading(cell, "1B2F50")

    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after = Pt(2)
            run = p.add_run(cell_text)
            run.font.name = FONT
            run.font.size = Pt(font_size)
            run.font.color.rgb = BODY_TEXT
            if r_idx % 2 == 1:
                set_cell_shading(cell, LIGHT_BG)

    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="B7D5ED"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" '
        f'w:color="B7D5ED"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" '
        f'w:color="B7D5ED"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" '
        f'w:color="B7D5ED"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" '
        f'w:color="B7D5ED"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" '
        f'w:color="B7D5ED"/>'
        f"</w:tblBorders>"
    )
    tblPr.append(borders)
    return table


def add_callout(doc, text, size=9):
    """Add a highlighted callout box with blue left border."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.right_indent = Cm(0.6)
    p.paragraph_format.line_spacing = Pt(13)
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{LIGHT_BG}" w:val="clear"/>')
    pPr.append(shading)
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}>' f'  <w:left w:val="single" w:sz="18" w:space="6" ' f'w:color="2A87C4"/>' f"</w:pBdr>")
    pPr.append(pBdr)
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.color.rgb = DELFT_BLUE
    run.italic = True
    return p


def add_meta_line(doc, label, value, size=9):
    """Add a label: value metadata line (e.g. 'Author: Erik Svilich')."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    lbl = p.add_run(f"{label}: ")
    lbl.font.name = FONT
    lbl.font.size = Pt(size)
    lbl.font.color.rgb = DELFT_BLUE
    lbl.bold = True
    val = p.add_run(value)
    val.font.name = FONT
    val.font.size = Pt(size)
    val.font.color.rgb = BODY_TEXT
    return p
