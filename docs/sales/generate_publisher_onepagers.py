#!/usr/bin/env python3
"""Generate branded 1-page publisher sales documents (Zoho Writer compatible).

Produces:
  - Encypher_Publisher_BusinessOutcome_v4.docx
  - Encypher_Publisher_Technical_v4.docx

Usage:
  python3 docs/sales/generate_publisher_onepagers.py
"""

import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Cm, Inches, Pt, RGBColor

# ---------------------------------------------------------------------------
# Brand constants
# ---------------------------------------------------------------------------
DELFT_BLUE = RGBColor(0x1B, 0x2F, 0x50)
BLUE_NCS = RGBColor(0x2A, 0x87, 0xC4)
COLUMBIA_BLUE = RGBColor(0xB7, 0xD5, 0xED)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BODY_TEXT = RGBColor(0x33, 0x33, 0x33)
LIGHT_BG = "EDF4F9"

FONT = "Roboto"
FONT_MONO = "Roboto Mono"

REPO_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), "../.."))
LOGO_PATH = os.path.join(REPO_ROOT, "apps/marketing-site/public/encypher_full_nobg.png")
OUTPUT_DIR = os.path.dirname(__file__)


# ---------------------------------------------------------------------------
# Helpers (Zoho-safe: no custom XML parts, no macros, simple borders only)
# ---------------------------------------------------------------------------


def _new_doc():
    """Create a doc with tight 1-page margins."""
    doc = Document()
    for section in doc.sections:
        section.top_margin = Cm(2.54)  # 1 inch standard
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(2.54)
        section.right_margin = Cm(2.54)
        section.page_height = Cm(27.94)  # US Letter
        section.page_width = Cm(21.59)
    return doc


def _set_cell_shading(cell, hex_color):
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def _add_logo_header(doc):
    """Add logo aligned left."""
    if not os.path.exists(LOGO_PATH):
        print(f"  WARNING: logo not found at {LOGO_PATH}")
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run()
    run.add_picture(LOGO_PATH, width=Inches(1.6))
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(0)


def _add_title(doc, text, size=18):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.color.rgb = DELFT_BLUE
    run.bold = True
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    # thin blue rule
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}>' f'  <w:bottom w:val="single" w:sz="6" w:space="2" w:color="2A87C4"/>' f'</w:pBdr>')
    pPr.append(pBdr)
    return p


def _add_heading(doc, text, level=2):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT
    run.bold = True
    if level == 2:
        run.font.size = Pt(10.5)
        run.font.color.rgb = BLUE_NCS
        p.paragraph_format.space_before = Pt(7)
        p.paragraph_format.space_after = Pt(2)
    elif level == 3:
        run.font.size = Pt(9)
        run.font.color.rgb = DELFT_BLUE
        p.paragraph_format.space_before = Pt(5)
        p.paragraph_format.space_after = Pt(1)
    return p


def _add_body(doc, text, size=8.5, space_after=3, bold=False):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.color.rgb = BODY_TEXT
    run.bold = bold
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(12)
    return p


def _add_body_mixed(doc, parts, size=8.5, space_after=3):
    """Add a paragraph with mixed bold/normal runs.

    parts: list of (text, bold) tuples.
    """
    p = doc.add_paragraph()
    for text, bold in parts:
        run = p.add_run(text)
        run.font.name = FONT
        run.font.size = Pt(size)
        run.font.color.rgb = BODY_TEXT
        run.bold = bold
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(12)
    return p


def _add_bullet(doc, text, bold_prefix=None, size=8.5):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(11)
    bullet_run = p.add_run("\u2022  ")
    bullet_run.font.name = FONT
    bullet_run.font.size = Pt(size)
    bullet_run.font.color.rgb = BLUE_NCS
    if bold_prefix:
        br = p.add_run(bold_prefix)
        br.font.name = FONT
        br.font.size = Pt(size)
        br.font.color.rgb = DELFT_BLUE
        br.bold = True
    body_run = p.add_run(text)
    body_run.font.name = FONT
    body_run.font.size = Pt(size)
    body_run.font.color.rgb = BODY_TEXT
    return p


def _add_numbered(doc, number, text, bold_prefix=None, size=8.5):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.6)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = Pt(11)
    num_run = p.add_run(f"{number}.  ")
    num_run.font.name = FONT
    num_run.font.size = Pt(size)
    num_run.font.color.rgb = BLUE_NCS
    num_run.bold = True
    if bold_prefix:
        br = p.add_run(bold_prefix)
        br.font.name = FONT
        br.font.size = Pt(size)
        br.font.color.rgb = DELFT_BLUE
        br.bold = True
    body_run = p.add_run(text)
    body_run.font.name = FONT
    body_run.font.size = Pt(size)
    body_run.font.color.rgb = BODY_TEXT
    return p


def _add_table(doc, headers, rows, col_widths=None, font_size=8):
    """Add a compact branded table. col_widths in inches."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Inches(w)

    # Header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after = Pt(2)
        run = p.add_run(header)
        run.font.name = FONT
        run.font.size = Pt(font_size)
        run.font.color.rgb = WHITE
        run.bold = True
        _set_cell_shading(cell, "1B2F50")

    # Data rows
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.space_after = Pt(1)
            run = p.add_run(cell_text)
            run.font.name = FONT
            run.font.size = Pt(font_size)
            run.font.color.rgb = BODY_TEXT
            if r_idx % 2 == 1:
                _set_cell_shading(cell, LIGHT_BG)

    # Light borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="B7D5ED"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="B7D5ED"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="B7D5ED"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="B7D5ED"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="B7D5ED"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="B7D5ED"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)
    return table


def _add_code_block(doc, text, size=7.5):
    """Add a monospace code block with light background."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.left_indent = Cm(0.4)
    p.paragraph_format.line_spacing = Pt(10)
    # Light background via paragraph shading
    pPr = p._p.get_or_add_pPr()
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{LIGHT_BG}" w:val="clear"/>')
    pPr.append(shading)
    run = p.add_run(text)
    run.font.name = FONT_MONO
    run.font.size = Pt(size)
    run.font.color.rgb = DELFT_BLUE
    return p


def _add_footer(doc):
    """Add a compact centered footer with contact info and links."""
    # Thin rule
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}>' f'  <w:top w:val="single" w:sz="4" w:space="2" w:color="B7D5ED"/>' f'</w:pBdr>')
    pPr.append(pBdr)
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)

    # Contact line 1
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    parts = [
        ("Erik Svilich", True),
        (" - CEO & Co-Chair, C2PA Text Task Force - erik.svilich@encypher.com", False),
    ]
    for text, bold in parts:
        run = p.add_run(text)
        run.font.name = FONT
        run.font.size = Pt(7.5)
        run.font.color.rgb = DELFT_BLUE if bold else BODY_TEXT
        run.bold = bold

    # Contact line 2
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.space_before = Pt(0)
    parts = [
        ("Matt Kaminsky", True),
        (" - Chief Commercial Officer - matt.kaminsky@encypher.com", False),
    ]
    for text, bold in parts:
        run = p.add_run(text)
        run.font.name = FONT
        run.font.size = Pt(7.5)
        run.font.color.rgb = DELFT_BLUE if bold else BODY_TEXT
        run.bold = bold

    # Links line
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(1)
    run = p.add_run("encypher.com  |  dashboard.encypher.com  |  docs.encypher.com")
    run.font.name = FONT
    run.font.size = Pt(7.5)
    run.font.color.rgb = BLUE_NCS


# ---------------------------------------------------------------------------
# Document 1: Business Outcome
# ---------------------------------------------------------------------------


def build_business_outcome():
    doc = _new_doc()
    _add_logo_header(doc)
    _add_title(doc, "Get Paid When AI Uses Your Content", size=16)

    # Tagline
    _add_body(
        doc,
        'Eliminate the "we didn\'t know it was yours" defense.',
        size=9.5,
        bold=True,
        space_after=4,
    )

    # The Problem
    _add_heading(doc, "The Problem")
    _add_body(
        doc,
        "Content on the open web has no proof of origin. When your articles, images, audio, "
        "video, and documents are copy-pasted, distributed through wire services, scraped by "
        "aggregators, or used to train AI models, the connection to your organization "
        "disappears. AI companies claim they didn't know. That defense works because "
        "ownership is unprovable. Encypher makes it provable.",
    )

    # The Legal Shift
    _add_heading(doc, "The Legal Shift")
    _add_table(
        doc,
        ["Without Encypher", "With Encypher"],
        [
            [
                "Innocent infringement. AI company claims it did not know "
                "the content was protected. You face expensive discovery, "
                "forensic analysis, and uncertain outcomes.",
                "Willful infringement. Invisible proof of ownership travels "
                "with every sentence. Formal notice is on the record. "
                "Continued unauthorized use is provable and intentional.",
            ]
        ],
        col_widths=[3.45, 3.45],
        font_size=7.5,
    )

    # Mark It. Track It. Own It.
    _add_heading(doc, "Mark It. Track It. Own It.")
    _add_table(
        doc,
        ["Step", "What Happens"],
        [
            [
                "1. Mark It",
                "Invisible, tamper-evident signatures embed directly into your content "
                "- text, images, audio, video, live streams, and documents. "
                "Survives copy-paste, scraping, and syndication. Audiences see nothing different.",
            ],
            [
                "2. Track It",
                "Your dashboard shows where signed content appears across "
                "competitor sites, aggregators, and AI outputs. "
                "Evidence trails build automatically.",
            ],
            [
                "3. Own It",
                "Set machine-readable licensing terms. Serve formal notice to "
                "AI companies using your content. License through the coalition "
                "or negotiate directly.",
            ],
        ],
        col_widths=[1.0, 5.6],
        font_size=7.5,
    )

    # Pricing summary
    _add_heading(doc, "Free to Sign. Paid to Enforce.")
    _add_table(
        doc,
        ["Free - Every Publisher", "Enforcement Add-Ons"],
        [
            [
                "Prove ownership of every piece you publish\n"
                "Full-stack: text, images, audio, video, streams, docs\n"
                "Invisible signatures that survive redistribution\n"
                "Machine-readable licensing terms\n"
                "Coalition membership\n"
                "WordPress plugin, Chrome extension, API",
                "Attribution Analytics (detection dashboard)\n"
                "Formal Notice packages\n"
                "Court-ready evidence packages\n"
                "Enterprise: unlimited volume, custom identity,\n"
                "dedicated SLA, bulk archive backfill ($0.01/doc)",
            ]
        ],
        col_widths=[3.5, 3.5],
        font_size=7.5,
    )

    # Two-track licensing
    _add_body_mixed(
        doc,
        [
            ("Two-Track Licensing: ", True),
            ("Coalition deals (we negotiate) or self-service deals (you negotiate). " "Choose per deal. We only win when you win.", False),
        ],
        size=8,
        space_after=2,
    )

    # What Makes Encypher Different
    _add_heading(doc, "What Makes Encypher Different")
    _add_table(
        doc,
        ["Approach", "How It Works", "Limitation"],
        [
            [
                "Access Gates\n(TollBit, Cloudflare)",
                "Block or meter AI crawlers at your server.",
                "Only works when AI companies opt in. " "Does nothing for content already distributed.",
            ],
            [
                "Opt-In Marketplaces\n(ProRata, Microsoft)",
                "Broker deals between publishers and AI companies.",
                "Revenue depends on platform adoption. " "Microsoft is simultaneously buyer and broker.",
            ],
            [
                "Encypher\n(Content Provenance)",
                "Invisible proof of origin embedded in the content itself "
                "- text, images, audio, video, streams, and documents. "
                "Survives redistribution.",
                "Works unilaterally. No AI company\n" "cooperation required.",
            ],
        ],
        col_widths=[1.7, 2.7, 2.5],
        font_size=7.5,
    )
    _add_body(
        doc,
        "These layers are complementary. Access gates handle cooperative crawlers. "
        "Marketplaces handle willing buyers. Encypher handles everything else "
        "- it is the only layer that works when AI companies do not opt in.",
        size=8,
        space_after=2,
    )

    # Why Encypher. Why Now.
    _add_heading(doc, "Why Encypher. Why Now.")
    _add_bullet(
        doc,
        "We authored the standard. Erik Svilich wrote the C2PA text provenance "
        "specification and co-chairs the Text Task Force. "
        "Encypher is the reference implementation.",
        size=8,
    )
    _add_bullet(doc, "Publishers who sign now define how AI content licensing works. " "Later adopters accept terms others set.", size=8)
    _add_bullet(doc, "Your archive is your most valuable asset. " "It should not exist on the open web without proof it is yours.", size=8)

    # Integration
    _add_heading(doc, "Get Started in Minutes")
    _add_body(
        doc,
        "WordPress plugin (one-click, auto-signs on publish), Ghost CMS webhook, "
        "REST API with Python and TypeScript SDKs, or bulk archive backfill at $0.01/doc. "
        "Average setup time: under 20 minutes. No engineering staff required.",
        size=8,
        space_after=2,
    )

    # CTA
    _add_heading(doc, "Next Step: See It In Action")
    _add_body(
        doc,
        "15-minute demo. Mark a subset of your content. See results in your dashboard. " "No commitment required.",
        bold=True,
    )

    _add_footer(doc)

    path = os.path.join(OUTPUT_DIR, "Encypher_Publisher_BusinessOutcome_v4.docx")
    doc.save(path)
    print(f"  Saved: {path}")
    return path


# ---------------------------------------------------------------------------
# Document 2: Technical Architecture
# ---------------------------------------------------------------------------


def build_technical():
    doc = _new_doc()
    _add_logo_header(doc)
    _add_title(doc, "Technical Architecture Overview", size=16)

    # Standards Foundation
    _add_heading(doc, "Standards Foundation")
    _add_body(
        doc,
        "C2PA 2.3 Section A.7 text provenance specification. "
        "Interoperable with RSL 1.0, ODRL, and JSON-LD. "
        "Patent-pending sentence-level extensions (Claims 1-20, 38-52) go beyond the public standard.",
    )

    # Supported Formats
    _add_heading(doc, "Full-Stack Content Provenance: 31+ Formats Across 7 Categories")
    _add_bullet(doc, " Plain text, HTML, Markdown, XML, JSON, CSV, source code", "Text: ", size=7.5)
    _add_bullet(doc, " JPEG, PNG, WebP, TIFF, GIF, SVG, AVIF, HEIF, DNG", "Images: ", size=7.5)
    _add_bullet(doc, " MP3, WAV, FLAC, AAC, OGG, AIFF, M4A", "Audio: ", size=7.5)
    _add_bullet(doc, " MP4, MOV, WebM, AVI, MKV + real-time live stream signing (WSS/SSE)", "Video: ", size=7.5)
    _add_bullet(doc, " PDF, DOCX, XLSX, PPTX, ODP, ODT, ODS  |  TTF, OTF, WOFF, WOFF2", "Documents & Fonts: ", size=7.5)

    # Architecture
    _add_heading(doc, "Architecture")
    _add_body(
        doc,
        "Microservices on PostgreSQL (database-per-service), Redis cache, Traefik gateway. "
        "<100ms verification, 1000+ req/s. SSL.com certificate integration for automated key lifecycle.",
    )

    # Signing Pipeline
    _add_heading(doc, "Signing Pipeline")
    _add_numbered(
        doc, 1, " Text segmented at configurable granularity (sentence, paragraph, section). " "Merkle tree constructed. C2PA manifest generated."
    )
    _add_numbered(doc, 2, " Unicode variation selectors encode per-segment provenance markers " "(invisible, survive copy-paste and redistribution).")
    _add_numbered(doc, 3, " Rights profile snapshot attached. Coalition index updated. Signed text returned.")

    # Quick Integration
    _add_heading(doc, "Quick Integration")
    _add_code_block(
        doc,
        'POST /api/v1/sign  { "text": "...", "options": { "segmentation_level": "sentence" } }\n'
        "Response (201): signed_text, document_id, verification_url, merkle_root, total_segments",
        size=6.5,
    )

    # Feature Matrix
    _add_heading(doc, "Feature Matrix")
    _add_table(
        doc,
        ["Capability", "Free", "Enterprise"],
        [
            ["C2PA 2.3 signing (all 31 formats + live streams)", "Yes", "Yes"],
            ["Sentence-level Merkle trees + invisible provenance markers", "Yes", "Yes"],
            ["Rights management (bronze / silver / gold)", "Yes", "Yes"],
            ["Public verification API + WordPress + Chrome ext + SDKs", "Yes", "Yes"],
            ["Monthly signing limit", "1,000", "Unlimited"],
            ["Word-level segmentation", "-", "Yes"],
            ["Cross-org search, fuzzy matching, fingerprinting", "-", "Yes"],
            ["Court-ready evidence generation", "-", "Yes"],
            ["BYOK, batch (100+ docs/req), SSO/SAML, webhooks, SLA", "-", "Yes"],
        ],
        col_widths=[3.8, 0.9, 0.9],
        font_size=7,
    )

    # Key Endpoints
    _add_heading(doc, "Key Endpoints")
    _add_table(
        doc,
        ["Method", "Path", "Description"],
        [
            ["POST", "/api/v1/sign", "Sign content with C2PA manifest + provenance markers + rights"],
            ["POST", "/api/v1/public/verify", "Verify signed content (no auth required)"],
            ["GET", "/api/v1/public/rights/{doc_id}", "Public rights resolution (discoverable by AI crawlers)"],
            ["POST", "/api/v1/enterprise/evidence/generate", "Court-ready evidence package (Enterprise)"],
        ],
        col_widths=[0.6, 2.8, 3.2],
        font_size=7,
    )

    _add_footer(doc)

    path = os.path.join(OUTPUT_DIR, "Encypher_Publisher_Technical_v4.docx")
    doc.save(path)
    print(f"  Saved: {path}")
    return path


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------

if __name__ == "__main__":
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    print("Generating publisher 1-pagers...")
    build_business_outcome()
    build_technical()
    print("Done.")
