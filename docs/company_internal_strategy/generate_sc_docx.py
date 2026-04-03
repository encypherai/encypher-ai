#!/usr/bin/env python3
"""Generate branded DOCX for C2PA Steering Committee Application."""

import os

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls
from docx.shared import Cm, Inches, Pt, RGBColor

# Encypher brand colors
DELFT_BLUE = RGBColor(0x1B, 0x2F, 0x50)
BLUE_NCS = RGBColor(0x2A, 0x87, 0xC4)
COLUMBIA_BLUE = RGBColor(0xB7, 0xD5, 0xED)
NEUTRAL_GRAY = RGBColor(0xA7, 0xAF, 0xBC)
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
BODY_TEXT = RGBColor(0x33, 0x33, 0x33)

FONT_NAME = "Roboto"
FONT_MONO = "Roboto Mono"


def set_cell_shading(cell, hex_color):
    """Set cell background color."""
    shading = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}" w:val="clear"/>')
    cell._tc.get_or_add_tcPr().append(shading)


def add_styled_paragraph(
    doc, text, style_name="Body", bold=False, color=None, size=None, space_after=None, space_before=None, alignment=None, font_name=None
):
    """Add a paragraph with custom styling."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font_name or FONT_NAME
    run.font.size = Pt(size or 10.5)
    run.font.color.rgb = color or BODY_TEXT
    run.bold = bold
    if space_after is not None:
        p.paragraph_format.space_after = Pt(space_after)
    if space_before is not None:
        p.paragraph_format.space_before = Pt(space_before)
    if alignment is not None:
        p.alignment = alignment
    return p


def add_heading_styled(doc, text, level=1):
    """Add a branded heading."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.bold = True
    if level == 1:
        run.font.size = Pt(18)
        run.font.color.rgb = DELFT_BLUE
        p.paragraph_format.space_before = Pt(24)
        p.paragraph_format.space_after = Pt(8)
        # Add a thin blue line under heading
        pPr = p._p.get_or_add_pPr()
        pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}>' f'  <w:bottom w:val="single" w:sz="6" w:space="4" w:color="2A87C4"/>' f'</w:pBdr>')
        pPr.append(pBdr)
    elif level == 2:
        run.font.size = Pt(13)
        run.font.color.rgb = BLUE_NCS
        p.paragraph_format.space_before = Pt(16)
        p.paragraph_format.space_after = Pt(6)
    elif level == 3:
        run.font.size = Pt(11)
        run.font.color.rgb = DELFT_BLUE
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
    return p


def add_body(doc, text, space_after=6):
    """Add body text paragraph."""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = FONT_NAME
    run.font.size = Pt(10.5)
    run.font.color.rgb = BODY_TEXT
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.line_spacing = Pt(15)
    return p


def add_bullet(doc, text, bold_prefix=None):
    """Add a bullet point, optionally with a bold prefix."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.line_spacing = Pt(15)

    # Add bullet character
    bullet_run = p.add_run("\u2022  ")
    bullet_run.font.name = FONT_NAME
    bullet_run.font.size = Pt(10.5)
    bullet_run.font.color.rgb = BLUE_NCS

    if bold_prefix:
        bold_run = p.add_run(bold_prefix)
        bold_run.font.name = FONT_NAME
        bold_run.font.size = Pt(10.5)
        bold_run.font.color.rgb = DELFT_BLUE
        bold_run.bold = True
        body_run = p.add_run(text)
    else:
        body_run = p.add_run(text)
    body_run.font.name = FONT_NAME
    body_run.font.size = Pt(10.5)
    body_run.font.color.rgb = BODY_TEXT
    return p


def add_branded_table(doc, headers, rows):
    """Add a branded table."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Style header row
    for i, header in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = ""
        p = cell.paragraphs[0]
        run = p.add_run(header)
        run.font.name = FONT_NAME
        run.font.size = Pt(9.5)
        run.font.color.rgb = WHITE
        run.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        set_cell_shading(cell, "1B2F50")

    # Style data rows
    for r_idx, row in enumerate(rows):
        for c_idx, cell_text in enumerate(row):
            cell = table.rows[r_idx + 1].cells[c_idx]
            cell.text = ""
            p = cell.paragraphs[0]
            run = p.add_run(cell_text)
            run.font.name = FONT_NAME
            run.font.size = Pt(9.5)
            run.font.color.rgb = BODY_TEXT
            if r_idx % 2 == 1:
                set_cell_shading(cell, "EDF4F9")

    # Set borders
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(f'<w:tblPr {nsdecls("w")}/>')
    borders = parse_xml(
        f'<w:tblBorders {nsdecls("w")}>'
        f'  <w:top w:val="single" w:sz="4" w:space="0" w:color="B7D5ED"/>'
        f'  <w:left w:val="single" w:sz="4" w:space="0" w:color="B7D5ED"/>'
        f'  <w:bottom w:val="single" w:sz="4" w:space="0" w:color="B7D5ED"/>'
        f'  <w:right w:val="single" w:sz="4" w:space="0" w:color="B7D5ED"/>'
        f'  <w:insideH w:val="single" w:sz="4" w:space="0" w:color="B7D5ED"/>'
        f'  <w:insideV w:val="single" w:sz="4" w:space="0" w:color="B7D5ED"/>'
        f'</w:tblBorders>'
    )
    tblPr.append(borders)

    doc.add_paragraph()  # spacing
    return table


def add_info_field(doc, label, value):
    """Add a label: value pair for company info."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    label_run = p.add_run(f"{label}: ")
    label_run.font.name = FONT_NAME
    label_run.font.size = Pt(10.5)
    label_run.font.color.rgb = DELFT_BLUE
    label_run.bold = True
    value_run = p.add_run(value)
    value_run.font.name = FONT_NAME
    value_run.font.size = Pt(10.5)
    value_run.font.color.rgb = BODY_TEXT
    return p


def build_document():
    doc = Document()

    # Set default margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.5)
        section.right_margin = Cm(2.5)

    # --- COVER / HEADER ---
    # Try to add logo
    png_logo_path = os.path.join(os.path.dirname(__file__), "../../apps/marketing-site/public/encypher_full_nobg.png")
    if os.path.exists(png_logo_path):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run()
        run.add_picture(png_logo_path, width=Inches(2.2))
        p.paragraph_format.space_after = Pt(20)

    # Title
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("C2PA Steering Committee Application")
    run.font.name = FONT_NAME
    run.font.size = Pt(26)
    run.font.color.rgb = DELFT_BLUE
    run.bold = True
    p.paragraph_format.space_after = Pt(4)

    # Subtitle
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("Encypher Corporation")
    run.font.name = FONT_NAME
    run.font.size = Pt(16)
    run.font.color.rgb = BLUE_NCS
    p.paragraph_format.space_after = Pt(20)

    # Thin rule
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}>' f'  <w:bottom w:val="single" w:sz="8" w:space="1" w:color="2A87C4"/>' f'</w:pBdr>')
    pPr.append(pBdr)
    p.paragraph_format.space_after = Pt(16)

    # --- COMPANY INFORMATION ---
    add_heading_styled(doc, "Company Information", level=1)

    add_info_field(doc, "Legal Entity Name", "Encypher Corporation")
    add_info_field(doc, "Industry Served", "Content provenance infrastructure (publishing, media, AI)")
    add_info_field(doc, "Headquarters", "Dover, DE (fully remote team)")
    add_info_field(doc, "Primary Contact", "Erik Svilich, Founder & CEO, erik.svilich@encypher.com")
    add_info_field(
        doc,
        "Secondary Contact",
        "Eddan Katz, Policy Advisor (policy, governmental affairs, and C2PA evangelism through government and legislative channels)",
    )

    add_heading_styled(doc, "Organizational Overview", level=2)
    add_body(
        doc,
        "Encypher Corporation builds content provenance infrastructure on the C2PA standard. "
        "Erik Svilich, Encypher's founder, authored Section A.7 of C2PA 2.3, the text "
        "provenance specification published January 8, 2026, and co-chairs the C2PA Text "
        "Provenance Task Force. The company built the production reference implementation "
        "of that specification: an enterprise signing API covering every content MIME type "
        "in the C2PA specification outside of AI/ML models (31 formats across text, images, "
        "audio, video, documents, and fonts) with C2PA manifests, a WordPress plugin "
        "(PR #294) bringing text provenance to the web's dominant CMS layer, a Chrome "
        "verification extension, and an open-source reference implementation under AGPL-3.0. "
        "Encypher has submitted to the C2PA Conformance Program with 20 MIME types approved "
        "across generate and validate, and is working with the conformance committee to "
        "expand coverage to 11 additional formats including text and live video. The "
        "C2PA team includes Eddan Katz (policy advisor; AI governance, copyright law, former "
        "Meta C2PA policy support) and Madhav Chinnappa (strategic advisor; former Google "
        "News director, BBC News, Reuters Institute Fellow).",
    )

    # --- Q1 ---
    add_heading_styled(doc, "Q1. Why is your organization seeking a Steering Committee role?", level=1)
    add_body(
        doc, "C2PA faces two scaling challenges: available implementations of the specification " "and mass media adoption. Encypher addresses both."
    )
    add_body(
        doc,
        "Erik Svilich has contributed to C2PA weekly since June 2025, authoring Section A.7 "
        "of the specification, co-chairing the Text Provenance Task Force, and building "
        "the production reference implementation. That contribution has achieved conformance "
        "across 20 MIME types, and Encypher is now working with the conformance committee "
        "to expand the program into text, live video, and other formats where no conformance "
        "criteria currently exist.",
    )
    add_body(
        doc,
        "We are seeking a Steering Committee role to scale this contribution. SC membership "
        "would expand our involvement from specification and implementation into governance, "
        "subcommittee participation, and adoption strategy. We plan to grow our dedicated "
        "C2PA team with 1-2 full-time hires focused on standards development, policy "
        "engagement, and ecosystem coordination, building sustained institutional capacity "
        "to match the scope of SC responsibilities.",
    )
    add_body(
        doc,
        "The transition from published spec to deployed infrastructure is the phase where "
        "governance decisions determine whether C2PA's expansion beyond visual media "
        "succeeds. The organization that wrote the spec, built the broadest implementation, "
        "and is driving publisher adoption should be at the table.",
    )

    # --- Q2 ---
    add_heading_styled(doc, "Q2. How will your organization support the mission, Guiding Principles, and Charter of the C2PA?", level=1)
    add_body(
        doc,
        "We already do. Erik Svilich has attended C2PA working group sessions consistently "
        "since June 2025, 1-3 meetings per week, as Co-Chair of the Text Provenance Task "
        "Force. Our support for C2PA's mission is embedded in our business model, not a "
        "side activity:",
    )
    add_bullet(doc, " Authored Section A.7, extending C2PA into unstructured text.", "Specification development.")
    add_bullet(
        doc,
        " Open-source implementation (AGPL-3.0) ensures every developer starts with a standards-compliant foundation.",
        "Reference implementation.",
    )
    add_bullet(
        doc,
        " Our public verification API covers all media types at no cost, removing friction from the trust chain. Every verification strengthens the C2PA ecosystem.",
        "Free verification.",
    )
    add_bullet(
        doc,
        " Submitted formal public comment on the EU Code of Practice on AI-Generated Content, advocating mandatory C2PA signing.",
        "Regulatory advocacy.",
    )
    add_bullet(
        doc,
        " Educating news media organizations (NMA, SPUR, enterprise publishers) on C2PA and driving real-world implementation.",
        "Publisher education and adoption.",
    )
    add_body(doc, "SC membership would expand this work into governance, subcommittee participation, " "and cross-domain coordination.")

    # --- Q3 ---
    add_heading_styled(doc, "Q3. Describe your organization's current involvement in provenance, authenticity, or related standards.", level=1)

    add_heading_styled(doc, "Current C2PA involvement:", level=3)
    add_bullet(doc, " C2PA Text Provenance Task Force (active since June 2025)", "Co-Chair,")
    add_bullet(doc, " Section A.7 of C2PA 2.3 (published January 8, 2026)", "Specification author,")
    add_bullet(doc, " (1-3 meetings/week) since June 2025", "Consistent weekly attendance")
    add_bullet(
        doc,
        " participant (approved for 20 MIME types; working with conformance committee to expand program to 11 additional formats including text)",
        "Conformance Program",
    )

    add_heading_styled(doc, "Standards and policy engagement beyond C2PA:", level=3)
    add_bullet(
        doc,
        " on AI-Generated Content: formal public comment submitted, advocating mandatory C2PA signing for AI-generated text",
        "EU Code of Practice",
    )
    add_bullet(doc, " Presented to AI Licensing Working Group (February 2026)", "News Media Alliance (NMA):")
    add_bullet(doc, " Engaged on publisher standards convergence with BBC, FT, Guardian, Sky News, Telegraph", "SPUR Coalition (UK):")
    add_bullet(doc, " Standards stack alignment mapping C2PA alongside IAB Tech Lab, NMA, and AAM frameworks", "Scott Cunningham / AAM:")

    # --- Q4 ---
    add_heading_styled(doc, "Q4. Describe your technical expertise relevant to C2PA.", level=1)
    add_body(doc, "Our technical expertise comes from building the specification and implementing it " "at production scale:")

    add_heading_styled(doc, "Specification-level expertise:", level=3)
    add_bullet(doc, "Authored the C2PA text provenance specification (Section A.7), including manifest embedding architecture for unstructured text")
    add_bullet(
        doc, "Deep familiarity with JUMBF/COSE manifest structures, trust list management, certificate chain validation, and C2PA assertion schemas"
    )
    add_bullet(doc, "Live video stream signing per C2PA 2.3 Section 19 (per-segment manifests with backwards-linked provenance chain)")

    add_heading_styled(doc, "Implementation expertise:", level=3)
    add_bullet(
        doc,
        "Production signing across 31 MIME types in 6 media categories (text, images, audio, video, documents, fonts), every content MIME type in the C2PA specification outside of AI/ML models",
    )
    add_bullet(
        doc,
        "Two verification pipelines: c2pa-python (c2pa-rs) for natively supported formats and custom JUMBF/COSE structural verification for extended formats",
    )
    add_bullet(doc, "Perceptual hash (pHash) attribution search for fuzzy derivative matching across resized and reformatted image variants")

    add_heading_styled(doc, "Proprietary extensions (standard-compatible):", level=3)
    add_bullet(doc, "Sentence-level content authentication extending C2PA's document-level approach")
    add_bullet(doc, "Per-character attribution for granular provenance tracking")
    add_bullet(doc, "Cryptographic evidence generation for provenance verification")

    # --- Q5 ---
    add_heading_styled(doc, "Q5. What C2PA-related products, prototypes, or implementations have you built or deployed?", level=1)
    add_body(doc, "All items below are production systems, not prototypes:")
    add_bullet(
        doc,
        " Production API generating C2PA JUMBF manifests for 31 MIME types across text, images (13 formats), audio (6 formats), video (4 formats), documents (5 formats), and fonts (3 formats), covering every content MIME type in the C2PA specification outside of AI/ML models. Includes live video stream signing.",
        "Enterprise Signing API.",
    )
    add_bullet(doc, " Document-level C2PA text provenance under AGPL-3.0. Available on GitHub.", "Open-Source Reference Implementation (c2pa-text).")
    add_bullet(
        doc,
        " C2PA text provenance integration for the CMS powering 40%+ of the web. PR submitted to WordPress core and under review.",
        "WordPress Plugin (PR #294).",
    )
    add_bullet(doc, " Browser-level Content Credentials verification for text content. Publicly available.", "Chrome Verification Extension.")
    add_bullet(
        doc, " No-authentication verification endpoint covering all supported media types. Free for any third party.", "Public Verification API."
    )
    add_bullet(doc, " Signing management, provenance monitoring, and attribution analytics.", "Enterprise Dashboard.")
    add_body(
        doc,
        "Submitted to the C2PA Conformance Program. Conformance approved for 20 MIME types "
        "across generate and validate. Working with the conformance committee to expand "
        "coverage to 11 additional formats including text provenance and live video.",
    )

    # --- Q6 ---
    add_heading_styled(doc, "Q6. Have you submitted or plan to submit any products or services to C2PA's Conformance Program?", level=1)
    add_body(doc, "Yes. Encypher has submitted to the C2PA Conformance Program and completed the " "product review process.")
    add_bullet(doc, " 1", "Assurance Level:")
    add_bullet(doc, " Passed.", "Generator product security review:")
    add_bullet(doc, " Complete. Approved by Scott Perry.", "Product implementation review:")
    add_bullet(
        doc,
        " 11 image formats (JPEG, JXL, PNG, SVG, GIF, DNG, TIFF, WebP, HEIC, HEIF, AVIF), 3 video (AVI, MP4, MOV), 5 audio (MPA, MPEG, WAV, AAC, MP4), 1 document (PDF).",
        "Conformance-approved scope (20 MIME types, generate + validate):",
    )
    add_bullet(
        doc,
        " Encypher has implemented 11 additional MIME types that the conformance program "
        "does not yet cover, including live video and the Section A.7 text provenance "
        "specification. Scott Perry has asked Encypher to assist in providing conformance "
        "infrastructure for these formats, given Encypher's leadership on these media types, "
        "and will bring in Conformance Task Force leaders to support the effort.",
        "Expanding the program:",
    )

    # --- Q7 ---
    add_heading_styled(
        doc, "Q7. Please identify how your organization has materially impacted the technical and/or advocacy goals of the C2PA.", level=1
    )

    add_heading_styled(doc, "Technical impact:", level=2)
    add_body(
        doc,
        'Erik Svilich authored Section A.7 of C2PA 2.3, "Embedding Manifests into '
        'Unstructured Text," published January 8, 2026. Text provenance had been an open '
        "problem within C2PA for two years before Erik proposed and delivered the solution. "
        "This extended C2PA's provenance framework from visual media into text for the "
        "first time. As Co-Chair of the Text Provenance Task Force, Erik led the "
        "specification development process through publication and continues to lead the "
        "working group through the adoption phase.",
    )
    add_body(
        doc,
        "Encypher built the production reference implementation of the specification, "
        "providing the starting point for developers implementing C2PA text provenance. "
        "Encypher's proprietary extensions add sentence-level authentication on top of "
        "C2PA's document-level approach, contributing capabilities that strengthen the "
        "ecosystem while remaining standard-compatible.",
    )

    add_heading_styled(doc, "Advocacy impact:", level=2)
    add_bullet(
        doc,
        "Submitted formal public comment on the EU Code of Practice on AI-Generated Content, advocating mandatory C2PA signing for AI-generated text content.",
    )
    add_bullet(
        doc,
        "Presented to the News Media Alliance AI Licensing Working Group (February 2026), connecting C2PA text provenance to publisher adoption pathways.",
    )
    add_bullet(
        doc,
        "Educating publisher organizations (NMA, SPUR, enterprise publishers) on C2PA and translating the specification into adoption infrastructure.",
    )

    # --- Q8 ---
    add_heading_styled(
        doc,
        "Q8. Please identify material contributions from your organization that were accepted into open-source software or standards development organizations.",
        level=1,
    )

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("1.  ")
    r.font.name = FONT_NAME
    r.font.size = Pt(10.5)
    r.font.color.rgb = BLUE_NCS
    r.bold = True
    b = p.add_run("C2PA Specification 2.3, Section A.7: ")
    b.font.name = FONT_NAME
    b.font.size = Pt(10.5)
    b.font.color.rgb = DELFT_BLUE
    b.bold = True
    t = p.add_run('"Embedding Manifests into Unstructured Text." Authored by Erik Svilich. Accepted and published January 8, 2026.')
    t.font.name = FONT_NAME
    t.font.size = Pt(10.5)
    t.font.color.rgb = BODY_TEXT

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("2.  ")
    r.font.name = FONT_NAME
    r.font.size = Pt(10.5)
    r.font.color.rgb = BLUE_NCS
    r.bold = True
    b = p.add_run("c2pa-text: ")
    b.font.name = FONT_NAME
    b.font.size = Pt(10.5)
    b.font.color.rgb = DELFT_BLUE
    b.bold = True
    t = p.add_run("Open-source C2PA text provenance reference implementation. AGPL-3.0.")
    t.font.name = FONT_NAME
    t.font.size = Pt(10.5)
    t.font.color.rgb = BODY_TEXT

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("3.  ")
    r.font.name = FONT_NAME
    r.font.size = Pt(10.5)
    r.font.color.rgb = BLUE_NCS
    r.bold = True
    b = p.add_run("WordPress PR #294: ")
    b.font.name = FONT_NAME
    b.font.size = Pt(10.5)
    b.font.color.rgb = DELFT_BLUE
    b.bold = True
    t = p.add_run("C2PA text provenance integration submitted to WordPress core.")
    t.font.name = FONT_NAME
    t.font.size = Pt(10.5)
    t.font.color.rgb = BODY_TEXT

    # --- Q9 ---
    add_heading_styled(
        doc,
        "Q9. Please identify up to 3 experts from your organization who will be collaborating to address the problems identified in C2PA's guiding documents.",
        level=1,
    )
    add_body(doc, "(100 words each)", space_after=10)

    # Erik
    add_heading_styled(doc, "Erik Svilich - Founder & CEO, Encypher Corporation", level=2)
    add_body(
        doc,
        "Author of C2PA 2.3 Section A.7, the text provenance specification published "
        "January 8, 2026. Co-Chair, C2PA Text Provenance Task Force. Built production "
        "reference implementation including enterprise signing API, WordPress integration "
        "(PR #294), and Chrome verification extension. Attends 1-3 C2PA meetings weekly. "
        "Submitted formal public comment on EU Code of Practice on AI-Generated Content "
        "advocating mandatory C2PA signing. Full decision-making authority as CEO. Covers "
        "Steering Committee, Text Provenance Task Force, and technical working groups.",
    )

    # Eddan
    add_heading_styled(doc, "Eddan Katz - Policy Advisor, Encypher Corporation", level=2)
    add_body(
        doc,
        "Research Fellow at UC Law SF focusing on copyright and AI Agent liability, the "
        "legal questions at the center of C2PA's mission. Directly supported Meta's C2PA "
        "engagement on the AI Policy team, conducting policy research for the company's "
        "standards negotiations. First Executive Director of Yale Law School's Information "
        "Society Project. International Affairs Director at the Electronic Frontier "
        "Foundation. AI Governance Platform Lead at the World Economic Forum's Centre for "
        "the Fourth Industrial Revolution. Co-founder, TechnoEthos (responsible AI "
        "governance). Covers C2PA's Governmental Affairs and Communications subcommittees.",
    )

    # Madhav
    add_heading_styled(doc, "Madhav Chinnappa - Strategic Advisor, Encypher Corporation", level=2)
    add_body(
        doc,
        "Reuters Institute Visiting Fellow (Oxford) and Future Insights Board member at "
        "Mediahuis. Former Director of News Ecosystem Development at Google (13 years), "
        "leading global publisher partnerships, innovation programs, and sustainability "
        "funds including the Digital News Innovation Fund and Google News Initiative. Former "
        "Head of Development & Rights at BBC News. Former Deputy Director at Associated "
        "Press Television. Three decades building publisher-technology ecosystem "
        "infrastructure globally. Expertise in content rights, news sustainability, and the "
        "publisher adoption dynamics critical to scaling C2PA text provenance. Advises on "
        "publisher ecosystem adoption workstreams.",
    )

    # --- Q10 ---
    add_heading_styled(doc, "Q10. Describe your organization's experience with standards bodies or multi-stakeholder governance.", level=1)

    add_heading_styled(doc, "Encypher (organizational):", level=3)
    add_bullet(doc, "C2PA Text Provenance Task Force: Co-Chair, active weekly participant")
    add_bullet(doc, "EU Code of Practice on AI-Generated Content: formal public comment")
    add_bullet(doc, "News Media Alliance: AI Licensing Working Group presenter and ongoing engagement")
    add_bullet(doc, "SPUR Coalition (UK): Engaged on multi-stakeholder publisher standards convergence")

    add_heading_styled(doc, "Team experience:", level=3)
    add_bullet(
        doc,
        " World Economic Forum global AI ethics initiatives. United Nations digital rights coalition representation. Meta standards negotiations support. Credo AI regulatory governance frameworks.",
        "Eddan Katz:",
    )
    add_bullet(
        doc,
        " 13 years leading Google's global publisher partnership and standards programs (Digital News Innovation Fund, Google News Initiative). BBC content rights frameworks. Reuters Institute research fellowship.",
        "Madhav Chinnappa:",
    )

    add_body(
        doc,
        "The team brings experience across technical standards development (C2PA), "
        "regulatory engagement (EU, UN), industry coalition governance (NMA, SPUR), and "
        "platform-publisher standards infrastructure (Google, BBC).",
    )

    # --- Q11 ---
    add_heading_styled(doc, "Q11. What resources do you plan to commit to the Steering Committee, Working Groups, and Task Forces?", level=1)
    add_body(
        doc,
        "Erik Svilich currently attends 1-3 C2PA meetings per week as Co-Chair of the "
        "Text Provenance Task Force and active working group participant. He has maintained "
        "this cadence since June 2025, contributing the specification authorship, reference "
        "implementation, conformance submission, and publisher adoption work described "
        "elsewhere in this application. The SC seat scales a contribution that is already "
        "sustained, not a new obligation.",
    )

    add_heading_styled(doc, "C2PA team structure with SC membership:", level=2)
    add_bullet(
        doc,
        " Primary representative. Steering Committee meetings, Text Provenance Task Force "
        "(Co-Chair, continuing), and technical working groups. 1-3 meetings/week currently; "
        "SC participation adds to this baseline. Full decision-making authority.",
        "Erik Svilich (CEO):",
    )
    add_bullet(
        doc,
        " Governmental Affairs and Communications subcommittees. Part-time dedicated "
        "commitment covering C2PA's policy and regulatory workstreams. Expertise in EU AI "
        "Act, US state AI legislation, and multi-stakeholder governance. SC membership "
        "formalizes a dedicated C2PA regulatory role within Encypher.",
        "Eddan Katz (Policy Advisor):",
    )
    add_bullet(
        doc,
        " Publisher ecosystem adoption workstreams. Available for specific initiatives, "
        "briefings, and working sessions where publisher-technology ecosystem expertise is "
        "required. Three decades of relationships across BBC, Google, AP, and the global "
        "news industry inform Encypher's approach to scaling C2PA text provenance adoption.",
        "Madhav Chinnappa (Strategic Advisor):",
    )

    add_body(
        doc,
        "The team covers three lanes: technical governance and standards development (Erik), "
        "policy, regulation, and governmental affairs (Eddan), and publisher ecosystem "
        "adoption and industry relationships (Madhav).",
    )
    add_body(
        doc,
        "Encypher is actively hiring additional technical team members who will dedicate "
        "hours to C2PA standards development and implementation work. SC membership supports "
        "this expansion: the governance commitment justifies 1-2 full-time roles focused on "
        "standards development, policy engagement, and ecosystem coordination, building "
        "sustained institutional capacity beyond the founding team's current allocation.",
    )

    # --- Q12 ---
    add_heading_styled(doc, "Q12. Describe your organization's reach (users, customers, partners).", level=1)
    add_body(
        doc,
        "Encypher's reach operates at infrastructure integration points rather than direct "
        "consumer endpoints. Our implementation is positioned at the CMS layer, the ad-tech "
        "layer, and the publisher platform layer, creating distribution leverage through "
        "integration rather than audience scale.",
    )

    add_heading_styled(doc, "Actively engaged:", level=2)
    add_bullet(doc, " C2PA text provenance plugin submitted as PR #294. WordPress powers 40%+ of websites globally.", "WordPress.org:")
    add_bullet(doc, " C2PA integration for ad-tech supply chain provenance. Prebid's ecosystem spans 70K+ publisher sites.", "Prebid:")
    add_bullet(doc, " Engaged with UK publishers (BBC, FT, Guardian, Sky News, Telegraph) on publisher standards convergence.", "SPUR Coalition:")
    add_bullet(
        doc,
        " Standards stack alignment mapping C2PA alongside IAB Tech Lab, NMA, and AAM frameworks. Scott founded both IAB Tech Lab and TAG.",
        "Scott Cunningham / IAB Tech Lab / AAM:",
    )
    add_bullet(
        doc,
        " Marketplace partnership agreement in legal review. Network of 700+ publishers including Reuters, Fortune, and Al Jazeera.",
        "Freestar / PubOS:",
    )
    add_bullet(doc, " C2PA education and standards awareness with US news publisher membership.", "NMA (News Media Alliance):")
    add_bullet(doc, " Engaged as fellow C2PA member on standards alignment.", "River Valley Technology:")

    add_heading_styled(doc, "In active discussion or initial pilot evaluation:", level=2)
    add_bullet(doc, " Evaluating CMS integration, audio provenance, and rights-aware marking across their media portfolio.", "NPR:")
    add_bullet(doc, " Enterprise WordPress hosting layer.", "Automattic / WordPress VIP:")
    add_bullet(doc, " Content provenance for global wire distribution.", "Associated Press (Paul Sarkis):")
    add_bullet(doc, " Digital media portfolio evaluation.", "Valnet:")
    add_bullet(doc, " Policy and government credibility channel.", "Trust in Media (Ellen McCarthy, former Asst. Secretary of State):")
    add_bullet(doc, " Academic publishing provenance.", "Taylor & Francis:")
    add_bullet(doc, " Legal infrastructure and regulatory compliance (Magic Circle firm).", "Linklaters:")
    add_bullet(doc, " C2PA education for local/regional news publishers.", "Local Media Alliance:")

    add_heading_styled(doc, "Open ecosystem:", level=2)
    add_bullet(doc, " publicly available, AGPL-3.0.", "Open-source reference implementation (c2pa-text):")
    add_bullet(doc, " publicly available.", "Chrome verification extension:")
    add_bullet(doc, " free, no authentication required, all media types.", "Public verification API:")

    # --- Q13 ---
    add_heading_styled(doc, "Q13. How would SC membership enable broader adoption of C2PA?", level=1)
    add_body(doc, "C2PA's impact scales with adoption. SC membership would accelerate that adoption across three fronts:")

    add_heading_styled(doc, "C2PA education and media industry adoption.", level=3)
    add_body(
        doc,
        "Encypher actively educates the publisher and news media ecosystem on C2PA and "
        "drives adoption. We have presented to the News Media Alliance AI Licensing Working "
        "Group, engaged with the SPUR Coalition (BBC, FT, Guardian, Sky News, Telegraph) on "
        "standards awareness, briefed enterprise media organizations (AP, NPR, Taylor & "
        "Francis) on C2PA implementation pathways, and built platform integrations (Freestar, "
        "700+ publishers) that bring Content Credentials to publishers at scale. SC "
        "membership would connect this adoption work directly to governance, ensuring the "
        "C2PA roadmap reflects implementation realities across the media industry.",
    )

    add_heading_styled(doc, "Implementation coverage.", level=3)
    add_body(
        doc,
        "Large portions of the C2PA specification lack any conforming implementation. "
        "Encypher is closing this gap across text, images, audio, video, live video, and "
        "document formats. SC membership would help coordinate implementation priorities "
        "with adoption strategy, ensuring the spec coverage expands where adoption demand "
        "is strongest.",
    )

    add_heading_styled(doc, "CMS-layer distribution.", level=3)
    add_body(
        doc,
        "The WordPress integration path could make C2PA Content Credentials available to "
        "40%+ of the web through a single plugin. SC coordination would align this "
        "distribution with C2PA's broader adoption strategy and conformance requirements.",
    )

    # --- Q14 ---
    add_heading_styled(doc, "Q14. What unique ecosystem leverage do you bring?", level=1)
    add_body(
        doc,
        "Encypher is the only organization that combines C2PA specification authorship, "
        "production reference implementation, and active publisher adoption relationships.",
    )
    add_body(doc, "No other applicant connects these three layers simultaneously:")

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("1.  ")
    r.font.name = FONT_NAME
    r.font.size = Pt(10.5)
    r.font.color.rgb = BLUE_NCS
    r.bold = True
    b = p.add_run("Standards authority. ")
    b.font.name = FONT_NAME
    b.font.size = Pt(10.5)
    b.font.color.rgb = DELFT_BLUE
    b.bold = True
    t = p.add_run("Authored Section A.7, solving a problem C2PA had been working on for two years. Co-Chair of the task force that developed it.")
    t.font.name = FONT_NAME
    t.font.size = Pt(10.5)
    t.font.color.rgb = BODY_TEXT

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("2.  ")
    r.font.name = FONT_NAME
    r.font.size = Pt(10.5)
    r.font.color.rgb = BLUE_NCS
    r.bold = True
    b = p.add_run("Implementation infrastructure. ")
    b.font.name = FONT_NAME
    b.font.size = Pt(10.5)
    b.font.color.rgb = DELFT_BLUE
    b.bold = True
    t = p.add_run(
        "Production signing API, open-source reference implementation, WordPress integration, and conformance program participation. Developers implementing C2PA text provenance start with our code."
    )
    t.font.name = FONT_NAME
    t.font.size = Pt(10.5)
    t.font.color.rgb = BODY_TEXT

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run("3.  ")
    r.font.name = FONT_NAME
    r.font.size = Pt(10.5)
    r.font.color.rgb = BLUE_NCS
    r.bold = True
    b = p.add_run("Adoption ecosystem. ")
    b.font.name = FONT_NAME
    b.font.size = Pt(10.5)
    b.font.color.rgb = DELFT_BLUE
    b.bold = True
    t = p.add_run(
        "Active C2PA education and adoption work across media organizations (NMA, SPUR, AP, NPR, Freestar), platform integration points (WordPress, Prebid), and standards stack alignment (C2PA + IAB Tech Lab + NMA + AAM) identified through collaboration with Scott Cunningham, founder of IAB Tech Lab and TAG."
    )
    t.font.name = FONT_NAME
    t.font.size = Pt(10.5)
    t.font.color.rgb = BODY_TEXT

    add_body(
        doc,
        "This combination means SC decisions about text provenance would be informed by "
        "someone who wrote the spec, built the implementation, and works daily with the "
        "publishers who will adopt it.",
    )

    # --- Q15 ---
    add_heading_styled(doc, "Q15. What gaps or opportunities do you believe C2PA should address in the next 12-24 months?", level=1)
    add_body(doc, "Three areas where C2PA governance can shape outcomes over the next 12-24 months:")

    add_heading_styled(doc, "1. Text-specific conformance criteria.", level=3)
    add_body(
        doc,
        "The Conformance Program was designed around visual media. Text provenance has "
        "different security surfaces: invisible character stripping by platforms, copy-paste "
        "transformation, CMS normalization, and format conversion. C2PA should develop "
        "text-specific conformance test suites that account for these realities. The "
        "reference implementation and our conformance submission provide a foundation for "
        "defining those criteria.",
    )

    add_heading_styled(doc, "2. AI ingestion-time provenance framework.", level=3)
    add_body(
        doc,
        "C2PA Content Credentials can currently be verified on the open web. The next "
        "adoption milestone is AI systems recognizing and preserving C2PA provenance signals "
        "during data ingestion. This requires a collaborative framework between content "
        "creators and AI platforms, not a technical specification alone. Multiple SC members "
        "operate infrastructure where this framework would apply. SC coordination is the "
        "right venue to develop it collaboratively.",
    )

    add_heading_styled(doc, "3. EU AI Act content transparency.", level=3)
    add_body(
        doc,
        "The enforcement deadline is August 2, 2026. The draft Code of Practice explicitly "
        "requires machine-readable marking and watermarking with a multilayered approach. "
        "C2PA should position proactively as compliance infrastructure for text content, not "
        "wait for regulatory interpretation to settle. Encypher has already engaged this "
        "process through formal public comment.",
    )

    # --- Q16 ---
    add_heading_styled(doc, "Q16. Links to relevant products, demos, documentation, or public statements.", level=1)

    add_branded_table(
        doc,
        ["Resource", "Link", "Verifies"],
        [
            [
                "C2PA 2.3 Section A.7",
                "https://spec.c2pa.org/specifications/specifications/2.3/specs/C2PA_Specification.html#embedding_manifests_into_unstructured_text",
                "Spec authorship (Q7)",
            ],
            ["Open-source reference implementation (c2pa-text)", "[GitHub repo URL]", "OSS contribution (Q8)"],
            ["WordPress PR #294", "[GitHub PR URL]", "CMS integration (Q5, Q8)"],
            ["EU Code of Practice public comment", "[Link to submitted comment]", "Regulatory advocacy (Q7)"],
            ["Chrome verification extension", "[Chrome Web Store URL]", "Verification tooling (Q5)"],
            ["Enterprise API documentation", "[docs URL]", "Production implementation (Q5)"],
            ["Company website", "https://encypher.com", "Organizational overview"],
        ],
    )

    # --- Q17 ---
    add_heading_styled(doc, "Q17. Ownership disclosure.", level=1)
    add_body(
        doc,
        "Encypher Corporation is a privately held Delaware corporation. The company is "
        "founder-controlled. Erik Svilich (CEO) holds majority equity and full "
        "decision-making authority. No external entity or government directs or controls "
        "the management or decisions of the organization.",
    )

    # --- Footer ---
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(30)
    pPr = p._p.get_or_add_pPr()
    pBdr = parse_xml(f'<w:pBdr {nsdecls("w")}>' f'  <w:top w:val="single" w:sz="6" w:space="4" w:color="B7D5ED"/>' f'</w:pBdr>')
    pPr.append(pBdr)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run("Encypher Corporation  |  Dover, DE  |  encypher.com  |  erik.svilich@encypher.com")
    run.font.name = FONT_NAME
    run.font.size = Pt(8.5)
    run.font.color.rgb = NEUTRAL_GRAY

    # Save
    output_path = os.path.join(os.path.dirname(__file__), "C2PA_Steering_Committee_Application_Encypher.docx")
    doc.save(output_path)
    print(f"Saved: {output_path}")
    return output_path


if __name__ == "__main__":
    build_document()
