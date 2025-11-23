"""
Script to create tampered example files with EncypherAI metadata for testing purposes.
This script creates examples with sophisticated tampering by modifying variation selectors
in the embedded metadata, which is more subtle than adding obvious text markers.
"""

import sys
from pathlib import Path

# Add the parent directory to the path so we can import the shared library
sys.path.append(str(Path(__file__).parent.parent.parent))

try:
    # Import from the core EncypherAI package as used in shared_commercial_libs
    from encypher.core.unicode_metadata import UnicodeMetadata, MetadataTarget
    from encypher.core.keys import load_public_key_from_data, load_private_key_from_data
    
    # Import from our shared commercial library
    from encypher_commercial_shared import EncypherAI
    
    from rich.console import Console
    console = Console()
    
    # Directory to save example files
    EXAMPLE_DIR = Path(__file__).parent / "tampered_examples"
    EXAMPLE_DIR.mkdir(exist_ok=True)
    
    # Create example key files for testing
    KEY_DIR = EXAMPLE_DIR / "keys"
    KEY_DIR.mkdir(exist_ok=True)
    
    # Generate a test private/public key pair for signing
    from cryptography.hazmat.primitives.asymmetric import ed25519
    from cryptography.hazmat.primitives import serialization
    
    def generate_key_pair(private_key_path, public_key_path):
        """Generate a test Ed25519 key pair for signing and verification."""
        # Generate a private key
        private_key = ed25519.Ed25519PrivateKey.generate()
        
        # Serialize private key to PEM format
        private_pem = private_key.private_bytes(
            encoding=serialization.Encoding.PEM,
            format=serialization.PrivateFormat.PKCS8,
            encryption_algorithm=serialization.NoEncryption()
        )
        
        # Serialize public key to PEM format
        public_key = private_key.public_key()
        public_pem = public_key.public_bytes(
            encoding=serialization.Encoding.PEM,
            format=serialization.PublicFormat.SubjectPublicKeyInfo
        )
        
        # Write keys to files
        with open(private_key_path, 'wb') as f:
            f.write(private_pem)
        
        with open(public_key_path, 'wb') as f:
            f.write(public_pem)
        
        return private_key, public_key
    
    # Generate keys for test-signer
    test_signer_private_key_path = KEY_DIR / "test-signer-private.pem"
    test_signer_public_key_path = KEY_DIR / "test-signer-public.pem"
    generate_key_pair(test_signer_private_key_path, test_signer_public_key_path)
    console.print(f"[green]Generated test signer keys at {KEY_DIR}[/green]")
    
    # Create EncypherAI instance with the test signer
    ea_test = EncypherAI(
        private_key_path=str(test_signer_private_key_path),
        public_key_path=str(test_signer_public_key_path),
        signer_id="test-signer",
        verbose=True
    )
    
    # We'll use UnicodeMetadata directly for more control over the embedding process
    # First, load the private key
    with open(test_signer_private_key_path, 'rb') as f:
        private_key_data = f.read()
    
    # Load the public key
    with open(test_signer_public_key_path, 'rb') as f:
        public_key_data = f.read()
    
    # Create a UnicodeMetadata instance directly
    um = UnicodeMetadata(
        private_key_data=private_key_data,
        signer_id="test-signer",
        verbose=True
    )
    
    # We'll use UnicodeMetadata directly instead of modifying EncypherAI methods
    
    # Create example files with metadata
    
    # Example 1: Create a file with valid metadata
    example_text = """This is an example file with valid EncypherAI metadata.
The metadata is embedded in this text file and can be verified
using the test-signer public key.

This file should pass verification when scanned by the audit-log-cli tool.
"""
    
    # Embed metadata
    try:
        from datetime import datetime, timezone
        current_time = datetime.now(timezone.utc)
        
        custom_metadata = {
            "purpose": "testing", 
            "example": "tampered_variation_selectors",
            "document_id": "doc-vs-123",
            "version": "1.0.0"
        }
        
        # Create original file with metadata using UnicodeMetadata directly
        original_with_metadata = um.embed_metadata(
            text=example_text,
            private_key=load_private_key_from_data(private_key_data),
            signer_id="test-signer",
            custom_metadata=custom_metadata,
            model_id="test-model-vs",
            metadata_format="basic"
            # timestamp is handled automatically
        )
        
        # Save the original file for reference
        original_path = EXAMPLE_DIR / "original.txt"
        with open(original_path, 'w', encoding='utf-8') as f:
            f.write(original_with_metadata)
        console.print(f"[green]Created original example file: {original_path}[/green]")
        
        # Now create a tampered version by modifying variation selectors in the metadata
        # This is a subtle tampering that won't be visible to the naked eye
        # but will break the signature verification
        
        # Find the metadata in the text (it's embedded as zero-width characters)
        # For this example, we'll modify some of the variation selectors (U+FE0F)
        # by replacing them with different variation selectors (U+FE0E)
        
        # Convert to a list of characters for easier manipulation
        chars = list(original_with_metadata)
        
        # Find and replace some variation selectors
        # U+FE0F (emoji style) to U+FE0E (text style)
        vs_emoji = '\uFE0F'
        vs_text = '\uFE0E'
        
        # Count and replace a few instances to simulate tampering
        count = 0
        max_replacements = 3  # Limit replacements to avoid breaking the format completely
        
        for i in range(len(chars)):
            if chars[i] == vs_emoji and count < max_replacements:
                chars[i] = vs_text
                count += 1
        
        # Reconstruct the tampered text
        tampered_text = ''.join(chars)
        
        # Save the tampered file
        tampered_path = EXAMPLE_DIR / "tampered_variation_selectors.txt"
        with open(tampered_path, 'w', encoding='utf-8') as f:
            f.write(tampered_text)
        
        console.print(f"[green]Created tampered example with modified variation selectors: {tampered_path}[/green]")
        console.print("[yellow]This tampering is invisible to the naked eye but breaks signature verification[/yellow]")
        
        # Create PDF example with the same approach
        try:
            from fpdf import FPDF
            
            # Create PDF with original content
            pdf = FPDF()
            pdf.add_page()
            pdf.set_font("Arial", size=12)
            
            # Split text into lines and add to PDF
            for line in example_text.split('\n'):
                pdf.cell(200, 10, txt=line, ln=True)
            
            # Save original PDF
            original_pdf_path = EXAMPLE_DIR / "original.pdf"
            pdf.output(str(original_pdf_path))
            
            # Now create a PDF with the tampered text
            # For PDF, we'll need to extract the text, tamper with it, and then create a new PDF
            
            # Create a new PDF with the tampered text
            pdf_tampered = FPDF()
            pdf_tampered.add_page()
            pdf_tampered.set_font("Arial", size=12)
            
            # Split tampered text into lines and add to PDF
            for line in tampered_text.split('\n'):
                pdf_tampered.cell(200, 10, txt=line, ln=True)
            
            # Save tampered PDF
            tampered_pdf_path = EXAMPLE_DIR / "tampered_variation_selectors.pdf"
            pdf_tampered.output(str(tampered_pdf_path))
            
            console.print("[green]Created original and tampered PDF examples[/green]")
            
        except ImportError:
            console.print("[yellow]FPDF not installed, skipping PDF example creation[/yellow]")
            console.print("[yellow]Install with: uv add fpdf[/yellow]")
        
        # Create DOCX example with the same approach
        try:
            from docx import Document
            
            # Create DOCX with original content
            doc = Document()
            for line in example_text.split('\n'):
                if line.strip():
                    doc.add_paragraph(line)
                else:
                    doc.add_paragraph()
            
            # Save original DOCX
            original_docx_path = EXAMPLE_DIR / "original.docx"
            doc.save(str(original_docx_path))
            
            # Create a new DOCX with the tampered text
            doc_tampered = Document()
            for line in tampered_text.split('\n'):
                if line.strip():
                    doc_tampered.add_paragraph(line)
                else:
                    doc_tampered.add_paragraph()
            
            # Save tampered DOCX
            tampered_docx_path = EXAMPLE_DIR / "tampered_variation_selectors.docx"
            doc_tampered.save(str(tampered_docx_path))
            
            console.print("[green]Created original and tampered DOCX examples[/green]")
            
        except ImportError:
            console.print("[yellow]python-docx not installed, skipping DOCX example creation[/yellow]")
            console.print("[yellow]Install with: uv add python-docx[/yellow]")
        
    except Exception as e:
        console.print(f"[red]Error creating examples: {e}[/red]")
    
    console.print("[bold green]Successfully created example files with variation selector tampering![/bold green]")
    console.print("\nTo test these files with the audit-log-cli tool, run:")
    console.print("uv run -- python -m audit_log_cli.app.main --target audit_log_cli/tests/tampered_examples --trusted-signers audit_log_cli/tests/tampered_examples/keys --verify-content-integrity")

except ImportError as e:
    print(f"Error importing required modules: {e}")
    print("Make sure you have installed all the required dependencies.")
    sys.exit(1)
except Exception as e:
    print(f"Unexpected error: {e}")
    sys.exit(1)
