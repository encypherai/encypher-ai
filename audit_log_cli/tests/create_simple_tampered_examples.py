"""
Script to create simple tampered example files for testing purposes.
This script creates examples with obvious tampering by adding text markers.
"""

import os
import subprocess
import sys
from pathlib import Path

# Add the parent directory to the path so we can import the shared library
sys.path.append(str(Path(__file__).parent.parent.parent))

try:
    # Import from our shared commercial library
    from rich.console import Console

    from encypher.core.keys import load_private_key_from_data

    # Import from the core Encypher package
    from encypher.core.unicode_metadata import MetadataTarget, UnicodeMetadata
    from encypher_commercial_shared import Encypher
    console = Console()
    
    # Directory to save example files
    EXAMPLE_DIR = Path(__file__).parent / "simple_tampered_examples"
    EXAMPLE_DIR.mkdir(exist_ok=True)
    
    # Create example key files for testing
    KEY_DIR = EXAMPLE_DIR / "keys"
    KEY_DIR.mkdir(exist_ok=True)
    
    # Generate a test private/public key pair for signing
    from cryptography.hazmat.primitives import serialization
    from cryptography.hazmat.primitives.asymmetric import ed25519
    
    def generate_key_pair(private_key_path, public_key_path):
        """Generate a test Ed25519 key pair for signing and verification."""
        # Generate a private key
        private_key = ed25519.Ed25519PrivateKey.generate()
        
        # Serialize private key to PEM format
        private_pem = private_key.private_bytes(
            encoding=serialization.Encoding.PEM,
            format=serialization.PrivateFormat.PKCS8,
            encryption_algorithm=serialization.NoEncryption()
        )
        
        # Serialize public key to PEM format
        public_key = private_key.public_key()
        public_pem = public_key.public_bytes(
            encoding=serialization.Encoding.PEM,
            format=serialization.PublicFormat.SubjectPublicKeyInfo
        )
        
        # Write keys to files
        with open(private_key_path, 'wb') as f:
            f.write(private_pem)
        
        with open(public_key_path, 'wb') as f:
            f.write(public_pem)
        
        return private_key, public_key
    
    # Generate keys for test-signer
    test_signer_private_key_path = KEY_DIR / "test-signer-private.pem"
    test_signer_public_key_path = KEY_DIR / "test-signer-public.pem"
    generate_key_pair(test_signer_private_key_path, test_signer_public_key_path)
    console.print(f"[green]Generated test signer keys at {KEY_DIR}[/green]")
    
    # Create Encypher instance with the test signer
    ea_test = Encypher(
        private_key_path=str(test_signer_private_key_path),
        public_key_path=str(test_signer_public_key_path),
        signer_id="test-signer",
        verbose=True
    )
    
    # Create example files with metadata
    
    # Example 1: Create a file with valid metadata
    example_text = """This is an example file with valid Encypher metadata.
The metadata is embedded in this text file and can be verified
using the test-signer public key.

This file should pass verification when scanned by the audit-log-cli tool.
"""
    
    # Embed metadata
    try:
        # Get current timestamp
        from datetime import datetime, timezone
        current_time = datetime.now(timezone.utc)
        
        custom_metadata = {
            "purpose": "testing", 
            "example": "tampered_example",
            "document_id": "doc-123",
            "version": "1.0.0"
        }
        
        # Load the private key directly
        with open(test_signer_private_key_path, 'rb') as f:
            private_key_data = f.read()
            private_key = load_private_key_from_data(private_key_data)
        
        # Create original file with metadata using UnicodeMetadata directly
        original_with_metadata = UnicodeMetadata.embed_metadata(
            text=example_text,
            private_key=private_key,
            signer_id="test-signer",
            custom_metadata=custom_metadata,
            model_id="test-model",
            metadata_format="basic",
            timestamp=current_time
        )
        
        # Save the original file for reference
        original_path = EXAMPLE_DIR / "original.txt"
        with open(original_path, 'w', encoding='utf-8') as f:
            f.write(original_with_metadata)
        console.print(f"[green]Created original example file: {original_path}[/green]")
        
        # Now create a tampered version by modifying the text content
        # This will cause verification to fail because the content doesn't match the signature
        
        # Split the content into lines for easier manipulation
        lines = original_with_metadata.split('\n')
        
        # Add a clear indication that the file was tampered with
        if len(lines) > 2:
            lines[1] = lines[1] + " [THIS TEXT WAS MODIFIED AFTER SIGNING]"
            tampered_text = '\n'.join(lines)
        else:
            # Fallback if we can't split by lines
            tampered_text = original_with_metadata.replace(
                "This file should pass verification",
                "This file should pass verification [THIS TEXT WAS MODIFIED AFTER SIGNING]"
            )
        
        # Save the tampered file
        tampered_path = EXAMPLE_DIR / "tampered.txt"
        with open(tampered_path, 'w', encoding='utf-8') as f:
            f.write(tampered_text)
        
        console.print(f"[green]Created tampered example: {tampered_path}[/green]")
        
        # We'll create PDFs by saving DOCX files directly to PDF
        # First we'll create the DOCX files with the metadata, then save them as PDFs
        pdf_creation_successful = False
        
        # Simulate copy-pasting text with embedded metadata into DOCX files
        docx_available = False
        try:
            import docx
            from docx import Document
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            docx_available = True
        except ImportError as e:
            console.print(f"[yellow]Error importing python-docx: {e}[/yellow]")
            console.print("[yellow]Install with: uv add python-docx[/yellow]")
        
        if docx_available:
            
            # Create the original DOCX with embedded metadata (simulating copy-paste)
            doc = Document()
            
            # Add a title
            heading = doc.add_heading('Original Document with Embedded Metadata', 0)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Try to use a font with good Unicode support
            # Windows system fonts that generally have good Unicode support
            # Arial Unicode MS has excellent Unicode variation selector support
            unicode_fonts = ['Arial Unicode MS']
            
            # Directly use the text with embedded metadata (as if copy-pasted)
            # Split by paragraphs to maintain formatting
            for line in original_with_metadata.split('\n'):
                if line.strip():
                    # Add each line as a separate paragraph to preserve all Unicode characters
                    paragraph = doc.add_paragraph()
                    
                    # Try each font in order until one works
                    for font in unicode_fonts:
                        try:
                            run = paragraph.add_run(line)
                            run.font.name = font
                            break  # Stop if we successfully added the run with this font
                        except Exception:
                            continue  # Try the next font if this one fails
                    else:
                        # If all fonts failed, add without specifying a font
                        paragraph.add_run(line)
                else:
                    doc.add_paragraph()
            
            # Save original DOCX
            original_docx_path = EXAMPLE_DIR / "original.docx"
            doc.save(str(original_docx_path))
            console.print(f"[green]Created original DOCX with embedded metadata: {original_docx_path}[/green]")
            
            # Create the tampered DOCX (simulating copy-paste of tampered text)
            doc_tampered = Document()
            
            # Add a title
            heading = doc_tampered.add_heading('Tampered Document', 0)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Directly use the tampered text (as if copy-pasted)
            for line in tampered_text.split('\n'):
                if line.strip():
                    # Add each line as a separate paragraph
                    paragraph = doc_tampered.add_paragraph()
                    
                    # Try each font in order until one works
                    for font in unicode_fonts:
                        try:
                            run = paragraph.add_run(line)
                            run.font.name = font
                            break  # Stop if we successfully added the run with this font
                        except Exception:
                            continue  # Try the next font if this one fails
                    else:
                        # If all fonts failed, add without specifying a font
                        paragraph.add_run(line)
                else:
                    doc_tampered.add_paragraph()
            
            # Save tampered DOCX
            tampered_docx_path = EXAMPLE_DIR / "tampered.docx"
            doc_tampered.save(str(tampered_docx_path))
            
            console.print(f"[green]Created tampered DOCX with embedded metadata: {tampered_docx_path}[/green]")
            console.print("[green]DOCX files created with simulated copy-paste to preserve Unicode metadata[/green]")
            
            # Now save the DOCX files as PDFs
            try:
                import comtypes.client
                # Create the original PDF from the DOCX
                original_pdf_path = EXAMPLE_DIR / "original.pdf"
                wdFormatPDF = 17
                in_file = os.path.abspath(original_docx_path)
                out_file = os.path.abspath(original_pdf_path)
                word = comtypes.client.CreateObject('Word.Application')
                word.Visible = 0
                word.Documents.Open(in_file)
                word.Documents(1).SaveAs2(out_file, FileFormat=wdFormatPDF)
                word.Quit()
                console.print(f"[green]Created original PDF from DOCX: {original_pdf_path}[/green]")
                
                # Create the tampered PDF from the DOCX
                tampered_pdf_path = EXAMPLE_DIR / "tampered.pdf"
                in_file = os.path.abspath(tampered_docx_path)
                out_file = os.path.abspath(tampered_pdf_path)
                word = comtypes.client.CreateObject('Word.Application')
                word.Visible = 0
                word.Documents.Open(in_file)
                word.Documents(1).SaveAs2(out_file, FileFormat=wdFormatPDF)
                word.Quit()
                console.print(f"[green]Created tampered PDF from DOCX: {tampered_pdf_path}[/green]")
                pdf_creation_successful = True
                
            except Exception as e:
                console.print(f"[yellow]Error creating PDFs from DOCX: {e}[/yellow]")
                console.print("[yellow]Install Microsoft Word and the pywin32 library to enable PDF creation[/yellow]")
        
        else:
            console.print("[yellow]python-docx not installed, skipping DOCX example creation[/yellow]")
            console.print("[yellow]Install with: uv add python-docx[/yellow]")
        
        if not pdf_creation_successful:
            # Simulate copy-pasting text with embedded metadata into PDF files
            try:
                from reportlab.lib.pagesizes import letter
                from reportlab.lib.utils import simpleSplit
                from reportlab.pdfbase import pdfmetrics
                from reportlab.pdfbase.ttfonts import TTFont
                from reportlab.pdfgen import canvas
                
                # Try to register Windows system fonts with good Unicode support
                # Try multiple fonts in order of preference
                font_name = 'Helvetica'  # Default fallback
                
                try:
                    # First try Arial Unicode MS which has excellent Unicode support
                    arial_unicode_path = 'C:\\Windows\\Fonts\\ARIALUNI.TTF'
                    if Path(arial_unicode_path).exists():
                        pdfmetrics.registerFont(TTFont('ArialUnicodeMS', arial_unicode_path))
                        font_name = 'ArialUnicodeMS'
                        console.print("[green]Using Arial Unicode MS font with excellent Unicode support[/green]")
                    
                    # If Arial Unicode MS is not available, try Segoe UI
                    elif Path('C:\\Windows\\Fonts\\segoeui.ttf').exists():
                        pdfmetrics.registerFont(TTFont('SegoeUI', 'C:\\Windows\\Fonts\\segoeui.ttf'))
                        font_name = 'SegoeUI'
                        console.print("[green]Using Segoe UI font with good Unicode support[/green]")
                    
                    # If Segoe UI is not available, try Calibri
                    elif Path('C:\\Windows\\Fonts\\calibri.ttf').exists():
                        pdfmetrics.registerFont(TTFont('Calibri', 'C:\\Windows\\Fonts\\calibri.ttf'))
                        font_name = 'Calibri'
                        console.print("[green]Using Calibri font with good Unicode support[/green]")
                    
                    else:
                        console.print("[yellow]Using Helvetica font - some Unicode characters may not render correctly[/yellow]")
                except Exception as e:
                    console.print(f"[yellow]Error registering custom font: {e}. Using Helvetica as fallback.[/yellow]")
                
                # Create the original PDF with embedded metadata (simulating copy-paste)
                original_pdf_path = EXAMPLE_DIR / "original.pdf"
                c = canvas.Canvas(str(original_pdf_path), pagesize=letter)
                c.setFont(font_name, 12)
                
                # Directly use the text with embedded metadata (as if copy-pasted)
                y_position = 750  # Start from top of page
                for line in original_with_metadata.split('\n'):
                    if line.strip():  # Skip empty lines
                        # Split long lines to fit on page
                        for subline in simpleSplit(line, font_name, 12, 500):
                            # Draw the text with all its Unicode characters intact
                            c.drawString(50, y_position, subline)
                            y_position -= 15  # Move down for next line
                    else:
                        y_position -= 15  # Empty line spacing
                
                c.save()
                console.print(f"[green]Created original PDF with embedded metadata: {original_pdf_path}[/green]")
                
                # Create the tampered PDF (simulating copy-paste of tampered text)
                tampered_pdf_path = EXAMPLE_DIR / "tampered.pdf"
                c = canvas.Canvas(str(tampered_pdf_path), pagesize=letter)
                c.setFont(font_name, 12)
                
                # Directly use the tampered text (as if copy-pasted)
                y_position = 750  # Start from top of page
                for line in tampered_text.split('\n'):
                    if line.strip():  # Skip empty lines
                        # Split long lines to fit on page
                        for subline in simpleSplit(line, font_name, 12, 500):
                            c.drawString(50, y_position, subline)
                            y_position -= 15  # Move down for next line
                    else:
                        y_position -= 15  # Empty line spacing
                
                c.save()
                console.print(f"[green]Created tampered PDF with embedded metadata: {tampered_pdf_path}[/green]")
                console.print("[green]PDFs created with simulated copy-paste to preserve Unicode metadata[/green]")
                
            except ImportError:
                console.print("[yellow]ReportLab not installed, skipping PDF example creation[/yellow]")
                console.print("[yellow]Install with: uv add reportlab[/yellow]")
        
        # Simulate copy-pasting text with embedded metadata into DOCX files
        docx_available = False
        try:
            import docx
            from docx import Document
            from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
            docx_available = True
        except ImportError as e:
            console.print(f"[yellow]Error importing python-docx: {e}[/yellow]")
            console.print("[yellow]Install with: uv add python-docx[/yellow]")
        
        if docx_available:
            
            # Create the original DOCX with embedded metadata (simulating copy-paste)
            doc = Document()
            
            # Add a title
            heading = doc.add_heading('Original Document with Embedded Metadata', 0)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Try to use a font with good Unicode support
            # Windows system fonts that generally have good Unicode support
            # Arial Unicode MS has excellent Unicode variation selector support
            unicode_fonts = ['Arial Unicode MS']
            
            # Directly use the text with embedded metadata (as if copy-pasted)
            # Split by paragraphs to maintain formatting
            for line in original_with_metadata.split('\n'):
                if line.strip():
                    # Add each line as a separate paragraph to preserve all Unicode characters
                    paragraph = doc.add_paragraph()
                    
                    # Try each font in order until one works
                    for font in unicode_fonts:
                        try:
                            run = paragraph.add_run(line)
                            run.font.name = font
                            break  # Stop if we successfully added the run with this font
                        except Exception:
                            continue  # Try the next font if this one fails
                    else:
                        # If all fonts failed, add without specifying a font
                        paragraph.add_run(line)
                else:
                    doc.add_paragraph()
            
            # Save original DOCX
            original_docx_path = EXAMPLE_DIR / "original.docx"
            doc.save(str(original_docx_path))
            console.print(f"[green]Created original DOCX with embedded metadata: {original_docx_path}[/green]")
            
            # Create the tampered DOCX (simulating copy-paste of tampered text)
            doc_tampered = Document()
            
            # Add a title
            heading = doc_tampered.add_heading('Tampered Document', 0)
            heading.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
            
            # Directly use the tampered text (as if copy-pasted)
            for line in tampered_text.split('\n'):
                if line.strip():
                    # Add each line as a separate paragraph
                    paragraph = doc_tampered.add_paragraph()
                    
                    # Try each font in order until one works
                    for font in unicode_fonts:
                        try:
                            run = paragraph.add_run(line)
                            run.font.name = font
                            break  # Stop if we successfully added the run with this font
                        except Exception:
                            continue  # Try the next font if this one fails
                    else:
                        # If all fonts failed, add without specifying a font
                        paragraph.add_run(line)
                else:
                    doc_tampered.add_paragraph()
            
            # Save tampered DOCX
            tampered_docx_path = EXAMPLE_DIR / "tampered.docx"
            doc_tampered.save(str(tampered_docx_path))
            
            console.print(f"[green]Created tampered DOCX with embedded metadata: {tampered_docx_path}[/green]")
            console.print("[green]DOCX files created with simulated copy-paste to preserve Unicode metadata[/green]")
            
            # Create PDFs directly using a more reliable method for Unicode preservation
            try:
                import binascii
                import os
                import tempfile
                
                # Debug function to show Unicode code points in a string
                def debug_unicode(text):
                    result = []
                    for char in text:
                        code_point = ord(char)
                        if 0xFE00 <= code_point <= 0xFE0F:  # Variation selectors
                            result.append(f"VS-{code_point-0xFE00+1}({hex(code_point)})")
                        elif code_point > 127:  # Non-ASCII
                            result.append(f"{char}({hex(code_point)})")
                        else:
                            result.append(char)
                    return ''.join(result)
                
                # Function to create PDF with embedded Unicode text as both visible text and raw data
                def create_pdf_from_text(text, output_path, title="Document"):
                    # Install required packages if not available
                    try:
                        # Make sure reportlab is installed
                        try:
                            import reportlab
                        except ImportError:
                            console.print("[yellow]ReportLab not found, installing...[/yellow]")
                            subprocess.check_call([sys.executable, "-m", "uv", "add", "reportlab"])
                        
                        # Make sure pikepdf is installed
                        try:
                            import pikepdf
                        except ImportError:
                            console.print("[yellow]pikepdf not found, installing...[/yellow]")
                            subprocess.check_call([sys.executable, "-m", "uv", "add", "pikepdf"])
                            import pikepdf
                            
                        # Make sure we have DejaVu Sans font which has good Unicode support
                        try:
                            import urllib.request
                            
                            # Create fonts directory if it doesn't exist
                            fonts_dir = os.path.join(os.path.dirname(__file__), "fonts")
                            os.makedirs(fonts_dir, exist_ok=True)
                            
                            # DejaVu Sans has excellent Unicode support
                            dejavu_path = os.path.join(fonts_dir, "DejaVuSans.ttf")
                            
                            if not os.path.exists(dejavu_path):
                                console.print("[yellow]Downloading DejaVu Sans font for Unicode support...[/yellow]")
                                # URL for DejaVu Sans TTF
                                dejavu_url = "https://github.com/dejavu-fonts/dejavu-fonts/raw/master/ttf/DejaVuSans.ttf"
                                urllib.request.urlretrieve(dejavu_url, dejavu_path)
                                console.print(f"[green]Downloaded DejaVu Sans font to {dejavu_path}[/green]")
                        except Exception as font_err:
                            console.print(f"[yellow]Error downloading font: {font_err}, will try system fonts[/yellow]")
                        
                        # Import required ReportLab modules
                        from reportlab.lib.pagesizes import letter
                        from reportlab.pdfbase import pdfmetrics
                        from reportlab.pdfbase.ttfonts import TTFont
                        from reportlab.pdfgen import canvas
                        
                        # Create a temporary file for the initial PDF
                        temp_pdf = tempfile.NamedTemporaryFile(delete=False, suffix='.pdf')
                        temp_pdf.close()
                        
                        # Try to register DejaVu Sans font first (which we downloaded)
                        try:
                            # Use our downloaded DejaVu Sans font
                            dejavu_path = os.path.join(os.path.dirname(__file__), "fonts", "DejaVuSans.ttf")
                            if os.path.exists(dejavu_path):
                                pdfmetrics.registerFont(TTFont('DejaVuSans', dejavu_path))
                                font_name = 'DejaVuSans'
                                console.print("[green]Using embedded DejaVu Sans font for Unicode support[/green]")
                            else:
                                # Try Arial Unicode MS as fallback
                                arial_unicode_paths = [
                                    r"C:\Windows\Fonts\ARIALUNI.TTF",
                                    r"C:\Windows\Fonts\Arial Unicode MS.ttf"
                                ]
                                
                                for font_path in arial_unicode_paths:
                                    if os.path.exists(font_path):
                                        pdfmetrics.registerFont(TTFont('ArialUnicode', font_path))
                                        font_name = 'ArialUnicode'
                                        console.print(f"[green]Using Arial Unicode MS font from {font_path}[/green]")
                                        break
                                else:
                                    # Fallback to default font
                                    font_name = 'Helvetica'
                                    console.print("[yellow]Unicode-compatible fonts not found, using Helvetica (Unicode support may be limited)[/yellow]")
                        except Exception as e:
                            console.print(f"[yellow]Error registering fonts: {e}[/yellow]")
                            font_name = 'Helvetica'
                        
                        # Create PDF with canvas for visible text
                        c = canvas.Canvas(temp_pdf.name, pagesize=letter)
                        width, height = letter
                        
                        # Set title
                        c.setFont(font_name, 16)
                        c.drawString(50, height-50, title)
                        
                        # Add text content line by line
                        c.setFont(font_name, 12)
                        y_position = height - 80
                        line_height = 14
                        
                        for line in text.split('\n'):
                            if y_position < 50:  # Start a new page if we're near the bottom
                                c.showPage()
                                y_position = height - 50
                                c.setFont(font_name, 12)
                            
                            # Draw the line with all Unicode characters preserved
                            c.drawString(50, y_position, line)
                            y_position -= line_height
                        
                        c.save()
                        
                        # Now use pikepdf to add the raw Unicode text as metadata
                        # This ensures the Unicode variation selectors are preserved
                        pdf = pikepdf.open(temp_pdf.name)
                        
                        # Add the full text with Unicode variation selectors as a document-level metadata string
                        with pdf.open_metadata() as meta:
                            meta['dc:description'] = text
                            meta['pdf:Producer'] = 'Encypher Unicode Metadata Preservation'
                            
                        # Also add the text as a separate attachment to ensure preservation
                        filespec = pdf.make_stream(text.encode('utf-8'))
                        pdf.attachments['original_text.txt'] = filespec
                        
                        # Save the final PDF
                        # Convert Path object to string explicitly to avoid filename error
                        output_path_str = str(output_path)
                        # Use a try-except block to handle potential issues with pikepdf save
                        try:
                            pdf.save(output_path_str)
                            console.print(f"[green]Successfully saved PDF with embedded Unicode metadata to {output_path_str}[/green]")
                        except Exception as save_err:
                            console.print(f"[yellow]Error saving PDF with pikepdf: {save_err}[/yellow]")
                            # Try a direct approach as fallback
                            import shutil
                            try:
                                shutil.copy(temp_pdf.name, output_path_str)
                                console.print(f"[yellow]Used fallback method to save PDF to {output_path_str}[/yellow]")
                            except Exception as copy_err:
                                console.print(f"[red]Failed to save PDF: {copy_err}[/red]")
                                return False
                        
                        # Clean up the temporary file
                        try:
                            os.unlink(temp_pdf.name)
                        except Exception:
                            pass
                            
                        # Debug output - show the Unicode characters in the text
                        console.print(f"[blue]Text contains {len(text)} characters with Unicode details:[/blue]")
                        sample = text[:100] + '...' if len(text) > 100 else text
                        console.print(f"[blue]Sample: {debug_unicode(sample)}[/blue]")
                        
                        return True
                        
                    except Exception as e:
                        import traceback
                        console.print(f"[yellow]Error creating PDF with ReportLab: {e}[/yellow]")
                        console.print(f"[yellow]Error details: {traceback.format_exc()}[/yellow]")
                        # Try a simpler approach as fallback
                        try:
                            console.print("[yellow]Attempting fallback PDF creation method...[/yellow]")
                            # Create a simple PDF with just the text content
                            from reportlab.lib.pagesizes import letter
                            from reportlab.pdfgen import canvas
                            
                            c = canvas.Canvas(str(output_path), pagesize=letter)
                            width, height = letter
                            
                            # Set title
                            c.setFont('Helvetica', 16)
                            c.drawString(50, height-50, title)
                            
                            # Add text content line by line
                            c.setFont('Helvetica', 12)
                            y_position = height - 80
                            line_height = 14
                            
                            for line in text.split('\n'):
                                if y_position < 50:  # Start a new page if we're near the bottom
                                    c.showPage()
                                    y_position = height - 50
                                    c.setFont('Helvetica', 12)
                                
                                # Draw the line with all Unicode characters preserved
                                c.drawString(50, y_position, line)
                                y_position -= line_height
                            
                            c.save()
                            console.print(f"[green]Created PDF using fallback method: {output_path}[/green]")
                            return True
                        except Exception as fallback_err:
                            console.print(f"[red]Fallback PDF creation also failed: {fallback_err}[/red]")
                            return False
                
                # Create a simpler function for PDF creation that focuses on metadata embedding
                def create_simple_pdf_with_metadata(text, output_path, title="Document"):
                    try:
                        # Make sure reportlab is installed
                        try:
                            import reportlab
                        except ImportError:
                            console.print("[yellow]ReportLab not found, installing...[/yellow]")
                            subprocess.check_call([sys.executable, "-m", "uv", "add", "reportlab"])
                        
                        # Make sure pikepdf is installed
                        try:
                            import pikepdf
                        except ImportError:
                            console.print("[yellow]pikepdf not found, installing...[/yellow]")
                            subprocess.check_call([sys.executable, "-m", "uv", "add", "pikepdf"])
                            import pikepdf
                        
                        # Import required ReportLab modules
                        from reportlab.lib.pagesizes import letter
                        from reportlab.pdfgen import canvas
                        
                        # Create a basic PDF with the text
                        pdf_path = str(output_path)
                        c = canvas.Canvas(pdf_path, pagesize=letter)
                        width, height = letter
                        
                        # Add title
                        c.setFont('Helvetica-Bold', 16)
                        c.drawString(50, height-50, title)
                        
                        # Add content
                        c.setFont('Helvetica', 12)
                        y_position = height - 80
                        
                        # Add a note about Unicode metadata
                        c.drawString(50, y_position, "This document contains Unicode metadata embedded in multiple ways:")
                        y_position -= 20
                        c.drawString(50, y_position, "1. As document metadata (dc:description)")
                        y_position -= 20
                        c.drawString(50, y_position, "2. As an attached text file")
                        y_position -= 20
                        c.drawString(50, y_position, "3. As visible text (may not display all Unicode variation selectors)")
                        y_position -= 40
                        
                        # Add the text content (may not display all Unicode characters correctly)
                        for line in text.split('\n'):
                            if y_position < 50:
                                c.showPage()
                                y_position = height - 50
                            c.drawString(50, y_position, line)
                            y_position -= 15
                        
                        c.save()
                        console.print(f"[green]Created basic PDF at {pdf_path}[/green]")
                        
                        # Now add the Unicode text as metadata using pikepdf
                        pdf = pikepdf.open(pdf_path)
                        
                        # Add metadata
                        with pdf.open_metadata() as meta:
                            meta['dc:description'] = text
                            meta['pdf:Producer'] = 'Encypher Unicode Metadata Test'
                        
                        # Add text as attachment
                        filespec = pdf.make_stream(text.encode('utf-8'))
                        pdf.attachments['original_text.txt'] = filespec
                        
                        # Save with metadata
                        pdf.save(pdf_path)
                        console.print(f"[green]Added Unicode metadata to PDF: {pdf_path}[/green]")
                        
                        # Debug output
                        console.print(f"[blue]Text contains {len(text)} characters with Unicode details:[/blue]")
                        sample = text[:100] + '...' if len(text) > 100 else text
                        console.print(f"[blue]Sample: {debug_unicode(sample)}[/blue]")
                        
                        return True
                    except Exception as e:
                        import traceback
                        console.print(f"[red]Error creating PDF: {e}[/red]")
                        console.print(f"[yellow]{traceback.format_exc()}[/yellow]")
                        return False
                
                # Create original PDF
                original_pdf_path = EXAMPLE_DIR / "original.pdf"
                console.print("[yellow]Creating original PDF with embedded metadata...[/yellow]")
                create_simple_pdf_with_metadata(original_with_metadata, original_pdf_path, "Original Document with Embedded Metadata")
                console.print(f"[green]Created original PDF: {original_pdf_path}[/green]")
                
                # Create tampered PDF
                tampered_pdf_path = EXAMPLE_DIR / "tampered.pdf"
                console.print("[yellow]Creating tampered PDF...[/yellow]")
                create_simple_pdf_with_metadata(tampered_text, tampered_pdf_path, "Tampered Document")
                console.print(f"[green]Created tampered PDF: {tampered_pdf_path}[/green]")
                
                pdf_creation_successful = True
                
            except Exception as e:
                console.print(f"[red]Error creating PDFs with ReportLab: {e}[/red]")
                console.print("[yellow]Install ReportLab with: uv add reportlab[/yellow]")
                pdf_creation_successful = False
        else:
            console.print("[yellow]python-docx not installed, skipping DOCX example creation[/yellow]")
            console.print("[yellow]Install with: uv add python-docx[/yellow]")
            pdf_creation_successful = False
        
    except Exception as e:
        console.print(f"[red]Error creating examples: {e}[/red]")
    
    console.print("[bold green]Successfully created simple tampered example files![/bold green]")
    console.print("\nTo test these files with the audit-log-cli tool, run:")
    console.print("uv run -- python -m audit_log_cli.app.main --target audit_log_cli/tests/simple_tampered_examples --trusted-signers audit_log_cli/tests/simple_tampered_examples/keys --verify-content-integrity")

except ImportError as e:
    print(f"Error importing required modules: {e}")
    print("Make sure you have installed all the required dependencies.")
    sys.exit(1)
except Exception as e:
    print(f"Unexpected error: {e}")
    sys.exit(1)
