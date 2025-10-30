"""
Binary file support for signing PDFs, DOCX, images, and other formats.

Provides text extraction and signing capabilities for binary files.
"""
from pathlib import Path
from typing import Optional, List, Dict, Any, BinaryIO
from dataclasses import dataclass
from datetime import datetime
import io


@dataclass
class BinaryFileInfo:
    """Information about a binary file."""
    file_path: Path
    file_type: str  # pdf, docx, image, video
    file_size: int
    mime_type: str
    extracted_text: Optional[str] = None
    metadata: Dict[str, Any] = None
    
    def to_dict(self) -> dict:
        """Convert to dictionary."""
        return {
            "file_path": str(self.file_path),
            "file_type": self.file_type,
            "file_size": self.file_size,
            "mime_type": self.mime_type,
            "extracted_text": self.extracted_text,
            "metadata": self.metadata or {}
        }


class TextExtractor:
    """
    Extract text from binary files.
    
    Example:
        >>> extractor = TextExtractor()
        >>> text = extractor.extract_from_pdf(Path("document.pdf"))
        >>> text = extractor.extract_from_docx(Path("document.docx"))
    """
    
    def extract_from_pdf(self, file_path: Path) -> str:
        """
        Extract text from PDF file.
        
        Args:
            file_path: Path to PDF file
        
        Returns:
            Extracted text
        """
        try:
            from PyPDF2 import PdfReader
            
            reader = PdfReader(str(file_path))
            text_parts = []
            
            for page in reader.pages:
                text = page.extract_text()
                if text:
                    text_parts.append(text)
            
            return "\n\n".join(text_parts)
        
        except ImportError:
            raise ImportError("PyPDF2 not installed. Install with: uv add PyPDF2")
        except Exception as e:
            raise ValueError(f"Failed to extract text from PDF: {e}")
    
    def extract_from_docx(self, file_path: Path) -> str:
        """
        Extract text from DOCX file.
        
        Args:
            file_path: Path to DOCX file
        
        Returns:
            Extracted text
        """
        try:
            from docx import Document
            
            doc = Document(str(file_path))
            text_parts = []
            
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_parts.append(paragraph.text)
            
            return "\n\n".join(text_parts)
        
        except ImportError:
            raise ImportError("python-docx not installed. Install with: uv add python-docx")
        except Exception as e:
            raise ValueError(f"Failed to extract text from DOCX: {e}")
    
    def extract_from_image(self, file_path: Path) -> Optional[str]:
        """
        Extract text from image using OCR (if available).
        
        Args:
            file_path: Path to image file
        
        Returns:
            Extracted text or None if OCR not available
        """
        # OCR would require pytesseract
        # For now, return None (OCR integration optional)
        return None
    
    def extract_metadata_from_pdf(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract metadata from PDF.
        
        Args:
            file_path: Path to PDF file
        
        Returns:
            Dict with PDF metadata
        """
        try:
            from PyPDF2 import PdfReader
            
            reader = PdfReader(str(file_path))
            metadata = {}
            
            if reader.metadata:
                for key, value in reader.metadata.items():
                    # Remove leading slash from keys
                    clean_key = key.lstrip('/')
                    metadata[clean_key] = value
            
            metadata['num_pages'] = len(reader.pages)
            
            return metadata
        
        except Exception:
            return {}
    
    def extract_metadata_from_docx(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract metadata from DOCX.
        
        Args:
            file_path: Path to DOCX file
        
        Returns:
            Dict with DOCX metadata
        """
        try:
            from docx import Document
            
            doc = Document(str(file_path))
            metadata = {}
            
            if doc.core_properties:
                props = doc.core_properties
                if props.author:
                    metadata['author'] = props.author
                if props.title:
                    metadata['title'] = props.title
                if props.subject:
                    metadata['subject'] = props.subject
                if props.created:
                    metadata['created'] = props.created.isoformat()
                if props.modified:
                    metadata['modified'] = props.modified.isoformat()
            
            metadata['num_paragraphs'] = len(doc.paragraphs)
            
            return metadata
        
        except Exception:
            return {}
    
    def extract_metadata_from_image(self, file_path: Path) -> Dict[str, Any]:
        """
        Extract metadata from image.
        
        Args:
            file_path: Path to image file
        
        Returns:
            Dict with image metadata
        """
        try:
            from PIL import Image
            
            with Image.open(file_path) as img:
                metadata = {
                    'format': img.format,
                    'mode': img.mode,
                    'width': img.width,
                    'height': img.height,
                    'size_bytes': file_path.stat().st_size
                }
                
                # Extract EXIF data if available
                if hasattr(img, '_getexif') and img._getexif():
                    exif = img._getexif()
                    if exif:
                        metadata['exif'] = {k: str(v) for k, v in exif.items()}
                
                return metadata
        
        except ImportError:
            raise ImportError("Pillow not installed. Install with: uv add Pillow")
        except Exception:
            return {}


class BinaryFileSigner:
    """
    Sign binary files with C2PA metadata.
    
    Example:
        >>> from encypher_enterprise import EncypherClient
        >>> client = EncypherClient(api_key="...")
        >>> signer = BinaryFileSigner(client)
        >>> 
        >>> # Sign PDF
        >>> result = signer.sign_pdf(Path("document.pdf"))
        >>> 
        >>> # Sign DOCX
        >>> result = signer.sign_docx(Path("document.docx"))
    """
    
    def __init__(self, client):
        """
        Initialize binary file signer.
        
        Args:
            client: Encypher client instance
        """
        self.client = client
        self.extractor = TextExtractor()
    
    def detect_file_type(self, file_path: Path) -> str:
        """
        Detect binary file type.
        
        Args:
            file_path: Path to file
        
        Returns:
            File type (pdf, docx, image, video, unknown)
        """
        suffix = file_path.suffix.lower()
        
        if suffix == '.pdf':
            return 'pdf'
        elif suffix in ['.docx', '.doc']:
            return 'docx'
        elif suffix in ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.webp', '.svg']:
            return 'image'
        elif suffix in ['.mp4', '.avi', '.mov', '.wmv', '.flv', '.webm']:
            return 'video'
        else:
            return 'unknown'
    
    def get_mime_type(self, file_path: Path) -> str:
        """
        Get MIME type for file.
        
        Args:
            file_path: Path to file
        
        Returns:
            MIME type string
        """
        suffix = file_path.suffix.lower()
        
        mime_types = {
            '.pdf': 'application/pdf',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.doc': 'application/msword',
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.png': 'image/png',
            '.gif': 'image/gif',
            '.bmp': 'image/bmp',
            '.webp': 'image/webp',
            '.svg': 'image/svg+xml',
            '.mp4': 'video/mp4',
            '.avi': 'video/x-msvideo',
            '.mov': 'video/quicktime',
        }
        
        return mime_types.get(suffix, 'application/octet-stream')
    
    def analyze_file(self, file_path: Path) -> BinaryFileInfo:
        """
        Analyze binary file and extract information.
        
        Args:
            file_path: Path to file
        
        Returns:
            BinaryFileInfo with file details
        """
        file_type = self.detect_file_type(file_path)
        mime_type = self.get_mime_type(file_path)
        file_size = file_path.stat().st_size
        
        # Extract text if possible
        extracted_text = None
        metadata = {}
        
        try:
            if file_type == 'pdf':
                extracted_text = self.extractor.extract_from_pdf(file_path)
                metadata = self.extractor.extract_metadata_from_pdf(file_path)
            elif file_type == 'docx':
                extracted_text = self.extractor.extract_from_docx(file_path)
                metadata = self.extractor.extract_metadata_from_docx(file_path)
            elif file_type == 'image':
                metadata = self.extractor.extract_metadata_from_image(file_path)
        except Exception:
            # Ignore extraction errors
            pass
        
        return BinaryFileInfo(
            file_path=file_path,
            file_type=file_type,
            file_size=file_size,
            mime_type=mime_type,
            extracted_text=extracted_text,
            metadata=metadata
        )
    
    def sign_pdf(
        self,
        file_path: Path,
        extract_text: bool = True,
        **metadata
    ) -> Any:
        """
        Sign PDF file.
        
        Args:
            file_path: Path to PDF file
            extract_text: Extract and sign text content
            **metadata: Additional metadata
        
        Returns:
            Signing result
        """
        info = self.analyze_file(file_path)
        
        if extract_text and info.extracted_text:
            # Sign extracted text
            return self.client.sign(
                text=info.extracted_text,
                document_type="pdf",
                **metadata
            )
        else:
            # Sign file hash (API enhancement needed for binary signing)
            raise NotImplementedError("Binary PDF signing requires API enhancement")
    
    def sign_docx(
        self,
        file_path: Path,
        extract_text: bool = True,
        **metadata
    ) -> Any:
        """
        Sign DOCX file.
        
        Args:
            file_path: Path to DOCX file
            extract_text: Extract and sign text content
            **metadata: Additional metadata
        
        Returns:
            Signing result
        """
        info = self.analyze_file(file_path)
        
        if extract_text and info.extracted_text:
            # Sign extracted text
            return self.client.sign(
                text=info.extracted_text,
                document_type="docx",
                **metadata
            )
        else:
            # Sign file hash (API enhancement needed)
            raise NotImplementedError("Binary DOCX signing requires API enhancement")
    
    def sign_image(
        self,
        file_path: Path,
        **metadata
    ) -> Any:
        """
        Sign image file.
        
        Args:
            file_path: Path to image file
            **metadata: Additional metadata
        
        Returns:
            Signing result
        """
        # Image signing requires API enhancement for binary files
        raise NotImplementedError("Image signing requires API enhancement")
    
    def sign_video(
        self,
        file_path: Path,
        **metadata
    ) -> Any:
        """
        Sign video file.
        
        Args:
            file_path: Path to video file
            **metadata: Additional metadata
        
        Returns:
            Signing result
        """
        # Video signing requires API enhancement for binary files
        raise NotImplementedError("Video signing requires API enhancement")


def is_binary_file(file_path: Path) -> bool:
    """
    Check if file is binary.
    
    Args:
        file_path: Path to file
    
    Returns:
        True if binary, False if text
    """
    try:
        with open(file_path, 'rb') as f:
            chunk = f.read(8192)
            
        # Check for null bytes (indicator of binary)
        if b'\x00' in chunk:
            return True
        
        # Try to decode as UTF-8
        try:
            chunk.decode('utf-8')
            return False
        except UnicodeDecodeError:
            return True
    
    except Exception:
        return True


def get_supported_formats() -> Dict[str, List[str]]:
    """
    Get supported binary file formats.
    
    Returns:
        Dict of category -> list of extensions
    """
    return {
        "documents": [".pdf", ".docx", ".doc"],
        "images": [".jpg", ".jpeg", ".png", ".gif", ".bmp", ".webp", ".svg"],
        "video": [".mp4", ".avi", ".mov", ".wmv", ".flv", ".webm"],
        "audio": [".mp3", ".wav", ".flac", ".aac", ".ogg"]
    }


def is_supported_format(file_path: Path) -> bool:
    """
    Check if file format is supported.
    
    Args:
        file_path: Path to file
    
    Returns:
        True if supported, False otherwise
    """
    supported = get_supported_formats()
    suffix = file_path.suffix.lower()
    
    for formats in supported.values():
        if suffix in formats:
            return True
    
    return False
