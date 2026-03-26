"""
Generate a professional DOCX from the C2PA Security Architecture Document markdown.

Usage (from enterprise_api directory):
    uv run python scripts/generate_security_doc.py
"""

import os
import re
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

# ---------------------------------------------------------------------------
# Paths
# ---------------------------------------------------------------------------
REPO_ROOT = Path(__file__).parent.parent.parent
SOURCE_MD = REPO_ROOT / "docs/c2pa/conformance/SECURITY_ARCHITECTURE_DOCUMENT.md"
OUTPUT_DIR = REPO_ROOT / "docs/c2pa/conformance/submission"
OUTPUT_DOCX = OUTPUT_DIR / "Encypher_Security_Architecture_Document.docx"

# ---------------------------------------------------------------------------
# Colour constants
# ---------------------------------------------------------------------------
LIGHT_GRAY = RGBColor(0xF2, 0xF2, 0xF2)  # alternating row / code background
WHITE = RGBColor(0xFF, 0xFF, 0xFF)
HEADER_BLUE = RGBColor(0x1F, 0x49, 0x7D)  # dark navy for H1


def _set_cell_bg(cell, color_hex: str) -> None:
    """Set the background fill of a table cell."""
    tc = cell._tc
    tcp = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcp.append(shd)


def _add_page_number_footer(doc: Document) -> None:
    """Add centred page-number footer to the default section."""
    section = doc.sections[0]
    footer = section.footer
    para = footer.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = para.add_run()
    # Insert PAGE field
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    run._r.append(instr)

    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_end)


def _add_header_text(doc: Document, text: str) -> None:
    """Add text to the default section header."""
    section = doc.sections[0]
    header = section.header
    para = header.paragraphs[0]
    para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = para.add_run(text)
    run.font.name = "Calibri"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def _set_margins(doc: Document, inches: float = 1.0) -> None:
    section = doc.sections[0]
    section.top_margin = Inches(inches)
    section.bottom_margin = Inches(inches)
    section.left_margin = Inches(inches)
    section.right_margin = Inches(inches)


# ---------------------------------------------------------------------------
# Style helpers
# ---------------------------------------------------------------------------


def _style_heading(para, level: int) -> None:
    """Apply heading font styling after adding the paragraph."""
    sizes = {1: 16, 2: 14, 3: 12, 4: 11}
    size = sizes.get(level, 11)
    for run in para.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run.font.bold = True
        if level == 1:
            run.font.color.rgb = HEADER_BLUE


def _apply_inline_formatting(run_container, text: str) -> None:
    """
    Parse **bold** and *italic* markers and add formatted runs to run_container.
    run_container is a paragraph object that supports .add_run().
    """
    # Split on bold (**...**) and italic (*...*)
    # Pattern: **text** or *text*
    pattern = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*)")
    parts = pattern.split(text)
    for part in parts:
        if part.startswith("**") and part.endswith("**"):
            r = run_container.add_run(part[2:-2])
            r.bold = True
        elif part.startswith("*") and part.endswith("*"):
            r = run_container.add_run(part[1:-1])
            r.italic = True
        else:
            run_container.add_run(part)


def _set_body_font(para) -> None:
    for run in para.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)


# ---------------------------------------------------------------------------
# Title page
# ---------------------------------------------------------------------------


def _add_title_page(doc: Document) -> None:
    section = doc.sections[0]
    section.different_first_page_header_footer = True
    # First-page header: blank
    first_header = section.first_page_header
    first_header.paragraphs[0].text = ""

    def _title_para(text: str, size: int, bold: bool = False, space_before: int = 0):
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(space_before)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(text)
        r.font.name = "Calibri"
        r.font.size = Pt(size)
        r.font.bold = bold

    # Spacer
    for _ in range(6):
        sp = doc.add_paragraph()
        sp.paragraph_format.space_after = Pt(0)

    _title_para("Generator Product Security Architecture Document", 22, bold=True, space_before=0)
    _title_para("C2PA Conformance Program Submission", 16, bold=False)

    sep = doc.add_paragraph()
    sep.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep.paragraph_format.space_after = Pt(18)

    _title_para("Encypher Corporation", 14, bold=True)
    _title_para("Record ID: 019d2241-eb97-7728-9ec0-cdaafba300c2", 11)
    _title_para("Version 1.0 -- 2026-03-25", 11)

    sep2 = doc.add_paragraph()
    sep2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sep2.paragraph_format.space_after = Pt(18)

    conf = doc.add_paragraph()
    conf.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = conf.add_run("Confidential")
    r.font.name = "Calibri"
    r.font.size = Pt(12)
    r.font.bold = True
    r.font.color.rgb = RGBColor(0xC0, 0x00, 0x00)

    # Page break after title page
    doc.add_page_break()


# ---------------------------------------------------------------------------
# Table rendering
# ---------------------------------------------------------------------------


def _render_table(doc: Document, raw_lines: list) -> None:
    """
    Render a markdown pipe-table into a docx table with alternating row shading.
    raw_lines: list of raw markdown table lines (including separator row).
    """
    rows_data = []
    for line in raw_lines:
        if re.match(r"^\s*\|[-:| ]+\|\s*$", line):
            continue  # separator row
        cells = [c.strip() for c in line.strip().strip("|").split("|")]
        rows_data.append(cells)

    if not rows_data:
        return

    col_count = max(len(r) for r in rows_data)
    # Pad rows to uniform width
    rows_data = [r + [""] * (col_count - len(r)) for r in rows_data]

    table = doc.add_table(rows=len(rows_data), cols=col_count)
    table.style = "Table Grid"

    for row_idx, row_cells in enumerate(rows_data):
        is_header = row_idx == 0
        bg_hex = "F2F2F2" if (row_idx % 2 == 0 and not is_header) else "FFFFFF"
        if is_header:
            bg_hex = "1F497D"

        for col_idx, cell_text in enumerate(row_cells):
            cell = table.cell(row_idx, col_idx)
            _set_cell_bg(cell, bg_hex)
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(3)
            para.paragraph_format.space_after = Pt(3)
            _apply_inline_formatting(para, cell_text)
            for run in para.runs:
                run.font.name = "Calibri"
                run.font.size = Pt(10)
                if is_header:
                    run.font.bold = True
                    run.font.color.rgb = WHITE

    doc.add_paragraph()  # spacing after table


# ---------------------------------------------------------------------------
# Code block rendering
# ---------------------------------------------------------------------------


def _render_code_block(doc: Document, lines: list) -> None:
    """Render lines as a monospace code block with light-gray background."""
    for line in lines:
        p = doc.add_paragraph(style="No Spacing")
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.left_indent = Inches(0.25)
        r = p.add_run(line)
        r.font.name = "Consolas"
        r.font.size = Pt(9)
        # Set gray background on the paragraph shading
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"), "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"), "F2F2F2")
        pPr.append(shd)
    doc.add_paragraph()  # spacing after block


# ---------------------------------------------------------------------------
# Body paragraph rendering
# ---------------------------------------------------------------------------


def _render_body_para(doc: Document, text: str, indent: int = 0) -> None:
    """Render a paragraph with optional indent (in 0.25in steps)."""
    if not text.strip():
        return
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    if indent > 0:
        p.paragraph_format.left_indent = Inches(0.25 * indent)
    _apply_inline_formatting(p, text)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)


def _render_list_item(doc: Document, text: str, ordered: bool, number: int, depth: int) -> None:
    """Render a bullet or numbered list item."""
    p = doc.add_paragraph(style="No Spacing")
    p.paragraph_format.space_before = Pt(1)
    p.paragraph_format.space_after = Pt(2)
    indent_in = 0.25 + depth * 0.25
    p.paragraph_format.left_indent = Inches(indent_in + 0.18)
    p.paragraph_format.first_line_indent = Inches(-0.18)

    if ordered:
        prefix = f"{number}. "
    else:
        prefix = "- "

    # First run: bullet/number prefix (plain)
    pr = p.add_run(prefix)
    pr.font.name = "Calibri"
    pr.font.size = Pt(11)

    # Remaining text with inline formatting
    _apply_inline_formatting(p, text)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(11)


# ---------------------------------------------------------------------------
# Heading rendering
# ---------------------------------------------------------------------------


def _render_heading(doc: Document, text: str, level: int) -> None:
    style_map = {1: "Heading 1", 2: "Heading 2", 3: "Heading 3", 4: "Heading 4"}
    style = style_map.get(level, "Heading 4")
    try:
        p = doc.add_heading(text, level=level)
    except Exception:
        p = doc.add_paragraph(style=style)
        p.add_run(text)

    # Apply font sizes directly on runs to override theme defaults
    sizes = {1: 16, 2: 14, 3: 12, 4: 11}
    size = sizes.get(level, 11)
    for run in p.runs:
        run.font.name = "Calibri"
        run.font.size = Pt(size)
        run.font.bold = True
        if level == 1:
            run.font.color.rgb = HEADER_BLUE
        elif level == 2:
            run.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
        elif level == 3:
            run.font.color.rgb = RGBColor(0x44, 0x72, 0xC4)
        else:
            run.font.color.rgb = RGBColor(0x26, 0x26, 0x26)


# ---------------------------------------------------------------------------
# Horizontal rule
# ---------------------------------------------------------------------------


def _render_hr(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pb = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "AAAAAA")
    pb.append(bottom)
    pPr.append(pb)


# ---------------------------------------------------------------------------
# Markdown parser / document builder
# ---------------------------------------------------------------------------


def _parse_and_build(doc: Document, md_text: str) -> None:
    """
    Walk through markdown line by line and emit DOCX content.
    Handles:
      - # / ## / ### / #### headings
      - --- horizontal rules
      - ``` code blocks
      - | pipe tables
      - - / * bullet lists (with depth via leading spaces)
      - 1. / a. / etc. numbered lists
      - blank lines
      - regular paragraphs (with **bold** and *italic*)
    """
    lines = md_text.splitlines()
    i = 0
    # Track ordered list counters per depth
    ordered_counters: dict = {}

    while i < len(lines):
        line = lines[i]
        raw = line  # preserve for display

        # --- Heading (# H1, ## H2, etc.)
        hm = re.match(r"^(#{1,4})\s+(.*)", line)
        if hm:
            level = len(hm.group(1))
            text = hm.group(2).strip()
            _render_heading(doc, text, level)
            # Reset ordered counters below this level
            ordered_counters = {k: v for k, v in ordered_counters.items() if k < level}
            i += 1
            continue

        # --- Horizontal rule
        if re.match(r"^---+\s*$", line):
            _render_hr(doc)
            i += 1
            continue

        # --- Code block (```)
        if line.strip().startswith("```"):
            code_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith("```"):
                code_lines.append(lines[i])
                i += 1
            i += 1  # consume closing ```
            _render_code_block(doc, code_lines)
            continue

        # --- Table (starts with |)
        if re.match(r"^\s*\|", line):
            table_lines = []
            while i < len(lines) and re.match(r"^\s*\|", lines[i]):
                table_lines.append(lines[i])
                i += 1
            _render_table(doc, table_lines)
            continue

        # --- Bullet list item (-, *, or indented -)
        bm = re.match(r"^(\s*)([-*])\s+(.*)", line)
        if bm:
            leading = bm.group(1)
            text = bm.group(3)
            depth = len(leading) // 2
            # Accumulate multi-line bullet
            i += 1
            while i < len(lines):
                next_line = lines[i]
                # Continuation: non-empty and no special prefix
                if (
                    next_line
                    and not re.match(r"^\s*[-*]\s", next_line)
                    and not re.match(r"^\s*\d+[.)]\s", next_line)
                    and not re.match(r"^#{1,4}\s", next_line)
                    and not re.match(r"^\s*\|", next_line)
                    and not next_line.strip().startswith("```")
                    and not re.match(r"^---+\s*$", next_line)
                    and next_line.strip() != ""
                ):
                    text += " " + next_line.strip()
                    i += 1
                else:
                    break
            _render_list_item(doc, text, ordered=False, number=0, depth=depth)
            continue

        # --- Ordered list item (1. / a. / etc.)
        om = re.match(r"^(\s*)(\d+|[a-z])[.)]\s+(.*)", line)
        if om:
            leading = om.group(1)
            marker = om.group(2)
            text = om.group(3)
            depth = len(leading) // 3

            if marker.isdigit():
                num = int(marker)
            else:
                num = ord(marker) - ord("a") + 1

            ordered_counters[depth] = num
            # Accumulate multi-line list item
            i += 1
            while i < len(lines):
                next_line = lines[i]
                if (
                    next_line
                    and not re.match(r"^\s*[-*]\s", next_line)
                    and not re.match(r"^\s*\d+[.)]\s", next_line)
                    and not re.match(r"^#{1,4}\s", next_line)
                    and not re.match(r"^\s*\|", next_line)
                    and not next_line.strip().startswith("```")
                    and not re.match(r"^---+\s*$", next_line)
                    and next_line.strip() != ""
                ):
                    # Only if indented more than the list item start
                    text += " " + next_line.strip()
                    i += 1
                else:
                    break
            _render_list_item(doc, text, ordered=True, number=num, depth=depth)
            continue

        # --- Blank line
        if not line.strip():
            i += 1
            continue

        # --- Regular paragraph / italic document footer
        # Collect continuation lines for multi-line paragraph
        para_text = line
        i += 1
        while i < len(lines):
            next_line = lines[i]
            if (
                next_line.strip() == ""
                or re.match(r"^#{1,4}\s", next_line)
                or re.match(r"^---+\s*$", next_line)
                or re.match(r"^\s*\|", next_line)
                or next_line.strip().startswith("```")
                or re.match(r"^\s*[-*]\s", next_line)
                or re.match(r"^\s*\d+[.)]\s", next_line)
            ):
                break
            para_text += " " + next_line.strip()
            i += 1

        _render_body_para(doc, para_text)


# ---------------------------------------------------------------------------
# Main
# ---------------------------------------------------------------------------


def main() -> None:
    OUTPUT_DIR.mkdir(parents=True, exist_ok=True)

    md_text = SOURCE_MD.read_text(encoding="utf-8")

    doc = Document()
    _set_margins(doc)
    _add_header_text(doc, "Encypher Corporation -- Confidential")
    _add_page_number_footer(doc)

    # Title page (must be added before body content)
    _add_title_page(doc)

    # Parse and build body
    _parse_and_build(doc, md_text)

    doc.save(str(OUTPUT_DOCX))
    size_kb = OUTPUT_DOCX.stat().st_size / 1024
    print(f"Written: {OUTPUT_DOCX}")
    print(f"Size: {size_kb:.1f} KB")


if __name__ == "__main__":
    main()
