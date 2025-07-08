#!/usr/bin/env python
"""
Test script for Word to PDF conversion with metadata embedding.
"""
import os
import sys
import tempfile
from cryptography.hazmat.primitives.asymmetric.ed25519 import Ed25519PrivateKey

# Use relative imports since we're in the project directory
from encypher.unicode_metadata import UnicodeMetadata
from encypher.pdf_generator import EncypherPDF

def generate_key_pair():
    """Generate a test key pair for metadata signing."""
    private_key = Ed25519PrivateKey.generate()
    return private_key

def test_docx_to_pdf_conversion():
    """Test converting a Word document to PDF with embedded metadata."""
    # Check if a docx file path was provided
    if len(sys.argv) < 2:
        print("Usage: python test_docx_to_pdf.py <path_to_docx_file>")
        print("Creating a sample docx file for testing...")
        
        # Create a simple docx file for testing
        from docx import Document
        doc = Document()
        doc.add_heading('Test Document with Unicode Characters', 0)
        doc.add_paragraph('This is a test document with some Unicode characters: ✓ ♥ ☺ ☻ ♦ ♣ ♠')
        doc.add_paragraph('Testing metadata embedding in PDF conversion.')
        doc.add_paragraph('The quick brown fox jumps over the lazy dog.')
        
        # Save the test document
        test_docx = os.path.join(tempfile.gettempdir(), 'test_document.docx')
        doc.save(test_docx)
        print(f"Created test document at: {test_docx}")
    else:
        test_docx = sys.argv[1]
        if not os.path.exists(test_docx):
            print(f"Error: File not found: {test_docx}")
            return
    
    # Generate output PDF path
    output_pdf = os.path.splitext(test_docx)[0] + '_output.pdf'
    
    # Generate a key pair for signing
    private_key = generate_key_pair()
    signer_id = "test-signer"
    key_id = "test-key-1"
    
    print(f"Converting {test_docx} to PDF with embedded metadata...")
    
    try:
        # Convert docx to PDF with metadata
        pdf_path = EncypherPDF.from_docx(
            docx_file=test_docx,
            output_file=output_pdf,
            private_key=private_key,
            signer_id=signer_id,
            key_id=key_id
        )
        
        print(f"Successfully generated PDF at: {pdf_path}")
        
        # Extract and verify the text from the PDF
        print("Extracting text from the generated PDF...")
        extracted_text = EncypherPDF.extract_text(pdf_path)
        
        # Check if metadata is present in the extracted text
        metadata = UnicodeMetadata.extract_metadata(extracted_text)
        if metadata:
            print("✅ Metadata successfully embedded and extracted from PDF!")
            print(f"Metadata: {metadata}")
        else:
            print("❌ No metadata found in the extracted text.")
            print("Checking if metadata was embedded as an attachment...")
            
            # Check if the PDF has an embedded metadata file
            import fitz
            doc = fitz.open(pdf_path)
            names = doc.embfile_names()
            
            if "embedded_metadata.txt" in names:
                buffer = doc.embfile_get("embedded_metadata.txt")
                if buffer:
                    embedded_text = buffer["content"].decode("utf-8")
                    metadata = UnicodeMetadata.extract_metadata(embedded_text)
                    if metadata:
                        print("✅ Metadata found in embedded attachment!")
                        print(f"Metadata: {metadata}")
                    else:
                        print("❌ No valid metadata found in embedded attachment.")
            else:
                print("❌ No embedded metadata attachment found in PDF.")
        
        print("\nTest completed.")
        
    except Exception as e:
        print(f"Error during conversion: {str(e)}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    test_docx_to_pdf_conversion()
