"""
Script to create example files with Encypher metadata for the dashboard app.
Includes text files, PDFs, and other document types with various metadata attributes.
"""

import base64
import json
import random
import sys
from datetime import datetime, timedelta, timezone
from pathlib import Path
from typing import Any, Dict, Optional

# Add the parent directory to the path so we can import the shared library
sys.path.append(str(Path(__file__).parent.parent.parent))

try:
    # Import from the core Encypher package as used in shared_commercial_libs
    # Import from our shared commercial library
    from reportlab.lib import colors

    # For document generation
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
    from reportlab.platypus import Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle

    from encypher.core.keys import load_private_key_from_data, load_public_key_from_data
    from encypher.core.unicode_metadata import MetadataTarget, UnicodeMetadata
    from encypher_commercial_shared import Encypher
    
    # For Word document generation
    try:
        import docx
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        from docx.shared import Inches, Pt
        DOCX_AVAILABLE = True
    except ImportError:
        DOCX_AVAILABLE = False
    
    from rich.console import Console
    from rich.progress import BarColumn, Progress, TaskProgressColumn, TextColumn
    
    console = Console()
    
    # Directory to save example files
    EXAMPLE_DIR = Path(__file__).parent / "files"
    EXAMPLE_DIR.mkdir(exist_ok=True)
    
    # Create example key files for testing
    KEY_DIR = EXAMPLE_DIR / "keys"
    KEY_DIR.mkdir(exist_ok=True)
    
    # Generate a test private/public key pair for signing
    from cryptography.hazmat.primitives import serialization
    from cryptography.hazmat.primitives.asymmetric import ed25519
    
    def generate_key_pair(private_key_path, public_key_path, force=False):
        """Generate a test Ed25519 key pair for signing and verification."""
        if not force and private_key_path.exists() and public_key_path.exists():
            console.print(f"[yellow]Keys already exist at {KEY_DIR}, skipping generation[/yellow]")
            return None, None
            
        # Generate a private key
        private_key = ed25519.Ed25519PrivateKey.generate()
        
        # Serialize private key to PEM format
        private_pem = private_key.private_bytes(
            encoding=serialization.Encoding.PEM,
            format=serialization.PrivateFormat.PKCS8,
            encryption_algorithm=serialization.NoEncryption()
        )
        
        # Serialize public key to PEM format
        public_key = private_key.public_key()
        public_pem = public_key.public_bytes(
            encoding=serialization.Encoding.PEM,
            format=serialization.PublicFormat.SubjectPublicKeyInfo
        )
        
        # Write keys to files
        with open(private_key_path, 'wb') as f:
            f.write(private_pem)
        
        with open(public_key_path, 'wb') as f:
            f.write(public_pem)
        
        return private_key, public_key
    
    # Generate keys for different signers
    signers = [
        {"id": "encypherai-official", "name": "Encypher Official"},
        {"id": "compliance-dept", "name": "Compliance Department"},
        {"id": "research-team", "name": "Research Team"},
        {"id": "untrusted-source", "name": "Untrusted Source"}
    ]
    
    # Generate keys for each signer
    for signer in signers:
        signer_id = signer["id"]
        private_key_path = KEY_DIR / f"{signer_id}-private.pem"
        public_key_path = KEY_DIR / f"{signer_id}-public.pem"
        generate_key_pair(private_key_path, public_key_path)
        console.print(f"[green]Generated {signer['name']} keys at {KEY_DIR}[/green]")
    
    # Create a list of Encypher instances for each signer
    encypherai_instances = {}
    for signer in signers:
        signer_id = signer["id"]
        private_key_path = KEY_DIR / f"{signer_id}-private.pem"
        public_key_path = KEY_DIR / f"{signer_id}-public.pem"
        
        # Create Encypher instance with timestamp support
        ea = Encypher(
            private_key_path=str(private_key_path),
            public_key_path=str(public_key_path),
            signer_id=signer_id,
            verbose=True
        )
        
        # Add timestamp support to embed_metadata method
        from types import MethodType
        
        def embed_metadata_with_timestamp(self, text, custom_metadata=None, model_id=None, 
                                         metadata_format="basic", target=MetadataTarget.WHITESPACE,
                                         timestamp=None):
            if not self._private_key:
                raise ValueError("Private key is required for embedding metadata")
            if not self._signer_id:
                raise ValueError("Signer ID is required for embedding metadata")
            
            try:
                from encypher.core.unicode_metadata import UnicodeMetadata
                result = UnicodeMetadata.embed_metadata(
                    text=text,
                    private_key=self._private_key,
                    signer_id=self._signer_id,
                    metadata_format=metadata_format,
                    model_id=model_id,
                    custom_metadata=custom_metadata,
                    target=target,
                    timestamp=timestamp
                )
                if self.verbose:
                    console.print("[green]Successfully embedded metadata[/green]")
                return result
            except Exception as e:
                if self.verbose:
                    console.print(f"[red]Error embedding metadata: {e}[/red]")
                raise
        
        # Replace the method
        ea.embed_metadata = MethodType(embed_metadata_with_timestamp, ea)
        
        # Store the instance
        encypherai_instances[signer_id] = ea
    
    # Helper function to create a Word document with embedded metadata
    def create_docx_with_metadata(
        output_path: Path,
        title: str,
        content: str,
        metadata: Dict[str, Any],
        signer_id: str,
        model_id: str,
        timestamp: Optional[datetime] = None,
        tamper: bool = False
    ):
        """Create a Word document with embedded Encypher metadata."""
        if not DOCX_AVAILABLE:
            console.print("[yellow]Skipping DOCX creation - docx library not available[/yellow]")
            return None
        
        # Create the Word document
        doc = docx.Document()
        
        # Add title
        doc.add_heading(title, level=1)
        
        # Add content
        doc.add_paragraph(content)
        
        # Add metadata section
        doc.add_heading("Document Metadata:", level=2)
        
        # Create a table for metadata
        table = doc.add_table(rows=1, cols=2)
        table.style = 'Table Grid'
        
        # Add header row
        header_cells = table.rows[0].cells
        header_cells[0].text = "Key"
        header_cells[1].text = "Value"
        
        # Add metadata rows
        for key, value in metadata.items():
            row_cells = table.add_row().cells
            row_cells[0].text = key
            if isinstance(value, dict) or isinstance(value, list):
                row_cells[1].text = json.dumps(value, indent=2)
            else:
                row_cells[1].text = str(value)
        
        # Save the document
        doc.save(output_path)
        
        # Now read the Word document and embed metadata
        with open(output_path, 'rb') as f:
            docx_content = f.read()
        
        # Convert binary to Base64 and then to UTF-8 string for metadata embedding
        docx_base64 = base64.b64encode(docx_content).decode('utf-8')
        
        # Create a text wrapper with the DOCX content encoded
        text_wrapper = f"""DOCX_BASE64_START
{docx_base64}
DOCX_BASE64_END"""
        
        # Embed metadata in the text wrapper
        ea = encypherai_instances[signer_id]
        text_with_metadata = ea.embed_metadata(
            text=text_wrapper,
            custom_metadata=metadata,
            model_id=model_id,
            metadata_format="basic",
            timestamp=timestamp
        )
        
        # If tampering is requested, modify the content after embedding
        if tamper:
            lines = text_with_metadata.split('\n')
            if len(lines) > 2:
                # Insert tampering marker
                tamper_line = "THIS DOCX WAS TAMPERED WITH AFTER SIGNING"
                lines.insert(2, tamper_line)
                text_with_metadata = '\n'.join(lines)
        
        # Write the text with metadata back to the file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text_with_metadata)
        
        console.print(f"[green]Created Word document with {'tampered ' if tamper else ''}metadata: {output_path}[/green]")
        return output_path

    # Helper function to create a PDF with embedded metadata
    def create_pdf_with_metadata(
        output_path: Path,
        title: str,
        content: str,
        metadata: Dict[str, Any],
        signer_id: str,
        model_id: str,
        timestamp: Optional[datetime] = None,
        tamper: bool = False
    ):
        """Create a PDF file with embedded Encypher metadata."""
        # Create the PDF
        doc = SimpleDocTemplate(str(output_path), pagesize=letter)
        styles = getSampleStyleSheet()
        story = []
        
        # Add title
        title_style = ParagraphStyle(
            'Title',
            parent=styles['Heading1'],
            fontSize=16,
            spaceAfter=12
        )
        story.append(Paragraph(title, title_style))
        story.append(Spacer(1, 12))
        
        # Add content
        content_style = styles["Normal"]
        story.append(Paragraph(content, content_style))
        story.append(Spacer(1, 12))
        
        # Add metadata section
        story.append(Paragraph("Document Metadata:", styles["Heading2"]))
        story.append(Spacer(1, 6))
        
        # Create a table for metadata
        metadata_data = [["Key", "Value"]]
        for key, value in metadata.items():
            if isinstance(value, dict) or isinstance(value, list):
                value = json.dumps(value, indent=2)
            metadata_data.append([key, str(value)])
        
        metadata_table = Table(metadata_data, colWidths=[120, 350])
        metadata_table.setStyle(TableStyle([
            ('BACKGROUND', (0, 0), (1, 0), colors.grey),
            ('TEXTCOLOR', (0, 0), (1, 0), colors.whitesmoke),
            ('ALIGN', (0, 0), (1, 0), 'CENTER'),
            ('FONTNAME', (0, 0), (1, 0), 'Helvetica-Bold'),
            ('FONTSIZE', (0, 0), (1, 0), 12),
            ('BOTTOMPADDING', (0, 0), (1, 0), 12),
            ('BACKGROUND', (0, 1), (-1, -1), colors.beige),
            ('GRID', (0, 0), (-1, -1), 1, colors.black)
        ]))
        story.append(metadata_table)
        
        # Build the PDF
        doc.build(story)
        
        # Now read the PDF file and embed metadata
        with open(output_path, 'rb') as f:
            pdf_content = f.read()
        
        # Convert binary to Base64 and then to UTF-8 string for metadata embedding
        pdf_base64 = base64.b64encode(pdf_content).decode('utf-8')
        
        # Create a text wrapper with the PDF content encoded
        text_wrapper = f"""PDF_BASE64_START
{pdf_base64}
PDF_BASE64_END"""
        
        # Embed metadata in the text wrapper
        ea = encypherai_instances[signer_id]
        text_with_metadata = ea.embed_metadata(
            text=text_wrapper,
            custom_metadata=metadata,
            model_id=model_id,
            metadata_format="basic",
            timestamp=timestamp
        )
        
        # If tampering is requested, modify the content after embedding
        if tamper:
            lines = text_with_metadata.split('\n')
            if len(lines) > 2:
                # Insert tampering marker
                tamper_line = "THIS PDF WAS TAMPERED WITH AFTER SIGNING"
                lines.insert(2, tamper_line)
                text_with_metadata = '\n'.join(lines)
        
        # Write the text with metadata back to the file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text_with_metadata)
        
        console.print(f"[green]Created PDF with {'tampered ' if tamper else ''}metadata: {output_path}[/green]")
        return output_path
    
    # Helper function to create a text file with embedded metadata
    def create_text_file_with_metadata(
        output_path: Path,
        content: str,
        metadata: Dict[str, Any],
        signer_id: str,
        model_id: str,
        timestamp: Optional[datetime] = None,
        tamper: bool = False,
        tamper_metadata: bool = False
    ):
        """Create a text file with embedded Encypher metadata."""
        ea = encypherai_instances[signer_id]
        
        # If tampering metadata is requested, create a modified copy
        if tamper_metadata:
            # Create a copy with modified values
            tampered_metadata = metadata.copy()
            for key in tampered_metadata:
                if isinstance(tampered_metadata[key], str) and len(tampered_metadata[key]) > 3:
                    tampered_metadata[key] = f"TAMPERED-{tampered_metadata[key]}"
                elif isinstance(tampered_metadata[key], dict):
                    for subkey in tampered_metadata[key]:
                        if isinstance(tampered_metadata[key][subkey], str):
                            tampered_metadata[key][subkey] = f"TAMPERED-{tampered_metadata[key][subkey]}"
            metadata = tampered_metadata
        
        # Embed metadata
        text_with_metadata = ea.embed_metadata(
            text=content,
            custom_metadata=metadata,
            model_id=model_id,
            metadata_format="basic",
            timestamp=timestamp
        )
        
        # If content tampering is requested, modify after embedding
        if tamper:
            lines = text_with_metadata.split('\n')
            if len(lines) > 2:
                # Add tampering marker
                lines[1] = lines[1] + " [THIS TEXT WAS MODIFIED AFTER SIGNING]"
                text_with_metadata = '\n'.join(lines)
            else:
                # Fallback if we can't split by lines
                text_with_metadata = text_with_metadata.replace(
                    "This document", 
                    "This document [THIS TEXT WAS MODIFIED AFTER SIGNING]"
                )
        
        # Write to file
        with open(output_path, 'w', encoding='utf-8') as f:
            f.write(text_with_metadata)
        
        tampering_type = ""
        if tamper:
            tampering_type = "content-tampered "
        elif tamper_metadata:
            tampering_type = "metadata-tampered "
            
        console.print(f"[green]Created text file with {tampering_type}metadata: {output_path}[/green]")
        return output_path
    
    # Create example files for the dashboard
    with Progress(
        TextColumn("[bold blue]{task.description}"),
        BarColumn(),
        TaskProgressColumn(),
        console=console
    ) as progress:
        # Define document categories for the examples
        categories = [
            {
                "name": "Financial Reports",
                "models": ["financial-report-v1", "quarterly-statement-v1"],
                "signers": ["encypherai-official", "compliance-dept"]
            },
            {
                "name": "Research Documents",
                "models": ["research-paper-v1", "experiment-results-v1"],
                "signers": ["research-team", "encypherai-official"]
            },
            {
                "name": "Compliance Documents",
                "models": ["compliance-check-v1", "audit-report-v1"],
                "signers": ["compliance-dept"]
            },
            {
                "name": "External Documents",
                "models": ["external-report-v1"],
                "signers": ["untrusted-source"]
            }
        ]
        
        # Create a task for overall progress
        total_files = 20  # Adjust based on how many files you want to create
        task = progress.add_task("Creating example files...", total=total_files)
        
        file_count = 0
        
        # Create example files for each category
        for category in categories:
            category_dir = EXAMPLE_DIR / category["name"].lower().replace(" ", "_")
            category_dir.mkdir(exist_ok=True)
            
            # Determine how many files to create for this category
            num_files = max(2, total_files // len(categories))
            
            for i in range(num_files):
                # Randomly select model and signer
                model_id = random.choice(category["models"])
                signer_id = random.choice(category["signers"])
                
                # Generate a timestamp within the last 30 days
                days_ago = random.randint(0, 30)
                timestamp = datetime.now(timezone.utc) - timedelta(days=days_ago)
                
                # Decide if this should be a PDF, DOCX, or text file
                file_type = random.choice(["pdf", "docx", "txt"])
                
                # Decide if this file should be tampered with
                tamper = random.random() < 0.2  # 20% chance of tampering
                tamper_metadata = not tamper and random.random() < 0.2  # 20% chance of metadata tampering if not content tampered
                
                # Generate a unique filename
                tamper_suffix = "_tampered" if tamper or tamper_metadata else ""
                filename = f"{category['name'].lower().replace(' ', '_')}_{i+1}{tamper_suffix}.{file_type}"
                file_path = category_dir / filename
                
                # Generate metadata based on category
                metadata = {
                    "category": category["name"],
                    "document_id": f"doc-{category['name'].lower().replace(' ', '-')}-{i+1}",
                    "version": "1.0.0",
                    "created_by": f"Encypher {category['name']} Generator",
                    "department": category["name"],
                    "confidentiality": random.choice(["Public", "Internal", "Confidential", "Restricted"]),
                    "tags": [f"tag-{random.randint(1, 10)}" for _ in range(random.randint(1, 3))]
                }
                
                # Add category-specific metadata
                if "Financial" in category["name"]:
                    metadata.update({
                        "financial_period": f"Q{random.randint(1, 4)} {2024 + random.randint(0, 1)}",
                        "currency": random.choice(["USD", "EUR", "GBP"]),
                        "total_amount": round(random.uniform(10000, 1000000), 2)
                    })
                elif "Research" in category["name"]:
                    metadata.update({
                        "research_id": f"R-{random.randint(1000, 9999)}",
                        "principal_investigator": random.choice(["Dr. Smith", "Dr. Johnson", "Dr. Williams"]),
                        "status": random.choice(["Draft", "Under Review", "Published"])
                    })
                elif "Compliance" in category["name"]:
                    metadata.update({
                        "compliance_check_id": f"CC-{random.randint(1000, 9999)}",
                        "compliance_status": random.choice(["Compliant", "Non-Compliant", "Pending Review"]),
                        "reviewer": random.choice(["Alice Johnson", "Bob Smith", "Carol Williams"])
                    })
                elif "External" in category["name"]:
                    metadata.update({
                        "source": random.choice(["Partner A", "Vendor B", "Client C"]),
                        "verification_status": random.choice(["Verified", "Unverified", "Pending"]),
                        "external_reference": f"EXT-{random.randint(10000, 99999)}"
                    })
                
                # Generate content based on category
                if "Financial" in category["name"]:
                    title = f"Financial Report - {metadata['financial_period']}"
                    content = f"""This document contains financial information for {metadata['financial_period']}.
                    
Total Amount: {metadata['currency']} {metadata['total_amount']}

This report was generated by the Encypher Financial Reporting System and is intended for internal use only.
The information contained herein is confidential and should not be shared outside the organization.

Document ID: {metadata['document_id']}
Version: {metadata['version']}
Confidentiality: {metadata['confidentiality']}
"""
                elif "Research" in category["name"]:
                    title = f"Research Document - {metadata['research_id']}"
                    content = f"""Research Document: {metadata['research_id']}
Principal Investigator: {metadata['principal_investigator']}
Status: {metadata['status']}

This document contains research findings and is protected by intellectual property rights.
All data and conclusions are preliminary and subject to peer review.

Document ID: {metadata['document_id']}
Version: {metadata['version']}
Confidentiality: {metadata['confidentiality']}
"""
                elif "Compliance" in category["name"]:
                    title = f"Compliance Check - {metadata['compliance_check_id']}"
                    content = f"""Compliance Check ID: {metadata['compliance_check_id']}
Status: {metadata['compliance_status']}
Reviewer: {metadata['reviewer']}

This document contains the results of a compliance check performed by the Encypher Compliance Department.
The findings in this report should be addressed according to company policy.

Document ID: {metadata['document_id']}
Version: {metadata['version']}
Confidentiality: {metadata['confidentiality']}
"""
                else:  # External
                    title = f"External Document - {metadata['external_reference']}"
                    content = f"""External Document Reference: {metadata['external_reference']}
Source: {metadata['source']}
Verification Status: {metadata['verification_status']}

This document was received from an external source and has been processed by the Encypher system.
The content has not been verified for accuracy or completeness.

Document ID: {metadata['document_id']}
Version: {metadata['version']}
Confidentiality: {metadata['confidentiality']}
"""
                
                # Create the file with metadata
                if file_type == "pdf":
                    create_pdf_with_metadata(
                        output_path=file_path,
                        title=title,
                        content=content,
                        metadata=metadata,
                        signer_id=signer_id,
                        model_id=model_id,
                        timestamp=timestamp,
                        tamper=tamper
                    )
                elif file_type == "docx":
                    create_docx_with_metadata(
                        output_path=file_path,
                        title=title,
                        content=content,
                        metadata=metadata,
                        signer_id=signer_id,
                        model_id=model_id,
                        timestamp=timestamp,
                        tamper=tamper
                    )
                else:  # txt
                    create_text_file_with_metadata(
                        output_path=file_path,
                        content=f"{title}\n\n{content}",
                        metadata=metadata,
                        signer_id=signer_id,
                        model_id=model_id,
                        timestamp=timestamp,
                        tamper=tamper,
                        tamper_metadata=tamper_metadata
                    )
                
                # Update progress
                file_count += 1
                progress.update(task, completed=file_count)
        
        # Create a few files without metadata
        for i in range(3):
            category_name = random.choice([c["name"] for c in categories])
            category_dir = EXAMPLE_DIR / category_name.lower().replace(" ", "_")
            category_dir.mkdir(exist_ok=True)
            
            file_type = random.choice(["txt", "pdf"])
            filename = f"no_metadata_{i+1}.{file_type}"
            file_path = category_dir / filename
            
            if file_type == "pdf":
                # Create a PDF without metadata
                doc = SimpleDocTemplate(str(file_path), pagesize=letter)
                styles = getSampleStyleSheet()
                story = []
                
                title = f"Document Without Metadata - {i+1}"
                content = f"""This is an example document without any Encypher metadata.
                
This file should be detected as having no metadata when scanned by the verification tools.

File ID: no-metadata-{i+1}
"""
                
                story.append(Paragraph(title, styles["Heading1"]))
                story.append(Spacer(1, 12))
                story.append(Paragraph(content, styles["Normal"]))
                
                doc.build(story)
            else:
                # Create a text file without metadata
                content = f"""Document Without Metadata - {i+1}

This is an example document without any Encypher metadata.
                
This file should be detected as having no metadata when scanned by the verification tools.

File ID: no-metadata-{i+1}
"""
                with open(file_path, 'w', encoding='utf-8') as f:
                    f.write(content)
            
            console.print(f"[green]Created file without metadata: {file_path}[/green]")
            
            # Update progress
            file_count += 1
            progress.update(task, completed=file_count)
    
    # Create a README file with information about the examples
    readme_path = EXAMPLE_DIR / "README.md"
    with open(readme_path, 'w', encoding='utf-8') as f:
        f.write(f"""# Encypher Dashboard Example Files

This directory contains example files with Encypher metadata for testing and demonstration purposes.

## File Categories

{chr(10).join([f"- **{category['name']}**: Example {category['name'].lower()} with metadata signed by {', '.join([s for s in category['signers']])}" for category in categories])}

## Metadata Status

- **Valid Files**: Files with valid metadata and signatures
- **Tampered Content**: Files where the content was modified after metadata was embedded
- **Tampered Metadata**: Files where the metadata was tampered with
- **No Metadata**: Files without any Encypher metadata

## Trusted Signers

{chr(10).join([f"- **{signer['name']}** (`{signer['id']}`): Public key at `keys/{signer['id']}-public.pem`" for signer in signers])}

## Usage

These files can be used to test the Encypher verification tools and dashboard functionality.
You can use the audit-log-cli tool to scan these files:

```
uv run -- python -m app.main --target examples/files --trusted-signers examples/files/keys --output dashboard_example_report.csv
```

Generated on: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}
""")
    
    console.print("[bold green]Successfully created all example files![/bold green]")
    console.print("\nTo use these files with the audit-log-cli tool, run:")
    console.print("uv run -- python -m app.main --target dashboard_app/examples/files --trusted-signers dashboard_app/examples/files/keys --output dashboard_example_report.csv")

except ImportError as e:
    print(f"Error importing required modules: {e}")
    print("Make sure you have installed all the required dependencies.")
    sys.exit(1)
except Exception as e:
    print(f"Unexpected error: {e}")
    sys.exit(1)
